# -*- coding: utf-8 -*-
"""
Created on Fri Jul 19 14:00:38 2024

@author: EDY
"""


import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text

        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module07'
def heat_island_season_report(all_result,df_day, main_st_ids,data_types,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    doc_path=os.path.join(cfg['report']['template'],'Module07','heat_island_season.docx')

    doc=DocxTemplate(doc_path)
    
    data_types_str=data_types[0]
    main_st = [df_day[df_day['Station_Id_C'] == main_id]['Station_Name'][0] for main_id in main_st_ids]

    data_1=pd.DataFrame(all_result[data_types_str].season['温度累年各季变化'])
    data_2=pd.DataFrame(all_result[data_types_str].season['热岛强度累年各季变化'] )
    data_3=pd.DataFrame(all_result[data_types_str].season['热岛强度等级累年各季变化'] )
    
    dic=dict()
    dic['data_1']='、'.join(main_st)
   
    i_str=''
    for i_name in main_st:
        columns_1 = data_2.filter(regex=i_name[0:2:])
        if i_name==main_st[-1]:
            i_str=i_str+str(columns_1.iloc[-2,:].min())+'~'+str(columns_1.iloc[-1,:].max())+'℃'
        else:
            i_str=i_str+str(columns_1.iloc[-2,:].min())+'~'+str(columns_1.iloc[-1,:].max())+'℃、'

    dic['data_2']=i_str

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'heat_island_season.docx')
    doc.save(report)
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,data_1,'温度累年各季变化表')
    creat_table(document,data_2,'热岛强度累年各季变化表')
    creat_table(document,data_3,'热岛强度等级累年各季变化表')

    document.save(report)
    
    return report
