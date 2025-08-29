# -*- coding: utf-8 -*-
"""
Created on Mon Jul 29 17:28:09 2024

@author: EDY
"""

from docx import Document
import os
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from matplotlib.pyplot import MultipleLocator
from docxcompose.composer import Composer

def combine(new_docx_path,result_path,all_file_path):


    try:
        master = Document(all_file_path[0])
        middle_new_docx = Composer(master)
        for word in all_file_path[1:]:  # 从第二个文档开始追加
            word_document = Document(word)
            word_document.add_page_break()
            # 删除手动添加分页符的代码
            middle_new_docx.append(word_document)
        middle_new_docx.save(new_docx_path)
    except Exception as e:
        print(f"发生错误：{e}")
    
    document = Document(new_docx_path)
    
    for p in document.paragraphs:
        name= p.style.name
        print(name, )
        
    
    head1=0
    head2=0
    head3=0
    head4=0
    head5=0
    
    for para in document.paragraphs:
        style_name= para.style.name
        if style_name=="Heading 1":
            head1+=1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text, str(head1)+" "+para.text
                )
            head2=0
            head3=0
            head4=0
            head5=0            
            
        if style_name=="Heading 2":
            head2+=1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text, str(head1)+"."+str(head2)+" "+para.text
                )
                
            head3=0
            head4=0
            head5=0  
            
        if style_name=="Heading 3":
            head3+=1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text,
                    str(head1)+"."+str(head2)+"."+str(head3)+" "+para.text
                )
            head4=0
            head5=0  
        if style_name=="Heading 4":
            head4+=1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text,
                    str(head1)+"."+str(head2)+"."+str(head3)+" "+str(head4)+" "+para.text
                )
            head5=0                  
        if style_name=="Heading 5":
            head5+=1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(
                    para.text,
                    str(head1)+"."+str(head2)+"."+str(head3)+" "+str(head4)+" "+str(head5)+" "+para.text
                    )
    
    document.save(result_path)
    

    return result_path

if __name__=="__main__":
    
    new_docx_path = r"D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module_result\1.docx"
    result_path=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module_result\2.docx'
    original_docx_path = r"D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module_result\doc_test"
    
    all_word = [file for file in os.listdir(original_docx_path) if file.endswith('.docx')]

    all_file_path = [os.path.join(original_docx_path, file) for file in all_word]
    
    result_path=combine(new_docx_path,result_path,all_file_path)