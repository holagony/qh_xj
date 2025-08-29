# -*- coding: utf-8 -*-
"""
Created on Fri May 24 09:28:33 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            

    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module06'
def climate_livable_report(result,daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module06','climate_livable_report.docx')
    doc=DocxTemplate(doc_path)
    
    
    tables=result['tables']
    dic=dict()


    dic['start_year']=daily_df.index.year[0]
    dic['end_year']=daily_df.index.year[-1]
    dic['station_name']=daily_df['Station_Name'][0]
    
    #%% temp
    temp_1=pd.DataFrame(tables['历年适宜温度日数'])
    temp_2=pd.DataFrame(tables['历年7月平均最低气温'])
    temp_3=pd.DataFrame(tables['历年1月平均最高气温'])
    temp_4=pd.DataFrame(tables['历年平均气温日较差'])
    temp_5=pd.DataFrame(tables['历年夏季平均气温日较差'])
    temp_6=pd.DataFrame(tables['历年冬季平均气温日较差'])
    
    # 合并一下
    temp_1[temp_2.columns[1]]=temp_2[temp_2.columns[1]]
    temp_1[temp_3.columns[1]]=temp_3[temp_3.columns[1]]
    temp_1[temp_4.columns[1]]=temp_4[temp_4.columns[1]]
    temp_1[temp_5.columns[1]]=temp_5[temp_5.columns[1]]
    temp_1[temp_6.columns[1]]=temp_6[temp_6.columns[1]]

    try:
        dic['temp_1_mean']=round(temp_1['历年适宜温度日数'].mean(),2)
        dic['temp_1_min']=temp_1['历年适宜温度日数'].min()
        dic['temp_1_min_day']=temp_1['日期'][temp_1['历年适宜温度日数'].idxmin()]
        dic['temp_1_max']=temp_1['历年适宜温度日数'].max()
        dic['temp_1_miax_day']=temp_1['日期'][temp_1['历年适宜温度日数'].idxmax()]
    except:
        dic['temp_1_mean']=np.NAN
        dic['temp_1_min']=np.nan
        dic['temp_1_min_day']=np.nan
        dic['temp_1_max']=np.nan
        dic['temp_1_miax_day']=np.nan

    try:
        dic['temp_2_mean']=round(temp_1['历年7月平均最低气温'].mean(),2)
        dic['temp_2_min']=temp_1['历年7月平均最低气温'].min()
        dic['temp_2_min_day']=temp_1['日期'][temp_1['历年7月平均最低气温'].idxmin()]
        dic['temp_2_max']=temp_1['历年7月平均最低气温'].max()
        dic['temp_2_miax_day']=temp_1['日期'][temp_1['历年7月平均最低气温'].idxmax()]
    except:
        dic['temp_2_mean']=np.nan
        dic['temp_2_min']=np.nan
        dic['temp_2_min_day']=np.nan
        dic['temp_2_max']=np.nan
        dic['temp_2_miax_day']=np.nan
            
    try:
        dic['temp_3_mean']=round(temp_1['历年1月平均最高气温'].mean(),2)
        dic['temp_3_min']=temp_1['历年1月平均最高气温'].min()
        dic['temp_3_min_day']=temp_1['日期'][temp_1['历年1月平均最高气温'].idxmin()]
        dic['temp_3_max']=temp_1['历年1月平均最高气温'].max()
        dic['temp_3_miax_day']=temp_1['日期'][temp_1['历年1月平均最高气温'].idxmax()]
    except:
        dic['temp_3_mean']=np.nan
        dic['temp_3_min']=np.nan
        dic['temp_3_min_day']=np.nan
        dic['temp_3_max']=np.nan
        dic['temp_3_miax_day']=np.nan
    
    try:
        dic['temp_4_mean']=round(temp_1['历年平均气温日较差'].mean(),2)
        dic['temp_4_min']=temp_1['历年平均气温日较差'].min()
        dic['temp_4_min_day']=temp_1['日期'][temp_1['历年平均气温日较差'].idxmin()]
        dic['temp_4_max']=temp_1['历年平均气温日较差'].max()
        dic['temp_4_miax_day']=temp_1['日期'][temp_1['历年平均气温日较差'].idxmax()]
    except:
        dic['temp_4_mean']=np.nan
        dic['temp_4_min']=np.nan
        dic['temp_4_min_day']=np.nan
        dic['temp_4_max']=np.nan
        dic['temp_4_miax_day']=np.nan
        
        
    try:        
        dic['temp_5_mean']=round(temp_1['历年夏季平均气温日较差'].mean(),2)
        dic['temp_5_min']=temp_1['历年夏季平均气温日较差'].min()
        dic['temp_5_min_day']=temp_1['日期'][temp_1['历年夏季平均气温日较差'].idxmin()]
        dic['temp_5_max']=temp_1['历年夏季平均气温日较差'].max()
        dic['temp_5_miax_day']=temp_1['日期'][temp_1['历年夏季平均气温日较差'].idxmax()]
    except:
        dic['temp_5_mean']=np.nan
        dic['temp_5_min']=np.nan
        dic['temp_5_min_day']=np.nan
        dic['temp_5_max']=np.nan
        dic['temp_5_miax_day']=np.nan
      
    try:        

        dic['temp_6_mean']=round(temp_1['历年冬季平均气温日较差'].mean(),2)
        dic['temp_6_min']=temp_1['历年冬季平均气温日较差'].min()
        dic['temp_6_min_day']=temp_1['日期'][temp_1['历年冬季平均气温日较差'].idxmin()]
        dic['temp_6_max']=temp_1['历年冬季平均气温日较差'].max()
        dic['temp_6_miax_day']=temp_1['日期'][temp_1['历年冬季平均气温日较差'].idxmax()]
    except:
        dic['temp_6_mean']=np.nan
        dic['temp_6_min']=np.nan
        dic['temp_6_min_day']=np.nan
        dic['temp_6_max']=np.nan
        dic['temp_6_miax_day']=np.nan
    
    # fig
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年适宜温度日数','日数（天）','d','历年适宜温度日数.png',2,5,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年适宜温度日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年适宜温度日数'].min()-2,  temp_1['历年适宜温度日数'].max()+5)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年适宜温度日数.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_1_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
    except:
        dic['temp_1_picture'] =np.nan
    
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年7月平均最低气温','温度（℃）','℃','历年7月平均最低气温.png',2,2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年7月平均最低气温'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('温度（℃）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年7月平均最低气温'].min()-2,  temp_1['历年7月平均最低气温'].max()+1)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年7月平均最低气温.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_2_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['temp_2_picture'] =np.nan  
        
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年1月平均最高气温','温度（℃）','℃','历年1月平均最高气温.png',2,2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年1月平均最高气温'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('温度（℃）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年1月平均最高气温'].min(),  temp_1['历年1月平均最高气温'].max()+1)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年1月平均最高气温.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_3_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['temp_3_picture'] =np.nan 
        
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年平均气温日较差','温度（℃）','℃','历年平均气温日较差.png',2,2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年平均气温日较差'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('温度（℃）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年平均气温日较差'].min()-2,  temp_1['历年平均气温日较差'].max()+0.5)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年平均气温日较差.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_4_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['temp_4_picture'] =np.nan 
        
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年夏季平均气温日较差','温度（℃）','℃','历年夏季平均气温日较差.png',2,2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年夏季平均气温日较差'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('温度（℃）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年夏季平均气温日较差'].min()-2,  temp_1['历年夏季平均气温日较差'].max()+0.5)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年夏季平均气温日较差.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_5_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['temp_5_picture'] =np.nan 
        
    try:
        average_pre_picture_hournum=plot_picture(temp_1, '日期','历年冬季平均气温日较差','温度（℃）','℃','历年冬季平均气温日较差.png',2,2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(temp_1['日期'], temp_1['历年冬季平均气温日较差'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('温度（℃）')
        # plt.xticks(temp_1['日期'][::3])
        # plt.ylim( temp_1['历年冬季平均气温日较差'].min()-2,  temp_1['历年冬季平均气温日较差'].max()+0.5)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年冬季平均气温日较差.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['temp_6_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['temp_6_picture'] =np.nan 
    
    #%% 降水 
    pre_1=pd.DataFrame(tables['历年总降水量'])
    pre_2=pd.DataFrame(tables['历年降水变差系数'])
    pre_3=pd.DataFrame(tables['历年季节降水均匀度'])
    pre_4=pd.DataFrame(tables['历年适宜降水日数'])

    
    # 合并一下
    pre_1[pre_2.columns[1]]=pre_2[pre_2.columns[1]]
    pre_1[pre_3.columns[1]]=pre_3[pre_3.columns[1]]
    pre_1[pre_4.columns[1]]=pre_4[pre_4.columns[1]]


    try:
        dic['pre_1_mean']=round(pre_1['历年总降水量'].mean(),2)
        dic['pre_1_min']=pre_1['历年总降水量'].min()
        dic['pre_1_min_day']=pre_1['日期'][pre_1['历年总降水量'].idxmin()]
        dic['pre_1_max']=pre_1['历年总降水量'].max()
        dic['pre_1_miax_day']=pre_1['日期'][pre_1['历年总降水量'].idxmax()]
    except:
        dic['pre_1_mean']=np.nan
        dic['pre_1_min']=np.nan
        dic['pre_1_min_day']=np.nan
        dic['pre_1_max']=np.nan
        dic['pre_1_miax_day']=np.nan
        
    try:
        dic['pre_2_mean']=round(pre_1['历年降水变差系数'].mean(),2)
        dic['pre_2_min']=pre_1['历年降水变差系数'].min()
        dic['pre_2_min_day']=pre_1['日期'][pre_1['历年降水变差系数'].idxmin()]
        dic['pre_2_max']=pre_1['历年降水变差系数'].max()
        dic['pre_2_miax_day']=pre_1['日期'][pre_1['历年降水变差系数'].idxmax()]
    except:
        dic['pre_2_mean']=np.nan
        dic['pre_2_min']=np.nan
        dic['pre_2_min_day']=np.nan
        dic['pre_2_max']=np.nan
        dic['pre_2_miax_day']=np.nan
            
    try:
        dic['pre_3_mean']=round(pre_1['历年季节降水均匀度'].mean(),2)
        dic['pre_3_min']=pre_1['历年季节降水均匀度'].min()
        dic['pre_3_min_day']=pre_1['日期'][pre_1['历年季节降水均匀度'].idxmin()]
        dic['pre_3_max']=pre_1['历年季节降水均匀度'].max()
        dic['pre_3_miax_day']=pre_1['日期'][pre_1['历年季节降水均匀度'].idxmax()]
    except:
        dic['pre_3_mean']=np.nan
        dic['pre_3_min']=np.nan
        dic['pre_3_min_day']=np.nan
        dic['pre_3_max']=np.nan
        dic['pre_3_miax_day']=np.nan

    try:
            
        dic['pre_4_mean']=round(pre_1['历年适宜降水日数'].mean(),2)
        dic['pre_4_min']=pre_1['历年适宜降水日数'].min()
        dic['pre_4_min_day']=pre_1['日期'][pre_1['历年适宜降水日数'].idxmin()]
        dic['pre_4_max']=pre_1['历年适宜降水日数'].max()
        dic['pre_4_miax_day']=pre_1['日期'][pre_1['历年适宜降水日数'].idxmax()]
    except:
        dic['pre_4_mean']=np.nan
        dic['pre_4_min']=np.nan
        dic['pre_4_min_day']=np.nan
        dic['pre_4_max']=np.nan
        dic['pre_4_miax_day']=np.nan
    
    
    # fig
    try:
        average_pre_picture_hournum=plot_picture(pre_1, '日期','历年总降水量','降雨量（mm）','mm','历年总降水量.png',50,50,data_dir)
    
        # plt.figure(figsize=(10, 6))
        # plt.bar(pre_1['日期'], pre_1['历年总降水量'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('降雨量（mm）')
        # plt.xticks(pre_1['日期'][::3]) 
        # plt.ylim( pre_1['历年总降水量'].min()-50,  pre_1['历年总降水量'].max()+50)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年总降水量.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['pre_1_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    except:
        dic['pre_1_picture'] = np.nan
        
    try:
        average_pre_picture_hournum=plot_picture(pre_1, '日期','历年降水变差系数','降水变差系数','','历年降水变差系数.png',0.2,0.2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(pre_1['日期'], pre_1['历年降水变差系数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('降水变差系数')
        # plt.xticks(pre_1['日期'][::3])
        # plt.ylim( pre_1['历年降水变差系数'].min()-0.2,  pre_1['历年降水变差系数'].max()+0.2)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年季节降水均匀度.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['pre_2_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))    

    except:
        dic['pre_2_picture'] = np.nan 
        
    try:
        average_pre_picture_hournum=plot_picture(pre_1, '日期','历年季节降水均匀度','降水均匀度','','历年季节降水均匀度.png',0.03,0.03,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(pre_1['日期'], pre_1['历年季节降水均匀度'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('降水均匀度')
        # plt.xticks(pre_1['日期'][::3])
        # plt.ylim( pre_1['历年季节降水均匀度'].min(),  pre_1['历年季节降水均匀度'].max()+0.03)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年季节降水均匀度.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['pre_3_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))  

    except:
        dic['pre_3_picture'] = np.nan  
        
    try:
        average_pre_picture_hournum=plot_picture(pre_1, '日期','历年适宜降水日数','日数（天）','d','历年适宜降水日数.png',30,30,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(pre_1['日期'], pre_1['历年适宜降水日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')
        # plt.xticks(pre_1['日期'][::3])
        # plt.ylim( pre_1['历年适宜降水日数'].min()-30,  pre_1['历年适宜降水日数'].max()+10)
    
        # average_pre_picture_hournum=os.path.join(data_dir,'历年适宜降水日数.png')
        # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['pre_4_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))    

    except:
        dic['pre_4_picture'] = np.nan  
    #%% 相对湿度 
    rhu_1=pd.DataFrame(tables['历年平均相对湿度'])
    rhu_2=pd.DataFrame(tables['历年夏季平均相对湿度'])
    rhu_3=pd.DataFrame(tables['历年适宜湿度日数'])
    rhu_4=pd.DataFrame(tables['累年各月适宜湿度日数'])

    
    # 合并一下
    rhu_1[rhu_2.columns[1]]=rhu_2[rhu_2.columns[1]]
    rhu_1[rhu_3.columns[1]]=rhu_3[rhu_3.columns[1]]
    rhu_1[rhu_4.columns[1]]=rhu_4[rhu_4.columns[1]]


    try:
        dic['rhu_1_mean']=round(rhu_1['历年平均相对湿度'].mean(),2)
        dic['rhu_1_min']=rhu_1['历年平均相对湿度'].min()
        dic['rhu_1_min_day']=rhu_1['日期'][rhu_1['历年平均相对湿度'].idxmin()]
        dic['rhu_1_max']=rhu_1['历年平均相对湿度'].max()
        dic['rhu_1_miax_day']=rhu_1['日期'][rhu_1['历年平均相对湿度'].idxmax()]
    except:
        dic['rhu_1_mean']=np.nan
        dic['rhu_1_min']=np.nan
        dic['rhu_1_min_day']=np.nan
        dic['rhu_1_max']=np.nan
        dic['rhu_1_miax_day']=np.nan

    try:
        dic['rhu_2_mean']=round(rhu_1['历年夏季平均相对湿度'].mean(),2)
        dic['rhu_2_min']=rhu_1['历年夏季平均相对湿度'].min()
        dic['rhu_2_min_day']=rhu_1['日期'][rhu_1['历年夏季平均相对湿度'].idxmin()]
        dic['rhu_2_max']=rhu_1['历年夏季平均相对湿度'].max()
        dic['rhu_2_miax_day']=rhu_1['日期'][rhu_1['历年夏季平均相对湿度'].idxmax()]
    except:
        dic['rhu_2_mean']=np.nan
        dic['rhu_2_min']=np.nan
        dic['rhu_2_min_day']=np.nan
        dic['rhu_2_max']=np.nan
        dic['rhu_2_miax_day']=np.nan
    
    try:
        dic['rhu_3_mean']=round(rhu_1['历年适宜湿度日数'].mean(),2)
        dic['rhu_3_min']=rhu_1['历年适宜湿度日数'].min()
        dic['rhu_3_min_day']=rhu_1['日期'][rhu_1['历年适宜湿度日数'].idxmin()]
        dic['rhu_3_max']=rhu_1['历年适宜湿度日数'].max()
        dic['rhu_3_miax_day']=rhu_1['日期'][rhu_1['历年适宜湿度日数'].idxmax()]
    except:
        dic['rhu_3_mean']=np.nan
        dic['rhu_3_min']=np.nan
        dic['rhu_3_min_day']=np.nan
        dic['rhu_3_max']=np.nan
        dic['rhu_3_miax_day']=np.nan
        
    try:
        dic['rhu_4_mean']=round(rhu_1['累年各月适宜湿度日数'].mean(),2)
        dic['rhu_4_min']=rhu_1['累年各月适宜湿度日数'].min()
        dic['rhu_4_min_day']=rhu_1['日期'][rhu_1['累年各月适宜湿度日数'].idxmin()]
        dic['rhu_4_max']=rhu_1['累年各月适宜湿度日数'].max()
        dic['rhu_4_miax_day']=rhu_1['日期'][rhu_1['累年各月适宜湿度日数'].idxmax()]
    except:
        dic['rhu_4_mean']=np.nan
        dic['rhu_4_min']=np.nan
        dic['rhu_4_min_day']=np.nan
        dic['rhu_4_max']=np.nan
        dic['rhu_4_miax_day']=np.nan
    
    
    # fig
    try:
        average_rhu_picture_hournum=plot_picture(rhu_1, '日期','历年平均相对湿度','相对湿度（%）','%','历年平均相对湿度.png',10,10,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(rhu_1['日期'], rhu_1['历年平均相对湿度'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('相对湿度（%）')
        # plt.xticks(rhu_1['日期'][::3])
        # plt.ylim( rhu_1['历年平均相对湿度'].min()-10,  rhu_1['历年平均相对湿度'].max()+2)
    
        # average_rhu_picture_hournum=os.path.join(data_dir,'历年平均相对湿度.png')
        # plt.savefig(average_rhu_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['rhu_1_picture'] = InlineImage(doc, average_rhu_picture_hournum, width=Mm(130))
    except:
        dic['rhu_1_picture'] = np.nan
        
    try:
        average_rhu_picture_hournum=plot_picture(rhu_1, '日期','历年夏季平均相对湿度','相对湿度（%）','%','历年夏季平均相对湿度.png',10,10,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(rhu_1['日期'], rhu_1['历年夏季平均相对湿度'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('相对湿度（%）')
        # plt.xticks(rhu_1['日期'][::3])
        # plt.ylim( rhu_1['历年夏季平均相对湿度'].min()-10,  rhu_1['历年夏季平均相对湿度'].max()+2)
    
        # average_rhu_picture_hournum=os.path.join(data_dir,'历年夏季平均相对湿度.png')
        # plt.savefig(average_rhu_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['rhu_2_picture'] = InlineImage(doc, average_rhu_picture_hournum, width=Mm(130))  
    except:
        dic['rhu_2_picture'] = np.nan
    
    try:
        average_rhu_picture_hournum=plot_picture(rhu_1, '日期','历年适宜湿度日数','日数(天)','d','历年适宜湿度日数.png',50,20,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(rhu_1['日期'], rhu_1['历年适宜湿度日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数(天)')
        # plt.xticks(rhu_1['日期'][::3])
        # plt.ylim( rhu_1['历年适宜湿度日数'].min()-50,  rhu_1['历年适宜湿度日数'].max()+20)
    
        # average_rhu_picture_hournum=os.path.join(data_dir,'历年适宜湿度日数.png')
        # plt.savefig(average_rhu_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['rhu_3_picture'] = InlineImage(doc, average_rhu_picture_hournum, width=Mm(130))  
    except:
        dic['rhu_3_picture'] = np.nan    
        
    try:
        average_rhu_picture_hournum=plot_picture(rhu_1, '日期','累年各月适宜湿度日数','日数(天)','d','累年各月适宜湿度日数.png',10,20,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(rhu_1['日期'], rhu_1['累年各月适宜湿度日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')  
        # plt.xticks(rhu_1['日期'][::3])
        # plt.ylim( rhu_1['累年各月适宜湿度日数'].min()-5,  rhu_1['累年各月适宜湿度日数'].max()+20)
    
        # average_rhu_picture_hournum=os.path.join(data_dir,'累年各月适宜湿度日数.png')
        # plt.savefig(average_rhu_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['rhu_4_picture'] = InlineImage(doc, average_rhu_picture_hournum, width=Mm(130))    
    except:
        dic['rhu_4_picture'] = np.nan
    #%% 风
    wind_1=pd.DataFrame(tables['历年平均风速'])
    wind_2=pd.DataFrame(tables['历年适宜风日数'])
    wind_3=pd.DataFrame(tables['累年各月适宜风日数'])

    
    # 合并一下
    wind_1[wind_2.columns[1]]=wind_2[wind_2.columns[1]]
    wind_1[wind_3.columns[1]]=wind_3[wind_3.columns[1]]


    try:
        dic['wind_1_mean']=round(wind_1['历年平均风速'].mean(),2)
        dic['wind_1_min']=wind_1['历年平均风速'].min()
        dic['wind_1_min_day']=wind_1['日期'][wind_1['历年平均风速'].idxmin()]
        dic['wind_1_max']=wind_1['历年平均风速'].max()
        dic['wind_1_miax_day']=wind_1['日期'][wind_1['历年平均风速'].idxmax()]
    except:
        dic['wind_1_mean']=np.nan
        dic['wind_1_min']=np.nan
        dic['wind_1_min_day']=np.nan
        dic['wind_1_max']=np.nan
        dic['wind_1_miax_day']=np.nan   

    try:
        dic['wind_2_mean']=round(wind_1['历年适宜风日数'].mean(),2)
        dic['wind_2_min']=wind_1['历年适宜风日数'].min()
        dic['wind_2_min_day']=wind_1['日期'][wind_1['历年适宜风日数'].idxmin()]
        dic['wind_2_max']=wind_1['历年适宜风日数'].max()
        dic['wind_2_miax_day']=wind_1['日期'][wind_1['历年适宜风日数'].idxmax()]
    except:
        dic['wind_2_mean']=np.nan
        dic['wind_2_min']=np.nan
        dic['wind_2_min_day']=np.nan
        dic['wind_2_max']=np.nan
        dic['wind_2_miax_day']=np.nan   
        
    try:
        dic['wind_3_mean']=round(wind_1['累年各月适宜风日数'].mean(),2)
        dic['wind_3_min']=wind_1['累年各月适宜风日数'].min()
        dic['wind_3_min_day']=wind_1['日期'][wind_1['累年各月适宜风日数'].idxmin()]
        dic['wind_3_max']=wind_1['累年各月适宜风日数'].max()
        dic['wind_3_miax_day']=wind_1['日期'][wind_1['累年各月适宜风日数'].idxmax()]
    except:
        dic['wind_3_mean']=np.nan
        dic['wind_3_min']=np.nan
        dic['wind_3_min_day']=np.nan
        dic['wind_3_max']=np.nan
        dic['wind_3_miax_day']=np.nan      
    
    # fig
    try:
        average_wind_picture_hournum=plot_picture(wind_1, '日期','历年平均风速','风速（m/s）','m/s','历年平均风速.png',1,1,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(wind_1['日期'], wind_1['历年平均风速'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('风速（m/s）')
        # plt.xticks(wind_1['日期'][::3])
        # plt.ylim(0,  wind_1['历年平均风速'].max()+0.5)
    
        # average_wind_picture_hournum=os.path.join(data_dir,'历年平均风速.png')
        # plt.savefig(average_wind_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['wind_1_picture'] = InlineImage(doc, average_wind_picture_hournum, width=Mm(130))
    except:
        dic['wind_1_picture'] = np.nan
    
    try:
        average_wind_picture_hournum=plot_picture(wind_1, '日期','历年适宜风日数','日数（天）','d','历年适宜风日数.png',10,5,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(wind_1['日期'], wind_1['历年适宜风日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')
        # plt.xticks(wind_1['日期'][::3])
        # plt.ylim( wind_1['历年适宜风日数'].min()-10,  wind_1['历年适宜风日数'].max()+2)
    
        # average_wind_picture_hournum=os.path.join(data_dir,'历年适宜风日数.png')
        # plt.savefig(average_wind_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['wind_2_picture'] = InlineImage(doc, average_wind_picture_hournum, width=Mm(130))    
    except:
        dic['wind_2_picture'] = np.nan  
        
    try:
        average_wind_picture_hournum=plot_picture(wind_1, '日期','累年各月适宜风日数','日数（天）','d','累年各月适宜风日数.png',10,5,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(wind_1['日期'], wind_1['累年各月适宜风日数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')
        # plt.xticks(wind_1['日期'][::3])
        # plt.ylim( wind_1['累年各月适宜风日数'].min()-10,  wind_1['累年各月适宜风日数'].max()+2)
    
        # average_wind_picture_hournum=os.path.join(data_dir,'累年各月适宜风日数.png')
        # plt.savefig(average_wind_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['wind_3_picture'] = InlineImage(doc, average_wind_picture_hournum, width=Mm(130))  
    except:
        dic['wind_3_picture'] = np.nan   
    
    #%% 日照
    sun_1=pd.DataFrame(tables['历年夏季日照时数'])
    sun_2=pd.DataFrame(tables['历年冬季日照时数'])

    
    # 合并一下
    sun_1[sun_2.columns[1]]=sun_2[sun_2.columns[1]]

    try:
        dic['sun_1_mean']=round(sun_1['历年夏季日照时数'].mean(),2)
        dic['sun_1_min']=sun_1['历年夏季日照时数'].min()
        dic['sun_1_min_day']=sun_1['日期'][sun_1['历年夏季日照时数'].idxmin()]
        dic['sun_1_max']=sun_1['历年夏季日照时数'].max()
        dic['sun_1_miax_day']=sun_1['日期'][sun_1['历年夏季日照时数'].idxmax()]
    except:
        dic['sun_1_mean']=np.nan
        dic['sun_1_min']=np.nan
        dic['sun_1_min_day']=np.nan
        dic['sun_1_max']=np.nan
        dic['sun_1_miax_day']=np.nan        

    try:
        dic['sun_2_mean']=round(sun_1['历年冬季日照时数'].mean(),2)
        dic['sun_2_min']=sun_1['历年冬季日照时数'].min()
        dic['sun_2_min_day']=sun_1['日期'][sun_1['历年冬季日照时数'].idxmin()]
        dic['sun_2_max']=sun_1['历年冬季日照时数'].max()
        dic['sun_2_miax_day']=sun_1['日期'][sun_1['历年冬季日照时数'].idxmax()]
    except:
        dic['sun_2_mean']=np.nan
        dic['sun_2_min']=np.nan
        dic['sun_2_min_day']=np.nan
        dic['sun_2_max']=np.nan
        dic['sun_2_miax_day']=np.nan  
       
    # fig
    try:
        average_sun_picture_hournum=plot_picture(sun_1, '日期','历年夏季日照时数','日照时数（h）','h','历年夏季日照时数.png',10,10,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(sun_1['日期'], sun_1['历年夏季日照时数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日照时数（h）')
        # plt.xticks(sun_1['日期'][::3])
        # plt.ylim(sun_1['历年夏季日照时数'].min()-10,  sun_1['历年夏季日照时数'].max()+10)
    
        # average_sun_picture_hournum=os.path.join(data_dir,'历年夏季日照时数.png')
        # plt.savefig(average_sun_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['sun_1_picture'] = InlineImage(doc, average_sun_picture_hournum, width=Mm(130))
    except:
        dic['sun_1_picture'] = np.nan 
    
    try:
        average_sun_picture_hournum=plot_picture(sun_1, '日期','历年冬季日照时数','日照时数（h）','h','历年冬季日照时数.png',10,10,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(sun_1['日期'], sun_1['历年冬季日照时数'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日照时数（h）')
        # plt.xticks(sun_1['日期'][::3])
        # plt.ylim(sun_1['历年冬季日照时数'].min()-10,  sun_1['历年冬季日照时数'].max()+10)
    
        # average_sun_picture_hournum=os.path.join(data_dir,'历年冬季日照时数.png')
        # plt.savefig(average_sun_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['sun_2_picture'] = InlineImage(doc, average_sun_picture_hournum, width=Mm(130))    
    except:
        dic['sun_2_picture'] = np.nan 
    #%% 气候季节
    season_1=pd.DataFrame(tables['历年大气含氧量'])
    season_2=pd.DataFrame(tables['历年春秋季总长'])

    
    # 合并一下
    season_1[season_2.columns[1]]=season_2[season_2.columns[1]]

    try:
        dic['season_1_mean']=round(season_1['历年大气含氧量'].mean(),2)
        dic['season_1_min']=season_1['历年大气含氧量'].min()
        dic['season_1_min_day']=season_1['日期'][season_1['历年大气含氧量'].idxmin()]
        dic['season_1_max']=season_1['历年大气含氧量'].max()
        dic['season_1_miax_day']=season_1['日期'][season_1['历年大气含氧量'].idxmax()]
    except:
        dic['season_1_mean']=np.nan
        dic['season_1_min']=np.nan
        dic['season_1_min_day']=np.nan
        dic['season_1_max']=np.nan
        dic['season_1_miax_day']=np.nan      

    try:
        dic['season_2_mean']=round(season_1['历年春秋季总长'].mean(),2)
        dic['season_2_min']=season_1['历年春秋季总长'].min()
        dic['season_2_min_day']=season_1['日期'][season_1['历年春秋季总长'].idxmin()]
        dic['season_2_max']=season_1['历年春秋季总长'].max()
        dic['season_2_miax_day']=season_1['日期'][season_1['历年春秋季总长'].idxmax()]
    except:
        dic['season_2_mean']=np.nan
        dic['season_2_min']=np.nan
        dic['season_2_min_day']=np.nan
        dic['season_2_max']=np.nan
        dic['season_2_miax_day']=np.nan    
       
    # fig
    try:
        average_season_picture_hournum=plot_picture(season_1, '日期','历年大气含氧量','大气含氧量（%）','%','历年大气含氧量.png',0.5,0.2,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(season_1['日期'], season_1['历年大气含氧量'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('大气含氧量（%）')
        # plt.xticks(season_1['日期'][::3])
        # plt.ylim(season_1['历年大气含氧量'].min()-0.5,  season_1['历年大气含氧量'].max()+0.2)
    
        # average_season_picture_hournum=os.path.join(data_dir,'历年大气含氧量.png')
        # plt.savefig(average_season_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['season_1_picture'] = InlineImage(doc, average_season_picture_hournum, width=Mm(130))
    except:
        dic['season_1_picture'] = np.nan 
    
    try:
        average_season_picture_hournum=plot_picture(season_1, '日期','历年春秋季总长','日数（天）','d','历年春秋季总长.png',10,10,data_dir)

        # plt.figure(figsize=(10, 6))
        # plt.bar(season_1['日期'], season_1['历年春秋季总长'], width=0.4, color='skyblue')
        # plt.grid(axis='y', linestyle='--', alpha=0.7)    
        # plt.xlabel('年')
        # plt.ylabel('日数（天）')
        # plt.xticks(season_1['日期'][::3])
        # plt.ylim(season_1['历年春秋季总长'].min()-10,  season_1['历年春秋季总长'].max()+10)
    
        # average_season_picture_hournum=os.path.join(data_dir,'历年春秋季总长.png')
        # plt.savefig(average_season_picture_hournum, bbox_inches='tight', dpi=200)
        # plt.clf()
        # plt.close('all')
        dic['season_2_picture'] = InlineImage(doc, average_season_picture_hournum, width=Mm(130))         
    except:
        dic['season_2_picture'] = np.nan  
   
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'climate_livable_report.docx')
    doc.save(report)
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,temp_1,'历年各温度要素统计表')
    creat_table(document,pre_1,'历年各降水要素统计表')
    creat_table(document,rhu_1,'历年各相对湿度要素统计表')
    creat_table(document,wind_1,'历年各风要素统计表')
    creat_table(document,sun_1,'历年各日照要素统计表')
    creat_table(document,season_1,'历年各气候季节要素统计表')

    document.save(report)

    
    return report