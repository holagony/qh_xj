# -*- coding: utf-8 -*-
"""
Created on Tue Jul  9 15:34:18 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            

    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)

def table_make(table):
    columns_1=table.columns[0]
    columns_2=table.columns[1]

    half_length =int(np.ceil(len(table)/2))
    table['A1'] = table.iloc[:half_length:,0]
    table['B1'] =table.iloc[:half_length:,1]
    table['A2'] = table.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    table['B2'] =table.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    table.drop([columns_1, columns_2], axis=1, inplace=True)
    table.columns=[columns_1, columns_2,columns_1, columns_2]
    table.dropna(how='all',inplace=True)

    
    return table
    
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module06'
def climate_comfort_report(result,daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module06','climate_comfort_report.docx')
    doc=DocxTemplate(doc_path)
    
    
    tables=result['tables']
    dic=dict()


    dic['start_year']=daily_df.index.year[0]
    dic['end_year']=daily_df.index.year[-1]
    dic['station_name']=daily_df['Station_Name'][0]
    
    #%% 人体舒适度
    descrip_1=dict()
    descrip_1['1']='寒冷, 不舒适'
    descrip_1['2']='较冷, 大部分人不舒适'
    descrip_1['3']='清凉, 少部分人不舒适'
    descrip_1['4']='偏凉, 大部分人舒适'
    descrip_1['5']='最为舒适'
    descrip_1['6']='偏热, 大部分人舒适'
    descrip_1['7']='闷热, 少部分人不舒适'
    descrip_1['8']='炎热, 大部分人不舒适'
    descrip_1['9']='暑热, 不舒适'
    descrip_1['10']='酷热,很不舒适'

    data_1=result['tables']['人体舒适度指数']['规范统计']
    data_12=pd.DataFrame(data_1['历年舒适及以上(等级4/5/6)的月数'])
    data_13=pd.DataFrame(data_1['历年平均人体舒适指数等级'])
    data_14=pd.DataFrame(data_1['累年各月平均人体舒适指数等级'])
    
    dic['data_11']=data_13['平均人体舒适度等级'].mode()[0]
    dic['data_12']=data_12['舒适及以上等级的月数'].mode()[0]
    dic['data_13']=data_12['舒适及以上等级的月数'].min()
    dic['data_14']=data_12['年'][data_12['舒适及以上等级的月数'].idxmin()]

    dic['data_15']='、'.join(data_14['日期'][data_14['累年各月平均人体舒适指数等级']==data_14['累年各月平均人体舒适指数等级'].max()])
    dic['data_16']=int(data_14['累年各月平均人体舒适指数等级'].max())
    dic['data_17']=descrip_1[str(dic['data_16'])]
    
    dic['data_18']='、'.join(data_14['日期'][data_14['累年各月平均人体舒适指数等级']==data_14['累年各月平均人体舒适指数等级'].min()])
    dic['data_19']=int(data_14['累年各月平均人体舒适指数等级'].min())
    dic['data_20']=descrip_1[str(dic['data_19'])]
    
    #%% 气候度假
    descrip_2=dict()
    descrip_2['1']='特别不适宜'
    descrip_2['2']='很不适宜'
    descrip_2['3']='不适宜'
    descrip_2['4']='一般'
    descrip_2['5']='可以接受'
    descrip_2['6']='适宜'
    descrip_2['7']='很适宜'
    descrip_2['8']='特别适宜'
    descrip_2['9']='理想状态'
    
    data_2=result['tables']['气候度假指数']['规范统计']
    data_21=pd.DataFrame(result['tables']['气候度假指数']['规范统计']['历年平均气候度假指数等级'])
    data_22=pd.DataFrame(result['tables']['气候度假指数']['规范统计']['累年各月平均气候度假指数等级'])

    dic['data_21']=int(data_21['平均气候度假指数'].mode()[0])
    dic['data_22']=descrip_2[str(dic['data_21'])]
    dic['data_23']='、'.join(data_22['日期'][data_22['累年各月平均气候度假指数等级']==data_22['累年各月平均气候度假指数等级'].max()])
    dic['data_24']=int(data_22['累年各月平均气候度假指数等级'].max())
    dic['data_25']=descrip_2[str(dic['data_24'])]
    
    #%% 气候旅游
    descrip_3=dict()
    descrip_3['0']='极度不舒适'
    descrip_3['1']='非常不舒适'
    descrip_3['2']='不太舒适'
    descrip_3['3']='不舒适'
    descrip_3['4']='一般舒适'
    descrip_3['5']='较舒适'
    descrip_3['6']='舒适'
    descrip_3['7']='很舒适'
    descrip_3['8']='非常舒适'
    descrip_3['9']='特别舒适'


    data_3=result['tables']['气候旅游指数']['规范统计']
    data_31=pd.DataFrame(data_3['历年平均气候旅游指数等级'])
    data_32=pd.DataFrame(data_3['累年各月平均气候旅游指数等级'])

    dic['data_26']=int(data_31['平均气候旅游指数'].mode()[0])
    dic['data_27']=descrip_3[str(dic['data_26'])]
    dic['data_28']='、'.join(data_32['日期'][data_32['累年各月平均气候旅游指数等级']==data_32['累年各月平均气候旅游指数等级'].max()])
    dic['data_29']=int(data_32['累年各月平均气候旅游指数等级'].max())
    dic['data_30']=descrip_3[str(dic['data_29'])]
    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'climate_comfort_report.docx')
    doc.save(report)
    
    document = Document(report)
    

    data_13_new=table_make(data_13)
    data_12_new=table_make(data_12)
    data_14_new=table_make(data_14)
    
    data_21_new=table_make(data_21)
    data_22_new=table_make(data_22)

    data_31_new=table_make(data_31)
    data_32_new=table_make(data_32)    
    
    
    # 填充表格数据
    creat_table(document,data_13_new,'历年人体舒适度等级')
    creat_table(document,data_12_new,'历年舒适及以上等级的月数')
    creat_table(document,data_14_new,'累年各月平均人体舒适指数等级')
    creat_table(document,data_21_new,'历年平均气候度假指数')
    creat_table(document,data_22_new,'累年各月平均气候度假指数等级')
    creat_table(document,data_31_new,'历年平均气候旅游指数')
    creat_table(document,data_32_new,'累年各月平均气候旅游指数等级')

    document.save(report)

    plt.cla()
    plt.close('all')
    
    return report