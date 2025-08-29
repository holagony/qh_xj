# -*- coding: utf-8 -*-
"""
Created on Tue Jul  9 17:51:18 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)

def table_make(dic,data_11,data_12,num,ele,data_dir,doc):
    
    name_1='历年'+ele+'日数'
    name_2='累年各月平均'+ele+'日数'
    name_3=ele+'日数'
    num_str='data_'+num
    
    dic[num_str+'_1'] = round(data_11[name_1].mean(),2)
    dic[num_str+'_2'] = data_11[name_1].min()
    dic[num_str+'_3'] = data_11[name_1].max()
    
    # years=data_11['日期'][:]
    mask = ~np.isnan(data_11[name_1])
    valid_years =data_11['日期'][mask]
    valid_preperatures = data_11[name_1][mask]
    slope, intercept = np.polyfit(valid_years, valid_preperatures, 1)   
    if slope> 0:
        dic[num_str+'_4']='上升'
    else:
        dic[num_str+'_4']='下降'
      
    average_pre_picture_hournum=plot_picture(data_11, '日期',name_1,name_3+'（d）','d','历年平均'+name_3+'变化.png',100,40,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(data_11['日期'],data_11[name_1], width=0.4, color='skyblue', label=name_3)
    # plt.plot(years, slope * years + intercept, color='red', label='线性（'+name_3+'）')
    
    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('年')
    # plt.ylabel(name_3+'（d）')
    # plt.xticks(years[::3])
    # # plt.ylim(dic['average_pre_1'], dic['average_pre_2']+10)
    # plt.legend()
    
    # average_pre_picture_hournum=os.path.join(data_dir,'历年平均'+name_3+'变化.png')
    # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
    # 
    # plt.clf()
    # plt.close('all')
    dic[num_str+'_5'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
    dic[num_str+'_7'] = data_12[name_2].min()
    dic[num_str+'_8'] = data_12[name_2].max()
    
    # 计算春夏秋冬
    spring = data_12.iloc[2:5,1] 
    summer = data_12.iloc[5:8,1]  
    autumn = data_12.iloc[8:11,1] 
    winter = pd.concat([data_12.iloc[11:12,1],data_12.iloc[0:2,1]])  
    
    average_spring = sum(spring) / len(spring)
    average_summer = sum(summer) / len(summer)
    average_autumn = sum(autumn) / len(autumn)
    average_winter = sum(winter) / len(winter)
    
    seasons = {
    "春季": average_spring,
    "夏季": average_summer,
    "秋季": average_autumn,
    "冬季": average_winter}
    
    sorted_seasons = sorted(seasons.items(), key=lambda x: x[1], reverse=True)
    
    dic[num_str+'_9'] = sorted_seasons[0][0]  
    dic[num_str+'_20'] = sorted_seasons[-1][0]  
    
    # 图像绘制
    months=data_12['日期']
    average_pre_picture_month=plot_picture_2(months,data_12[name_2],dic,'平均'+name_3+'（d）','average_pre_m1','average_pre_m2','平均'+name_3+'逐月变化.png',100,10,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(months, data_12[name_2], width=0.4, color='skyblue')
    
    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('月')
    # plt.ylabel('平均'+name_3+'（d）')
    # plt.xticks(months)
    
    # average_pre_picture_month=os.path.join(data_dir,'平均'+name_3+'逐月变化.png')
    # plt.savefig(average_pre_picture_month, bbox_inches='tight', dpi=200)
    # plt.clf()
    # plt.close('all')
    
    dic[num_str+'_21'] = InlineImage(doc, average_pre_picture_month, width=Mm(130))

    return dic

def table_make_2(table,num):
    columns_1=table.columns[0]
    columns_2=table.columns[1]
    if num==1:
        table[columns_1]=table[columns_1].astype(str)
    half_length =int(np.ceil(len(table)/2))
    table['A1'] = table.iloc[:half_length:,0]
    table['B1'] =table.iloc[:half_length:,1]
    table['A2'] = table.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    table['B2'] =table.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    table.drop([columns_1, columns_2], axis=1, inplace=True)
    table.columns=[columns_1, columns_2,columns_1, columns_2]
    table.dropna(how='all',inplace=True)

    
    return table

# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module06'
def climate_disadvantage_report(result,daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module06','climate_disadvantage_report.docx')
    doc=DocxTemplate(doc_path)
    
    
    tables=result['tables']
    dic=dict()


    dic['start_year']=daily_df.index.year[0]
    dic['end_year']=daily_df.index.year[-1]
    dic['station_name']=daily_df['Station_Name'][0]
    
    data_11=pd.DataFrame(tables['历年高温日数'])
    data_12=pd.DataFrame(tables['累年各月平均高温日数'])
    data_21=pd.DataFrame(tables['历年寒冷日数'])
    data_22=pd.DataFrame(tables['累年各月平均寒冷日数'])
    data_31=pd.DataFrame(tables['历年大雨日数'])
    data_32=pd.DataFrame(tables['累年各月平均大雨日数'])
    data_41=pd.DataFrame(tables['历年无雨日数'])
    data_42=pd.DataFrame(tables['累年各月平均无雨日数'])
    data_51=pd.DataFrame(tables['历年强风日数'])
    data_52=pd.DataFrame(tables['累年各月平均强风日数'])
    data_61=pd.DataFrame(tables['历年静风日数'])
    data_62=pd.DataFrame(tables['累年各月平均静风日数'])
    data_71=pd.DataFrame(tables['历年沙尘日数'])
    data_72=pd.DataFrame(tables['累年各月平均沙尘日数'])
    data_81=pd.DataFrame(tables['历年霾日数'])
    data_82=pd.DataFrame(tables['累年各月平均霾日数'])
    data_91=pd.DataFrame(tables['历年强对流日数'])
    data_92=pd.DataFrame(tables['累年各月平均强对流日数'])

    
    dic=table_make(dic,data_11,data_12,'1','高温',data_dir,doc)
    dic=table_make(dic,data_21,data_22,'2','寒冷',data_dir,doc)
    dic=table_make(dic,data_31,data_32,'3','大雨',data_dir,doc)
    dic=table_make(dic,data_41,data_42,'4','无雨',data_dir,doc)
    dic=table_make(dic,data_51,data_52,'5','强风',data_dir,doc)
    dic=table_make(dic,data_61,data_62,'6','静风',data_dir,doc)
    dic=table_make(dic,data_71,data_72,'7','沙尘',data_dir,doc)
    dic=table_make(dic,data_81,data_82,'8','霾',data_dir,doc)
    dic=table_make(dic,data_91,data_92,'9','强对流',data_dir,doc)


    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'climate_disadvantage_report.docx')
    doc.save(report)
    
    document = Document(report)
    
    data_11_new=table_make_2(data_11,1)
    data_12_new=table_make_2(data_12,0)
    data_21_new=table_make_2(data_21,1)
    data_22_new=table_make_2(data_22,0)
    data_31_new=table_make_2(data_31,1)
    data_32_new=table_make_2(data_32,0)
    data_41_new=table_make_2(data_41,1)
    data_42_new=table_make_2(data_42,0)
    data_51_new=table_make_2(data_51,1)
    data_52_new=table_make_2(data_52,0)
    data_61_new=table_make_2(data_61,1)
    data_62_new=table_make_2(data_62,0)
    data_71_new=table_make_2(data_71,1)
    data_72_new=table_make_2(data_72,0)
    data_81_new=table_make_2(data_81,1)
    data_82_new=table_make_2(data_82,0)
    data_91_new=table_make_2(data_91,1)
    data_92_new=table_make_2(data_92,0)    
    
    # 填充表格数据
    creat_table(document,data_11_new,'高温日数年变化（单位：d）')
    creat_table(document,data_12_new,'高温日数月变化（单位：d）')
    creat_table(document,data_21_new,'寒冷日数年变化（单位：d）')
    creat_table(document,data_22_new,'寒冷日数月变化（单位：d）')
    creat_table(document,data_31_new,'大雨日数年变化（单位：d）')
    creat_table(document,data_32_new,'大雨日数月变化（单位：d）')
    creat_table(document,data_41_new,'无雨日数年变化（单位：d）')
    creat_table(document,data_42_new,'无雨日数月变化（单位：d）')
    creat_table(document,data_51_new,'强风日数年变化（单位：d）')
    creat_table(document,data_52_new,'强风日数月变化（单位：d）')
    creat_table(document,data_61_new,'静风日数年变化（单位：d）')
    creat_table(document,data_62_new,'静风日数月变化（单位：d）')
    creat_table(document,data_71_new,'沙尘日数年变化（单位：d）')
    creat_table(document,data_72_new,'沙尘日数月变化（单位：d）')
    creat_table(document,data_81_new,'霾日数年变化（单位：d）')
    creat_table(document,data_82_new,'霾日数月变化（单位：d）')
    creat_table(document,data_91_new,'强对流日数年变化（单位：d）')
    creat_table(document,data_92_new,'强对流日数月变化（单位：d）')

    document.save(report)
    
    return report