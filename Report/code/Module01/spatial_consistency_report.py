# -*- coding: utf-8 -*-
"""
Created on Mon Jun 17 15:20:43 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docxcompose.composer import Composer

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)
            

            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)

unit = {'平均气压': 'hPa', 
                '最高气压': 'hPa', 
                '最低气压': 'hPa', 
                '平均气温': '℃', 
                '最低气温': '℃', 
                '最高气温': '℃', 
                '平均湿度': '%', 
                '降水量': 'mm', 
                '平均风速': 'm/s', 
                '最大风速': 'm/s'}
    
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module01'
def spatial_consistency_report(result_dict,main_sta_ids,sub_sta_ids,data_dir):
    
    num=0
    report_path=[]
    report_path.append(os.path.join(cfg['report']['template'],'Module01','spatial_consistency_title.docx'))
    for keys,values in result_dict.items():
        if keys != 'day':
            num=num+1
            doc_path=os.path.join(cfg['report']['template'],'Module01','spatial_consistency.docx')
            doc=DocxTemplate(doc_path)
            
            mean_tem=pd.DataFrame(result_dict[keys])
            dic=dict()
            dic['ele_unit']=unit[keys]
            dic['ele']=keys
            dic['num']=num
            dic['main_station']=mean_tem[mean_tem['站号']==main_sta_ids].iloc[0,1]
            dic['sub_station']=mean_tem[mean_tem['站号']==sub_sta_ids[0]].iloc[0,1]
            dic['num_station']=len(sub_sta_ids)
            
            dic['max_tem_mon']=mean_tem.iloc[0,3:-1:].astype(float).idxmax()
            dic['max_tem_1']=mean_tem[dic['max_tem_mon']].min()
            dic['max_tem_2']=mean_tem[dic['max_tem_mon']].max()
            
            sorted_data = mean_tem.iloc[0,3:-1].astype(float).sort_values()
            
            dic['second_tem_mon']=sorted_data.index[-2]
            dic['third_tem_mon']=sorted_data.index[-3]
            dic['min_tem_mon']=sorted_data.index[0]
            dic['min_tem_1']=mean_tem[dic['min_tem_mon']].min()
            dic['min_tem_2']=mean_tem[dic['min_tem_mon']].max()
            
            
            dic['max_diff'] =round( mean_tem['平均'].max()-mean_tem['平均'].min(),2)
            
            df=mean_tem.iloc[:,3:-1:]
            counter_dict = {}
            
            for column in df.columns:
                first_value = df.iloc[0, df.columns.get_loc(column)]
                
                column_counter = 0
                
                for index in range(1, len(df)):
                    current_value = df.iloc[index, df.columns.get_loc(column)]
                    
                    if current_value < first_value:
                        column_counter += 1
                    elif current_value > first_value:
                        column_counter -= 1
                
                counter_dict[column] = column_counter
                
            big_mon=[]
            small_mon=[]
            middle_mon=[]
        
            for key,values in counter_dict.items():
                if values == np.size(mean_tem,0)-1:
                    big_mon.append(key)
                elif values*(-1) == np.size(mean_tem,0)-1:
                    small_mon.append(key)
                else:
                    middle_mon.append(key)
                    
            if len(big_mon)==12: 
                dic['big_mon']=dic['main_station']+'1-12月'+dic['ele']+'均高于各区域自动站'+dic['ele']
            elif len(small_mon)==12: 
                dic['big_mon']=dic['main_station']+'1-12月'+dic['ele']+'均低于各区域自动站'+dic['ele']
            elif len(middle_mon)==12: 
                dic['big_mon']=dic['main_station']+'1-12月'+dic['ele']+'均介于各区域自动站'+dic['ele']+'之间'
            elif not big_mon: 
                dic['big_mon']=dic['main_station']+'、'.join(middle_mon)+dic['ele']+'介于各区域自动站'+dic['ele']+'之间'
            elif not middle_mon: 
                dic['big_mon']= dic['main_station']+'、'.join(big_mon)+dic['ele']+'均高于各区域自动站'+dic['ele']
            else:
                dic['big_mon']= dic['main_station']+'、'.join(big_mon)+dic['ele']+'均高于各区域自动站'+dic['ele']+'，'+dic['main_station']+'、'.join(middle_mon)+dic['ele']+'介于各区域自动站'+dic['ele']+'之间'
        
            diff_month = mean_tem.iloc[:,3:-1].max()-mean_tem.iloc[:,3:-1].min()
            dic['change_tem_1'] =round(diff_month.min(),2)
            dic['change_tem_2'] =round(diff_month.max(),2)
            
            data=mean_tem.iloc[:,3:-1].T
            data.columns=mean_tem.iloc[:,1]
            
            plt.figure(figsize=(10, 6))
            data.plot(kind='bar')    
            plt.grid(axis='y', linestyle='--', alpha=0.7)    
            plt.xlabel('月份')
            plt.ylabel(dic['ele']+'（'+dic['ele_unit']+'）')
            plt.legend(loc='upper center', bbox_to_anchor=(0.5, 1.08), ncol=len(data.columns), fontsize='small')
        
            y_min, y_max = data.min().min(), data.max().max()
            y_range = y_max - y_min
            margin = y_range * 0.05 if y_range > 0 else 0.1
            plt.ylim(y_min - margin, y_max + margin)

            average_gst_picture_hournum=os.path.join(data_dir,dic['ele']+'.png')
            plt.savefig(average_gst_picture_hournum, bbox_inches='tight', dpi=200)
            plt.clf()
            plt.close('all')
            dic['mean_tem'] = InlineImage(doc, average_gst_picture_hournum, width=Mm(130))

            # 模版文件读取写入字典
            doc.render(dic)
            # 保存结果到新的docx文件
            name2='spatial_consistency_'+keys+'.docx'
            report=os.path.join(data_dir,name2)
            doc.save(report)
            
            document = Document(report)
            
            # 填充表格数据
            mean_tem_x=mean_tem.iloc[:,3:-1]
            name='周边气象站月'+dic['ele']+'分布（单位：'+dic['ele_unit']+'）'
            creat_table(document,mean_tem_x,name)
        
            document.save(report)
            report_path.append(report)
        if num==len(result_dict.keys())-1:
            doc_path=os.path.join(cfg['report']['template'],'Module01','spatial_consistency_test.docx')
            doc=DocxTemplate(doc_path)
            
            mean_tem=pd.DataFrame(result_dict[keys])
            dic=dict()
            dic['main_station']=mean_tem[mean_tem['站号']==main_sta_ids].iloc[0,1]
            dic['sub_station']=mean_tem[mean_tem['站号']==sub_sta_ids[0]].iloc[0,1]
            dic['num_station']=len(sub_sta_ids)
            
            day=result_dict['day']
            
            if len(day.keys())==2:
                dic['test_methods']='T检验和F检验'
                dic['num_test']='表'+str(len(result_dict.keys()))+'和表'+str(len(result_dict.keys())+1)
                dic['t_test']='表'+str(len(result_dict.keys()))+' 厂址站与其他气象站时检验结果(T检验方法)'
                dic['f_test']='表'+str(len(result_dict.keys())+1)+' 厂址站与其他气象站时检验结果(F检验方法)'
            else:
                dic['test_methods']= next(iter(day))[0].upper()+'检验'
                dic['num_test']='表'+str(len(result_dict.keys()))
                dic['t_test']='表'+str(len(result_dict.keys()))+' 厂址站与其他气象站时检验结果'
                dic['f_test']=[]
        
            
            # 模版文件读取写入字典
            doc.render(dic)
            # 保存结果到新的docx文件
            name2='spatial_consistency_'+'day'+'.docx'
            report=os.path.join(data_dir,name2)
            doc.save(report)
            
            document = Document(report)
            
            # 填充表格数据
            for keys,values in day.items():
                values=pd.DataFrame(values)
                columns=values.columns
                
                for i in columns[1::]:
                    values.rename(columns={i:mean_tem[mean_tem['站号']==i].iloc[0,1]}, inplace=True)
                values=values.T
                                
                values.columns = values.iloc[0]
                values = values.drop(values.index[0])
                values = values.reset_index()
                values.rename(columns={'index':'要素'}, inplace=True)
                creat_table(document,values,dic[keys])
        
            document.save(report)
            report_path.append(report)
    try:
        new_docx_path=os.path.join(data_dir,'spatial_consistency.docx')
        master = Document(report_path[0])
        middle_new_docx = Composer(master)
        for word in report_path[1:]:  # 从第二个文档开始追加
            word_document = Document(word)
            # 删除手动添加分页符的代码
            middle_new_docx.append(word_document)
        middle_new_docx.save(new_docx_path)
    except Exception as e:
        print(f"发生错误：{e}")
    
    return new_docx_path
