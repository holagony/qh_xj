# -*- coding: utf-8 -*-
"""
Created on Wed Jun 19 10:34:06 2024

@author: EDY
"""

import matplotlib

matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt

plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document, data, expect_text):
    data = data.astype(str)
    data = data.transpose()
    data = data.reset_index()
    data = data.transpose()

    table = document.add_table(rows=data.shape[0], cols=data.shape[1])

    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i, j]

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)

    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break

    move_table_after(table, target)


def check_variable_exists(var_name):
    return var_name in locals() or var_name in globals()


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module01'
def correlation_analysis_report(result_dict, main_sta_ids, sub_sta_ids, daily_df, data_dir, method):

    doc_path = os.path.join(cfg['report']['template'], 'Module01', 'correlation_analysis.docx')
    doc = DocxTemplate(doc_path)

    result_dict = dict(result_dict)
    day = result_dict['day']
    for keys in day.keys():
        values = day[keys]

        for key in values.keys():
            value = values[key]

            if key == 'regression':
                value = pd.DataFrame(value)

                try:
                    regressions = pd.concat([regressions, value])
                except:
                    regressions = value
            if key == 'ratio':
                value = pd.DataFrame(value)
                try:
                    ratios = pd.concat([ratios, value])
                except:
                    ratios = value

    dic = dict()
    dic['num_station'] = len(sub_sta_ids)
    dic['main_station'] = daily_df[daily_df['Station_Id_C'] == main_sta_ids].iloc[0, 0]
    dic['start_year'] = daily_df.index.year[0]
    dic['end_year'] = daily_df.index.year[-1]
    dic['ele'] = '、'.join(list(day.keys()))

    if len(method) == 2:
        ratios.reset_index(inplace=True)
        ratios.drop(['index'], axis=1, inplace=True)
        ratios.drop(['对比站X', '权重'], axis=1, inplace=True)
        ratios.columns = ['气象要素', '参证站Y', '样本数', '回归方程', '相关系数', 'X均值', 'Y均值']

        for i in np.arange(len(ratios)):
            ratios.iloc[i, 1] = daily_df[daily_df['Station_Id_C'] == ratios.iloc[i, 1]].iloc[0, 0]
            ratios.iloc[i, 4] = round(np.sqrt(ratios.iloc[i, 4]), 2)
        dic['table_2'] = '表2 项目周边国家基本气象站与参证站相关系数（ratios法）'
        dic['method'] = 'regression法'

    elif method[0] == 'regression':

        regressions.reset_index(inplace=True)
        regressions.drop(['index'], axis=1, inplace=True)

        choose_data = regressions[regressions['参证站Y'] == sub_sta_ids[0]]
        choose_data_desc = choose_data.sort_values(by='确定系数', ascending=False)
        choose_data_desc.reset_index(inplace=True)
        choose_data_desc.drop(['index'], axis=1, inplace=True)

        dic['ele_1'] = '、'.join(choose_data_desc.iloc[:int(len(choose_data_desc) / 3):, 0].to_list())
        dic['ele_2'] = '、'.join(choose_data_desc.iloc[int(len(choose_data_desc) / 3) + 1:int(len(choose_data_desc) / 3) * 2:, 0].to_list())
        dic['ele_3'] = choose_data_desc.iloc[-1, 0]
        dic['ele_3_num'] = round(np.sqrt(choose_data_desc.iloc[-1, 7]), 2)

        regressions.drop(['对比站X', '权重', '偏差'], axis=1, inplace=True)

        regressions.columns = ['气象要素', '参证站Y', '样本数', '回归方程', '相关系数', 'X均值', 'Y均值']
        for i in np.arange(len(regressions)):
            regressions.iloc[i, 1] = daily_df[daily_df['Station_Id_C'] == regressions.iloc[i, 1]].iloc[0, 0]
            regressions.iloc[i, 4] = round(np.sqrt(regressions.iloc[i, 4]), 2)

        dic['method'] = 'regression法'
        dic['table_2'] = ''

    elif method[0] == 'ratio':
        ratios.reset_index(inplace=True)
        ratios.drop(['index'], axis=1, inplace=True)

        choose_data = ratios[ratios['参证站Y'] == sub_sta_ids[0]]
        choose_data_desc = choose_data.sort_values(by='确定系数', ascending=False)
        choose_data_desc.reset_index(inplace=True)
        choose_data_desc.drop(['index'], axis=1, inplace=True)

        dic['ele_1'] = '、'.join(choose_data_desc.iloc[:int(len(choose_data_desc) / 3):, 0].to_list())
        dic['ele_2'] = '、'.join(choose_data_desc.iloc[int(len(choose_data_desc) / 3) + 1:int(len(choose_data_desc) / 3) * 2:, 0].to_list())
        dic['ele_3'] = choose_data_desc.iloc[-1, 0]
        dic['ele_3_num'] = round(np.sqrt(choose_data_desc.iloc[-1, 7]), 2)

        dic['method'] = 'ratios法'
        dic['table_2'] = ''

    doc.render(dic)
    # 保存结果到新的docx文件
    report = os.path.join(data_dir, 'correlation_analysis.docx')
    doc.save(report)

    document = Document(report)

    # 填充表格数据
    if len(method) == 2:
        creat_table(document, regressions, 'regression法）')
        creat_table(document, ratios, '（ratios法）')
    elif method[0] == 'regression':
        creat_table(document, regressions, 'regression法）')
    else:
        creat_table(document, ratios, '（ratios法）')

    document.save(report)

    return report
