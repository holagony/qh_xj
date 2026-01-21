# -*- coding: utf-8 -*-
"""
Created on Wed Sep  3 10:46:16 2025

@author: hx

时间一致性分析报告
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docxcompose.composer import Composer

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)
            

            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)

unit = {'平均气压': 'hPa', 
                '最高气压': 'hPa', 
                '最低气压': 'hPa', 
                '平均气温': '℃', 
                '最低气温': '℃', 
                '最高气温': '℃', 
                '平均湿度': '%', 
                '降水量': 'mm', 
                '平均风速': 'm/s', 
                '最大风速': 'm/s'}
    
# data_dir=r'D:\Project\3_项目\13_新疆气候可行性论证\Data\data'
def time_consistency_report(result_dict,sta_ids,sta_name,data_dir):
    
    num=0
    report_path=[]
    report_path.append(os.path.join(cfg['report']['template'],'Module01','time_consistency_title.docx'))
    result_dict =result_dict[sta_ids[0]]
    for keys,values in result_dict.items():
        
        num=num+1
        doc_path=os.path.join(cfg['report']['template'],'Module01','time_consistency.docx')
        doc=DocxTemplate(doc_path)
        
        mean_tem=pd.DataFrame(result_dict[keys].data)
        dic=dict()
        dic['units']=unit[keys]
        dic['ele']=keys
        dic['num']=num
        dic['station_name']=sta_name
        dic['start_year']=mean_tem['年份'].loc[0]
        dic['end_year']=mean_tem['年份'].loc[len(mean_tem)-1]

        dic['mean_data']=mean_tem['要素值'].mean().round(2)
        dic['min_data']=np.round(mean_tem['要素值'].min())
        dic['max_data']=np.round(mean_tem['要素值'].max())
        dic['differ_data']=dic['max_data']-dic['min_data']

        mutation_year_mk=result_dict[keys].mutation_year_mk
        mutation_year_t=result_dict[keys].mutation_year_t

        str_explain1=''
        str_explain2=''
        str_explain3=''
        if mutation_year_mk is not None:
            str_explain1=f"经过mk检验，发现突变年份：{'，'.join(map(str, mutation_year_mk))}。"
        if mutation_year_t is not None:
            str_explain2=f"经过t检验，发现突变年份：{'，'.join(map(str, mutation_year_mk))}。" 
        if mutation_year_mk is None and mutation_year_t is None:
            str_explain3="经过mk检验和t检验，没有发生明显的突变。"
        
        str_explain = str_explain1+str_explain2+str_explain3
        dic['explain']=str_explain
        
        if hasattr(result_dict[keys],'mk_result') and result_dict[keys].mk_result is not None:
            mk_df = pd.DataFrame(result_dict[keys].mk_result)
        else:
            in_seq = mean_tem['要素值'].values
            n = in_seq.shape[0]
            NUMI = np.arange(2, n)
            E = NUMI * (NUMI - 1) / 4
            VAR = NUMI * (NUMI - 1) * (2 * NUMI + 5) / 72
            Ri = [(in_seq[i] > in_seq[1:i]).sum() for i in NUMI]
            Sk = np.cumsum(Ri)
            UFk = np.pad((Sk - E) / np.sqrt(VAR), (2, 0))
            Bk = np.cumsum([(in_seq[i] > in_seq[i:-1]).sum() for i in -(NUMI + 1)])
            UBk = np.pad((-(Bk - E) / np.sqrt(VAR)), (2, 0))[::-1]
            mk_df = pd.DataFrame({'年份': mean_tem['年份'].values, 'UFk': UFk, 'UBk': UBk})
            if mutation_year_mk is None:
                point_idx = []
                diff = UFk - UBk
                for k in range(1, n):
                    if diff[k - 1] * diff[k] < 0:
                        point_idx.append(k)
                mutation_year_mk = mk_df.loc[point_idx, '年份'].tolist() if len(point_idx)>0 else None
        
        mk_explain_base = "采用Mann-Kendall突变检验，绘制UF与UB曲线，交叉处为可能突变点。图中虚线为双边0.01显著性水平（±2.576）。"
        if mutation_year_mk is not None:
            mk_explain_add = "可能突变年份："+'、'.join(map(str, mutation_year_mk))+"。"
        else:
            mk_explain_add = "未发现明显的突变。"
        dic['mk_explain'] = mk_explain_base + mk_explain_add
        
        fig, axes = plt.subplots(2, 1, figsize=(9, 6), sharex=True)
        axes[0].plot(mean_tem['年份'], mean_tem['要素值'], color='black', marker='o', markersize=5, linestyle='-', markerfacecolor='black', markeredgecolor='black', markeredgewidth=0)
        axes[0].tick_params(axis='both', direction='in')
        axes[0].set_ylabel(dic['ele']+'（'+dic['units']+'）')
        year_min = mean_tem['年份'].min()
        year_max = mean_tem['年份'].max()
        num_ticks = 10
        step = max(1, int((year_max - year_min) / (num_ticks - 1)))
        xtick_years = np.arange(year_min, year_max + 1, step)
        axes[0].set_xticks(xtick_years)
        axes[0].xaxis.set_major_locator(MaxNLocator(integer=True))
        axes[1].set_xticks(xtick_years)
        axes[0].set_xlim(np.min(mean_tem['年份']), np.max(mean_tem['年份']))
        y_min, y_max =  mean_tem['要素值'].min(),  mean_tem['要素值'].max()
        y_range = y_max - y_min
        margin = y_range * 0.05 if y_range > 0 else 0.1
        y_min_limit = 0 if y_min >= 0 and y_min - margin < 0 else y_min - margin
        axes[0].set_ylim(y_min_limit, y_max + margin)
        axes[0].grid(axis='y', linestyle='--', alpha=0.4)
        axes[0].grid(axis='x', linestyle='--', alpha=0.4)
        
        axes[1].plot(mk_df['年份'], mk_df['UFk'], color='red', marker='o', markersize=3, linestyle='-', label='UF')
        axes[1].plot(mk_df['年份'], mk_df['UBk'], color='blue', marker='o', markersize=3, linestyle='-', label='UB')
        axes[1].tick_params(axis='both', direction='in')
        axes[1].set_xlabel('年')
        axes[1].set_ylabel('MK统计量')
        axes[1].set_xlim(np.min(mk_df['年份']), np.max(mk_df['年份']))
        axes[1].axhline(2.576, color='gray', linestyle='--', linewidth=1, label='0.01显著水平')
        axes[1].axhline(-2.576, color='gray', linestyle='--', linewidth=1)
        axes[1].axhline(0, color='black', linestyle=':', linewidth=0.8)
        if mutation_year_mk is not None:
            for y in mutation_year_mk:
                axes[1].axvline(y, color='orange', linestyle='--', linewidth=0.8, alpha=0.6)
        axes[1].grid(axis='y', linestyle='--', alpha=0.4)
        axes[1].grid(axis='x', linestyle='--', alpha=0.4)
        axes[1].xaxis.set_major_locator(MaxNLocator(integer=True))
        axes[1].legend(loc='best')
        
        average_gst_picture_hournum=os.path.join(data_dir,dic['ele']+'.png')
        fig.savefig(average_gst_picture_hournum, bbox_inches='tight', dpi=200)
        plt.clf()
        plt.close('all')
        dic['picture'] = InlineImage(doc, average_gst_picture_hournum, width=Mm(130))
        
        plt.figure(figsize=(9, 4))
        plt.plot(mean_tem['年份'], mean_tem['要素值'], color='black', marker='o', markersize=5, linestyle='-', markerfacecolor='black', markeredgecolor='black', markeredgewidth=0)
        plt.tick_params(axis='both', direction='in')
        plt.xlabel('年')
        plt.ylabel(dic['ele']+'（'+dic['units']+'）')
        year_min = mean_tem['年份'].min()
        year_max = mean_tem['年份'].max()
        num_ticks = 10
        step = max(1, int((year_max - year_min) / (num_ticks - 1)))
        xtick_years = np.arange(year_min, year_max + 1, step)
        plt.xticks(xtick_years)
        plt.gca().xaxis.set_major_locator(MaxNLocator(integer=True))
        plt.xlim(np.min(mean_tem['年份']), np.max(mean_tem['年份']))
        y_min, y_max =  mean_tem['要素值'].min(),  mean_tem['要素值'].max()
        y_range = y_max - y_min
        margin = y_range * 0.05 if y_range > 0 else 0.1
        y_min_limit = 0 if y_min >= 0 and y_min - margin < 0 else y_min - margin
        plt.ylim(y_min_limit, y_max + margin)
        plt.grid(axis='y', linestyle='--', alpha=0.4)
        plt.grid(axis='x', linestyle='--', alpha=0.4)
        series_picture_path=os.path.join(data_dir,dic['ele']+'_series.png')
        plt.savefig(series_picture_path, bbox_inches='tight', dpi=200)
        plt.clf()
        plt.close('all')
        dic['series_picture'] = InlineImage(doc, series_picture_path, width=Mm(130))
        
        plt.figure(figsize=(9, 4))
        plt.plot(mk_df['年份'], mk_df['UFk'], color='red', marker='o', markersize=3, linestyle='-', label='UF')
        plt.plot(mk_df['年份'], mk_df['UBk'], color='blue', marker='o', markersize=3, linestyle='-', label='UB')
        plt.tick_params(axis='both', direction='in')
        plt.xlabel('年')
        plt.ylabel('MK统计量')
        plt.xlim(np.min(mk_df['年份']), np.max(mk_df['年份']))
        plt.gca().xaxis.set_major_locator(MaxNLocator(integer=True))
        plt.axhline(2.576, color='gray', linestyle='--', linewidth=1, label='0.01显著水平')
        plt.axhline(-2.576, color='gray', linestyle='--', linewidth=1)
        plt.axhline(0, color='black', linestyle=':', linewidth=0.8)
        if mutation_year_mk is not None:
            for y in mutation_year_mk:
                plt.axvline(y, color='orange', linestyle='--', linewidth=0.8, alpha=0.6)
        plt.grid(axis='y', linestyle='--', alpha=0.4)
        plt.grid(axis='x', linestyle='--', alpha=0.4)
        plt.legend(loc='best')
        mk_picture_path=os.path.join(data_dir,dic['ele']+'_mk.png')
        plt.savefig(mk_picture_path, bbox_inches='tight', dpi=200)
        plt.clf()
        plt.close('all')
        dic['mk_picture'] = InlineImage(doc, mk_picture_path, width=Mm(130))

        # 模版文件读取写入字典
        doc.render(dic)
        # 保存结果到新的docx文件
        name2='time_consistency_'+keys+'.docx'
        report=os.path.join(data_dir,name2)
        doc.save(report)
        
        report_path.append(report)
      
    try:
        new_docx_path=os.path.join(data_dir,'time_consistency.docx')
        master = Document(report_path[0])
        middle_new_docx = Composer(master)
        for word in report_path[1:]:  # 从第二个文档开始追加
            word_document = Document(word)
            # 删除手动添加分页符的代码
            middle_new_docx.append(word_document)
        middle_new_docx.save(new_docx_path)
    except Exception as e:
        print(f"发生错误：{e}")
    
    return new_docx_path
