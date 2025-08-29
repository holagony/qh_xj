# -*- coding: utf-8 -*-
"""
Created on Mon Jul 22 09:55:38 2024

@author: EDY
"""


import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            

    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)

 
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module07'
def airport_wind_ds_report(all_result,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module08','airport_wind_ds.docx')
 
    document = Document(doc_path)
    

    # 填充表格数据
    creat_table(document,all_result,'不同风速条件下风向次数统计')

    report=os.path.join(data_dir,'airport_wind_ds.docx')

    document.save(report)


    
    return report