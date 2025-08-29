# -*- coding: utf-8 -*-
"""
Created on Thu Jun 27 16:15:27 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH


plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text

        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13'
# doc_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13\rain_step_1.docx'
# data_flag=0
# step2_csv = r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13\multi_sample.csv' # 年最大值/年多样本
# mode=4
def merged(report_path,new_docx_path):
    master = Document(report_path[0])
    middle_new_docx = Composer(master)
    for word in report_path[1:]:  
        word_document = Document(word)
        middle_new_docx.append(word_document)
    middle_new_docx.save(new_docx_path) 
    return 1

def step_2_report(data_flag,step2_csv,mode,doc_dir,result,pre_data,result_formula, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    #%% 简单处理数据
    pre_data.drop(['Unnamed: 0'], axis=1, inplace=True)
   
    dic=dict()
    
    new_docx_path=os.path.join(data_dir,'temt_rain_step_2.docx')
    if mode==0:
        report_path=[doc_dir,os.path.join(cfg['report']['template'],'Module13','rain_step_2_g.docx')]
        merged(report_path,new_docx_path)

    elif mode==2:
        report_path=[doc_dir,os.path.join(cfg['report']['template'],'Module13','rain_step_2_p.docx')]
        merged(report_path,new_docx_path)
    else:
        report_path=[doc_dir,os.path.join(cfg['report']['template'],'Module13','rain_step_2_z.docx')]
        merged(report_path,new_docx_path)

    
    doc_path=os.path.join(data_dir,'temt_rain_step_2.docx')
    doc = DocxTemplate(doc_path)
      
    #%% picture

    dic['pic_1']=InlineImage(doc, result['5min']['img_save_path'], width=Mm(130))
    dic['pic_2']=InlineImage(doc, result['10min']['img_save_path'], width=Mm(130))
    dic['pic_3']=InlineImage(doc, result['15min']['img_save_path'], width=Mm(130))
    dic['pic_4']=InlineImage(doc, result['20min']['img_save_path'], width=Mm(130))
    dic['pic_5']=InlineImage(doc, result['30min']['img_save_path'], width=Mm(130))
    dic['pic_6']=InlineImage(doc, result['45min']['img_save_path'], width=Mm(130))
    dic['pic_7']=InlineImage(doc, result['60min']['img_save_path'], width=Mm(130))
    dic['pic_8']=InlineImage(doc, result['90min']['img_save_path'], width=Mm(130))
    dic['pic_9']=InlineImage(doc, result['120min']['img_save_path'], width=Mm(130))
    dic['pic_10']=InlineImage(doc, result['150min']['img_save_path'], width=Mm(130))
    dic['pic_11']=InlineImage(doc, result['180min']['img_save_path'], width=Mm(130))


    if data_flag==0:
        dic['li_shi']='5min、10min、15min、20min、30min、45min、60min、90min、120min、150min、180min'
        keys=['5min','10min','15min','20min','30min','45min','60min','90min','120min','150min','180min']
        dic['li_shi_num']=len(keys)
        dic['title_16']='图12'
    else:
        dic['li_shi']='5min、10min、15min、20min、30min、45min、60min、90min、120min、150min、180min、240min、360min、720min、1440min'
        keys=['5min','10min','15min','20min','30min','45min','60min','90min','120min','150min','180min','240min','360min','720min','1440min']
        dic['li_shi_num']=len(keys)

        dic['pic_12']=InlineImage(doc, result['240min']['img_save_path'], width=Mm(130))
        dic['pic_13']=InlineImage(doc, result['360min']['img_save_path'], width=Mm(130))
        dic['pic_14']=InlineImage(doc, result['720min']['img_save_path'], width=Mm(130))
        dic['pic_15']=InlineImage(doc, result['1440min']['img_save_path'], width=Mm(130))
        dic['title_16']='图16'

        if mode==0:
       
            dic['title_12']='图12 240min历时耿贝尔分布拟合曲线（包含适线）'
            dic['title_13']='图13 360min历时耿贝尔分布拟合曲线（包含适线）'
            dic['title_14']='图14 720min历时耿贝尔分布拟合曲线（包含适线）'
            dic['title_15']='图15 1440min历时耿贝尔分布拟合曲线（包含适线）'
        elif mode==4: 
            dic['title_12']='图12 240min历时皮尔逊分布拟合曲线（包含适线）'
            dic['title_13']='图13 360min历时皮尔逊分布拟合曲线（包含适线）'
            dic['title_14']='图14 720min历时皮尔逊分布拟合曲线（包含适线）'
            dic['title_15']='图15 1440min历时皮尔逊分布拟合曲线（包含适线）'
        elif mode==6: 
            dic['title_12']='图12 240min历时指数分布拟合曲线（包含适线）'
            dic['title_13']='图13 360min历时指数分布拟合曲线（包含适线）'
            dic['title_14']='图14 720min历时指数分布拟合曲线（包含适线）'
            dic['title_15']='图15 1440min历时指数分布拟合曲线（包含适线）'    

    filename = os.path.basename(step2_csv)
    if filename=='multi_sample.csv':
        dic['method_1']='年多样本法'
        dic['method_d']='年多样本法'       
        dic['method_d2']='每年前8个大值'

    else:
        dic['method_1']='年最大值法'
        dic['method_d']='年最大值法为每年选一个最大值，选样简单，独立性强，合理地考虑了气象现象以一年为循环的特点，其机率意义是一年发生一次的频率年值，理论上比较严密。'
        dic['method_d2']='每年的最大值'
      
    dic['num_years_s']=np.size(pre_data,0)
    dic['num_rain']=dic['num_years_s']* dic['li_shi_num']

    dic['pic_num']=len(dic['li_shi'])
    
    dic['A']=result_formula['full']['params']['A']
    dic['B']=result_formula['full']['params']['b']
    dic['C']=result_formula['full']['params']['C']
    dic['N']=result_formula['full']['params']['n']
    dic['mae_1']=result_formula['full']['error']['mae']
    dic['rse_1']=result_formula['full']['error']['rmse']
    dic['xd_1']=result_formula['full']['error']['rel_error']
   
    dic['A_2']=result_formula['2a']['params']['A']
    dic['B_2']=result_formula['2a']['params']['b']
    dic['N_2']=result_formula['2a']['params']['n']
    dic['mae_2']=result_formula['2a']['error']['mae']
    dic['rse_2']=result_formula['2a']['error']['rmse']
    dic['xd_2']=result_formula['2a']['error']['rel_error']
    
    dic['A_3']=result_formula['3a']['params']['A']
    dic['B_3']=result_formula['3a']['params']['b']
    dic['N_3']=result_formula['3a']['params']['n']
    dic['mae_3']=result_formula['3a']['error']['mae']
    dic['rse_3']=result_formula['3a']['error']['rmse']
    dic['xd_3']=result_formula['3a']['error']['rel_error']
    
    dic['A_4']=result_formula['5a']['params']['A']
    dic['B_4']=result_formula['5a']['params']['b']
    dic['N_4']=result_formula['5a']['params']['n']
    dic['mae_4']=result_formula['5a']['error']['mae']
    dic['rse_4']=result_formula['5a']['error']['rmse']
    dic['xd_4']=result_formula['5a']['error']['rel_error']
    
    dic['A_5']=result_formula['10a']['params']['A']
    dic['B_5']=result_formula['10a']['params']['b']
    dic['N_5']=result_formula['10a']['params']['n']
    dic['mae_5']=result_formula['10a']['error']['mae']
    dic['rse_5']=result_formula['10a']['error']['rmse']
    dic['xd_5']=result_formula['10a']['error']['rel_error']
    
    dic['A_6']=result_formula['20a']['params']['A']
    dic['B_6']=result_formula['20a']['params']['b']
    dic['N_6']=result_formula['20a']['params']['n']
    dic['mae_6']=result_formula['20a']['error']['mae']
    dic['rse_6']=result_formula['20a']['error']['rmse']
    dic['xd_6']=result_formula['20a']['error']['rel_error']
    
    dic['A_7']=result_formula['30a']['params']['A']
    dic['B_7']=result_formula['30a']['params']['b']
    dic['N_7']=result_formula['30a']['params']['n']
    dic['mae_7']=result_formula['30a']['error']['mae']
    dic['rse_7']=result_formula['30a']['error']['rmse']
    dic['xd_7']=result_formula['30a']['error']['rel_error']
    
    dic['A_8']=result_formula['50a']['params']['A']
    dic['B_8']=result_formula['50a']['params']['b']
    dic['N_8']=result_formula['50a']['params']['n']
    dic['mae_8']=result_formula['50a']['error']['mae']
    dic['rse_8']=result_formula['50a']['error']['rmse']
    dic['xd_8']=result_formula['50a']['error']['rel_error']
    
    dic['A_9']=result_formula['100a']['params']['A']
    dic['B_9']=result_formula['100a']['params']['b']
    dic['N_9']=result_formula['100a']['params']['n']
    dic['mae_9']=result_formula['100a']['error']['mae']
    dic['rse_9']=result_formula['100a']['error']['rmse']
    dic['xd_9']=result_formula['100a']['error']['rel_error']
    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'rain_step_2.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    data=[]
    for i in keys:
        a=result[i]['distribution_info']
        data.append(a)
        
    data=pd.DataFrame(data)
    data.index=keys
    data.reset_index(inplace=True)
    if mode==4:
        data.columns=['历时','期望','Cs','Cv','Cs/Cv','RMSE','MAE','相对误差']
    elif mode==6:
        data.columns=['历时','expon_loc','expon_scale','RMSE','MAE','相对误差']

    
    
    pit_table=pd.DataFrame(result['pit_table'])
    
    pre_data['year'] = pre_data['year'].astype(int)
    pre_data.iloc[:,1::]=round(pre_data.iloc[:,1::],2)
 
    creat_table(document,pre_data,'历年各历时降水量')
    creat_table(document,pit_table,'分布P-i-t表(单位：mm/min)')
    
    if mode==4: 
        creat_table(document,data,'各历时的Cv、Cs参数值')
    elif mode==6:
        creat_table(document,data,'指数分布拟合各历时的参数值')

    document.save(report)
    return report
