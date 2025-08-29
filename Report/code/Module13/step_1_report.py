# -*- coding: utf-8 -*-
"""
Created on Thu Jun 27 15:55:48 2024

@author: EDY
"""


import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13'
def step_1_report(result_1, start_year_out, end_year_out,result_2, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module13','rain_step_1.docx')
    doc=DocxTemplate(doc_path)
    
    
    dic=dict()
    dic['station_name']=''
    dic['start_year']=start_year_out
    dic['end_year']=end_year_out
    dic['num_years']=len(result_1.keys())-1

    dic['method_1']='{{method_1}}'
    dic['method_d']='{{method_d}}'
    dic['method_d2']='{{method_d2}}'
    dic['num_rain']='{{num_rain}}'
    dic['li_shi']='{{li_shi}}'
    dic['li_shi_num']='{{li_shi_num}}'
    dic['num_years_s']='{{num_years_s}}'

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'rain_step_1.docx')
    doc.save(report)
    
    return report
