# -*- coding: utf-8 -*-
"""
Created on Fri Jun 28 11:13:41 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH


plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text

        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13'
# doc_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules13\rain_step_2.docx'
# data_flag=0
def merged(report_path,new_docx_path):
    master = Document(report_path[0])
    middle_new_docx = Composer(master)
    for word in report_path[1:]:  
        word_document = Document(word)
        middle_new_docx.append(word_document)
    middle_new_docx.save(new_docx_path) 
    return 1

def step_3_report(data_flag,doc_dir,result_dict, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    #%% 简单处理数据
   
    dic=dict()
    
    new_docx_path=os.path.join(data_dir,'temt_rain_step_3.docx')
    report_path=[doc_dir,os.path.join(cfg['report']['template'],'Module13','rain_step_3.docx')]
    merged(report_path,new_docx_path)
    
    doc_path=os.path.join(data_dir,'temt_rain_step_3.docx')
    doc=DocxTemplate(doc_path)
      
    #%% data
    data1=pd.DataFrame(result_dict['不同历时平均雨峰系数'])
    dic['r']=data1.iloc[0,6]
    #%% picture
    dic['picture_17']=InlineImage(doc, result_dict['img_save_path']['60min雨型'], width=Mm(130))
    dic['picture_18']=InlineImage(doc, result_dict['img_save_path']['90min雨型'], width=Mm(130))
    dic['picture_19']=InlineImage(doc, result_dict['img_save_path']['120min雨型'], width=Mm(130))
    dic['picture_20']=InlineImage(doc, result_dict['img_save_path']['150min雨型'], width=Mm(130))
    dic['picture_21']=InlineImage(doc, result_dict['img_save_path']['180min雨型'], width=Mm(130))

    if data_flag==0:

        dic['title_17']='图13'
        dic['title_18']='图14'
        dic['title_19']='图15'
        dic['title_20']='图16'
        dic['title_21']='图17'

    else:

        dic['title_17']='图17'
        dic['title_18']='图18'
        dic['title_19']='图19'
        dic['title_20']='图20'
        dic['title_21']='图21'    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'rain_step_3.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    data60=pd.DataFrame(result_dict['rain_type']['60min雨型'])
    data60.reset_index(inplace=True)
    data60.rename(columns={'index': '历时'}, inplace=True)
    
    data90=pd.DataFrame(result_dict['rain_type']['90min雨型'])
    data90.reset_index(inplace=True)
    data90.rename(columns={'index': '历时'}, inplace=True)
    
    data120=pd.DataFrame(result_dict['rain_type']['120min雨型'])
    data120.reset_index(inplace=True)
    data120.rename(columns={'index': '历时'}, inplace=True)
    
    data150=pd.DataFrame(result_dict['rain_type']['150min雨型'])
    data150.reset_index(inplace=True)
    data150.rename(columns={'index': '历时'}, inplace=True)
    
    data180=pd.DataFrame(result_dict['rain_type']['180min雨型'])
    data180.reset_index(inplace=True)
    data180.rename(columns={'index': '历时'}, inplace=True)
    
    creat_table(document,data1,'各历时雨峰位置')
    creat_table(document,data60,'60min各时段瞬时降雨强度成果表（单位：mm/min）')
    creat_table(document,data90,'90min各时段瞬时降雨强度成果表（单位：mm/min）')
    creat_table(document,data120,'120min各时段瞬时降雨强度成果表（单位：mm/min）')
    creat_table(document,data150,'150min各时段瞬时降雨强度成果表（单位：mm/min）')
    creat_table(document,data180,'180min各时段瞬时降雨强度成果表（单位：mm/min）')

    document.save(report)
    
    return report
