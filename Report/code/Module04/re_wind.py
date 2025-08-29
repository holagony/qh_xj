# -*- coding: utf-8 -*-
"""
Created on Thu May 30 15:45:03 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# daily_df=df_sequence
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'
# wind_s_result=wind_result
def re_wind_report(wind_s_result,daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_WIND.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=daily_df['Station_Name'][0]

    return_years=wind_s_result['return_years']
    main_return_result=dict(wind_s_result['main_return_result'])
    img_save_path=dict(wind_s_result['img_save_path'])
    
    max_values=dict(main_return_result['max_values'])
    max_values=pd.DataFrame(max_values)
    
    wind_pressure=dict(main_return_result['wind_pressure'])
    wind_pressure=pd.DataFrame(wind_pressure)

    max_values_inst=dict(main_return_result['max_values_i'])
    max_values_inst=pd.DataFrame(max_values_inst)

    dic=dict()
    

    dic['station_name']=station_name
    
    dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
    dic['picture_p'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))
    dic['picture_gd_2'] = InlineImage(doc, img_save_path['P3_plot_i'], width=Mm(130))
    dic['picture_p_2'] = InlineImage(doc, img_save_path['Gumbel_plot_i'], width=Mm(130))

    dic['wind_50_G']=max_values['耿贝尔'][return_years.index(50)]
    dic['wind_50_P']=max_values['皮尔逊Ⅲ型'][return_years.index(50)]
    dic['wind_100_G']=max_values['耿贝尔'][return_years.index(100)]
    dic['wind_100_P']=max_values['皮尔逊Ⅲ型'][return_years.index(100)]
    dic['wind_50j_G']=max_values_inst['耿贝尔'][return_years.index(50)]
    dic['wind_50j_P']=max_values_inst['皮尔逊Ⅲ型'][return_years.index(50)]
    dic['wind_100j_G']=max_values_inst['耿贝尔'][return_years.index(100)]
    dic['wind_100j_P']=max_values_inst['皮尔逊Ⅲ型'][return_years.index(100)]
    
    if dic['wind_50_G']>= dic['wind_50_P']:
        dic['g_or_p']='GD法'
        dic['prs_wind']=wind_pressure['耿贝尔'][return_years.index(50)]
    else:
        dic['g_or_p']='P-Ⅲ法'
        dic['prs_wind']=wind_pressure['皮尔逊Ⅲ型'][return_years.index(50)]



    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_WIND.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    wind_data=pd.DataFrame(wind_s_result['wind_data'])
    wind_data_i=pd.DataFrame(wind_s_result['wind_data_i'])
    
    wind_data['年份']=wind_data['年份'].astype(str)
    wind_data_i['年份']=wind_data_i['年份'].astype(str)

    half_length =int(np.ceil(len(wind_data)/3))
   
    # 分割每个列并创建新列
    wind_data['A1'] = wind_data.iloc[:half_length:,0]
    wind_data['B1'] =wind_data.iloc[:half_length:,1]
    wind_data['A2'] = wind_data.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    wind_data['B2'] =wind_data.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    wind_data['A3'] =wind_data.iloc[half_length*2::,0].reset_index(drop=True)
    wind_data['B3'] = wind_data.iloc[half_length*2::,1].reset_index(drop=True)

    wind_data.drop(['年份', '最大风速(m/s)'], axis=1, inplace=True)
    
    wind_data.dropna(how='all',inplace=True)
    wind_data.columns=['年份', '最大风速','年份', '最大风速','年份', '最大风速']
    
    half_length =int(np.ceil(len(wind_data_i)/3))
   
    # 分割每个列并创建新列
    wind_data_i['A1'] = wind_data_i.iloc[:half_length:,0]
    wind_data_i['B1'] =wind_data_i.iloc[:half_length:,1]
    wind_data_i['A2'] = wind_data_i.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    wind_data_i['B2'] =wind_data_i.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    wind_data_i['A3'] =wind_data_i.iloc[half_length*2::,0].reset_index(drop=True)
    wind_data_i['B3'] = wind_data_i.iloc[half_length*2::,1].reset_index(drop=True)

    wind_data_i.drop(['年份', '极大风速(m/s)'], axis=1, inplace=True)
    
    wind_data_i.dropna(how='all',inplace=True)
    wind_data_i.columns=['年份', '极大风速','年份', '极大风速','年份', '极大风速']
     
    
    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]='GD'
    max_values.iloc[1,0]='P-Ⅲ'

    
    max_values_inst=max_values_inst.T.reset_index()
    max_values_inst.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values_inst.iloc[0,0]='GD'
    max_values_inst.iloc[1,0]='P-Ⅲ'
    
    wind_pressure=wind_pressure.T.reset_index()
    wind_pressure.columns=['重现期']+[str(year)+'a' for year in return_years]
    wind_pressure.iloc[0,0]='GD'
    wind_pressure.iloc[1,0]='P-Ⅲ'
    
    creat_table(document,wind_data,'最大风速')
    creat_table(document,wind_data_i,'极大风速')
    creat_table(document,max_values,'最大风速（单位：m/s）')
    creat_table(document,max_values_inst,'极大风速（单位：m/s）')
    creat_table(document,wind_pressure,'（单位：kN/m2）')

    document.save(report)
    
    
    return report

    
def re_wind_report_pg(wind_s_result,daily_df,methods,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_WIND_P.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=daily_df['Station_Name'][0]

    return_years=wind_s_result['return_years']
    main_return_result=dict(wind_s_result['main_return_result'])
    img_save_path=dict(wind_s_result['img_save_path'])
    
    max_values=dict(main_return_result['max_values'])
    max_values=pd.DataFrame(max_values)
    
    wind_pressure=dict(main_return_result['wind_pressure'])
    wind_pressure=pd.DataFrame(wind_pressure)

    max_values_inst=dict(main_return_result['max_values_i'])
    max_values_inst=pd.DataFrame(max_values_inst)
    
    
    if methods=='Gumbel':
        name_str = '耿贝尔'
    else:
        name_str = '皮尔逊Ⅲ型'
    
    
    dic=dict()
    dic['methods']=methods
    dic['station_name']=station_name
    

    if methods=='Gumbel':
        dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
        dic['picture_gd_2'] = InlineImage(doc, img_save_path['Gumbel_plot_i'], width=Mm(130))

    else:
        dic['picture_gd'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))
        dic['picture_gd_2'] = InlineImage(doc, img_save_path['P3_plot_i'], width=Mm(130))

    dic['wind_50_G']=max_values[name_str][return_years.index(50)]
    dic['wind_100_G']=max_values[name_str][return_years.index(100)]
    dic['wind_50j_G']=max_values_inst[name_str][return_years.index(50)]
    dic['wind_100j_G']=max_values_inst[name_str][return_years.index(100)]
    dic['prs_wind']=wind_pressure[name_str][return_years.index(50)]


    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_WIND.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    wind_data=pd.DataFrame(wind_s_result['wind_data'])
    wind_data_i=pd.DataFrame(wind_s_result['wind_data_i'])
    
    wind_data['年份']=wind_data['年份'].astype(str)
    wind_data_i['年份']=wind_data_i['年份'].astype(str)

    half_length =int(np.ceil(len(wind_data)/3))
   
    # 分割每个列并创建新列
    wind_data['A1'] = wind_data.iloc[:half_length:,0]
    wind_data['B1'] =wind_data.iloc[:half_length:,1]
    wind_data['A2'] = wind_data.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    wind_data['B2'] =wind_data.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    wind_data['A3'] =wind_data.iloc[half_length*2::,0].reset_index(drop=True)
    wind_data['B3'] = wind_data.iloc[half_length*2::,1].reset_index(drop=True)

    wind_data.drop(['年份', '最大风速(m/s)'], axis=1, inplace=True)
    
    wind_data.dropna(how='all',inplace=True)
    wind_data.columns=['年份', '最大风速','年份', '最大风速','年份', '最大风速']
    
    half_length =int(np.ceil(len(wind_data_i)/3))
   
    # 分割每个列并创建新列
    wind_data_i['A1'] = wind_data_i.iloc[:half_length:,0]
    wind_data_i['B1'] =wind_data_i.iloc[:half_length:,1]
    wind_data_i['A2'] = wind_data_i.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    wind_data_i['B2'] =wind_data_i.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    wind_data_i['A3'] =wind_data_i.iloc[half_length*2::,0].reset_index(drop=True)
    wind_data_i['B3'] = wind_data_i.iloc[half_length*2::,1].reset_index(drop=True)

    wind_data_i.drop(['年份', '极大风速(m/s)'], axis=1, inplace=True)
    
    wind_data_i.dropna(how='all',inplace=True)
    wind_data_i.columns=['年份', '极大风速','年份', '极大风速','年份', '极大风速']
     
    
    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]=methods

    
    max_values_inst=max_values_inst.T.reset_index()
    max_values_inst.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values_inst.iloc[0,0]=methods
    
    wind_pressure=wind_pressure.T.reset_index()
    wind_pressure.columns=['重现期']+[str(year)+'a' for year in return_years]
    wind_pressure.iloc[0,0]=methods
    
    creat_table(document,wind_data,'最大风速')
    creat_table(document,wind_data_i,'极大风速')
    creat_table(document,max_values,'最大风速（单位：m/s）')
    creat_table(document,max_values_inst,'极大风速（单位：m/s）')
    creat_table(document,wind_pressure,'（单位：kN/m2）')

    document.save(report)
    
    
    return report

if __name__ == '__main__':
    

    daily_df=df_sequence
    data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'
    methods=fitting_method[0]
    wind_s_result=wind_result
    report=re_wind_report_pg(wind_result,daily_df,methods,data_dir)