# -*- coding: utf-8 -*-
"""
Created on Fri Jun  7 15:26:26 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# year_df=df_sequence
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'

def re_snow_report(snow_result,year_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_SNOW.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=year_df['Station_Name'][0]

    return_years=snow_result['return_years']
    main_return_result=dict(snow_result['main_return_result'])
    img_save_path=dict(snow_result['img_save_path'])
    
    max_values=dict(main_return_result['max_values'])
    max_values=pd.DataFrame(max_values)
    
    snow_pressure=dict(main_return_result['max_values_snow_prs'])
    snow_pressure=pd.DataFrame(snow_pressure)



    dic=dict()
    

    dic['station_name']=station_name
    dic['start_year']=year_df.index.year[0]
    dic['end_year']=year_df.index.year[-1]
    dic['num_years']=len(year_df.index.year)
    dic['max_snow']=year_df['Snow_Depth_Max'].max()

    result_snow=pd.DataFrame(columns=['年份','雪深'])
    result_snow['年份']=year_df.index.year
    result_snow['雪深']=year_df.reset_index()['Snow_Depth_Max']


    dic['snow_50_G']=max_values['Gumbel_max_vals'][return_years.index(50)]
    dic['snow_50_P']=max_values['P3_max_vals'][return_years.index(50)]
    dic['snow_100_G']=max_values['Gumbel_max_vals'][return_years.index(100)]
    dic['snow_100_P']=max_values['P3_max_vals'][return_years.index(100)]

    
    if dic['snow_50_G']>= dic['snow_50_P']:
        dic['g_or_p']='耿贝尔法'
    else:
        dic['g_or_p']='皮尔逊-Ⅲ法'

    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]='GD'
    max_values.iloc[1,0]='P-Ⅲ'
    
    dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
    dic['picture_p'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))


    dic['snow_50p_G']=snow_pressure['Gumbel_max_vals'][return_years.index(50)]
    dic['snow_50p_P']=snow_pressure['P3_max_vals'][return_years.index(50)]
    dic['snow_100p_G']=snow_pressure['Gumbel_max_vals'][return_years.index(100)]
    dic['snow_100p_P']=snow_pressure['P3_max_vals'][return_years.index(100)] 

    if dic['snow_50p_G']>= dic['snow_50p_P']:
        dic['g_or_p_p']='耿贝尔法'
    else:
        dic['g_or_p_p']='皮尔逊-Ⅲ法'

    snow_pressure=snow_pressure.T.reset_index()
    snow_pressure.columns=['重现期']+[str(year)+'a' for year in return_years]
    snow_pressure.iloc[0,0]='GD'
    snow_pressure.iloc[1,0]='P-Ⅲ'

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_SNOW.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    
    
    creat_table(document,result_snow,'年）最大积雪深度（单位：cm）')
    creat_table(document,max_values,'期最大积雪深度（单位：cm）')
    creat_table(document,snow_pressure,'（单位：kN/m2）')

    document.save(report)
    
    
    return report

    
def re_snow_report_pg(snow_result,year_df,methods,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_SNOW_P.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=year_df['Station_Name'][0]

    return_years=snow_result['return_years']
    main_return_result=dict(snow_result['main_return_result'])
    img_save_path=dict(snow_result['img_save_path'])
    
    max_values=dict(main_return_result['max_values'])
    max_values=pd.DataFrame(max_values)
    
    snow_pressure=dict(main_return_result['max_values_snow_prs'])
    snow_pressure=pd.DataFrame(snow_pressure)

    if methods=='Gumbel':
        name_str = 'Gumbel_max_vals'
    else:
        name_str = 'P3_max_vals'

    dic=dict()
    
    dic['methods']=methods

    dic['station_name']=station_name
    dic['start_year']=year_df.index.year[0]
    dic['end_year']=year_df.index.year[-1]
    dic['num_years']=len(year_df.index.year)
    dic['max_snow']=year_df['Snow_Depth_Max'].max()

    result_snow=pd.DataFrame(columns=['年份','雪深'])
    result_snow['年份']=year_df.index.year
    result_snow['雪深']=year_df.reset_index()['Snow_Depth_Max']


    dic['snow_50_G']=max_values[name_str][return_years.index(50)]
    dic['snow_100_G']=max_values[name_str][return_years.index(100)]

    


    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]=methods
    
    if methods=='Gumbel':
        dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
    else:
        dic['picture_gd'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))

    dic['snow_50p_G']=snow_pressure[name_str][return_years.index(50)]
    dic['snow_100p_G']=snow_pressure[name_str][return_years.index(100)]



    snow_pressure=snow_pressure.T.reset_index()
    snow_pressure.columns=['重现期']+[str(year)+'a' for year in return_years]
    snow_pressure.iloc[0,0]=methods

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_SNOW.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    
    
    creat_table(document,result_snow,'年）最大积雪深度（单位：cm）')
    creat_table(document,max_values,'期最大积雪深度（单位：cm）')
    creat_table(document,snow_pressure,'（单位：kN/m2）')

    document.save(report)
    
    
    return report

if __name__ == '__main__':
    year_df=df_sequence
    data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'

    methods=fitting_method[0]
    report=re_snow_report_pg(snow_result,year_df,methods,data_dir)