# -*- coding: utf-8 -*-
"""
Created on Tue Jun 11 11:15:51 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# daily_df=df_sequence
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'

def re_pre_report(pre_result,daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_PRE.docx')
    doc=DocxTemplate(doc_path)
    

    station_name=daily_df['Station_Name'][0]

    return_years=pre_result['return_years']
    
    PRE_Max_Day=dict(pre_result['PRE_Max_Day'])
    img_save_path=dict(PRE_Max_Day['img_save_path'])
    
    return_result=dict(PRE_Max_Day['return_result'])
    max_values=dict(return_result['max_values'])
    max_values=pd.DataFrame(max_values)

    data=PRE_Max_Day['data']
    data=pd.DataFrame(data)



    dic=dict()
    

    dic['station_name']=station_name

    dic['max_year']=data[data['最大日降水量(mm)']==data['最大日降水量(mm)'].max()].iloc[0,0]
    dic['max_year_pre']=data['最大日降水量(mm)'].max()
    dic['min_year']=data[data['最大日降水量(mm)']==data['最大日降水量(mm)'].min()].iloc[0,0]
    dic['min_year_pre']=data['最大日降水量(mm)'].min()
    
    dic['min_y']=min(return_years)
    dic['max_y']=max(return_years)
        
    # years=data['年份'][:]
    mask = ~np.isnan(data['最大日降水量(mm)'])
    valid_years = data['年份'][mask]
    valid_preperatures = data['最大日降水量(mm)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_preperatures, 1)   
    if slope> 0:
        dic['slope']='上升'
    else:
        dic['slope']='下降'
    
    pre_picture=plot_picture(data, '年份','最大日降水量(mm)','最大日降水量(mm)','mm','最大日降水量年际变化图.png',10,10,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(data['年份'], data['最大日降水量(mm)'], width=0.4, color='skyblue', label='最大日降水量')
    # plt.plot(years, slope * years + intercept, color='red', label='最大日降水量')

    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('年')
    # plt.ylabel('最大日降水量（mm）')
    # plt.xticks(years[::3])
    # # plt.ylim(dic['average_pre_1']-5, dic['average_pre_2']+5)
    # plt.legend()
    
    # pre_picture=os.path.join(data_dir,'最大日降水量年际变化图.png')
    # plt.savefig(pre_picture, bbox_inches='tight', dpi=200)
    # plt.clf()
    # plt.close('all')
    dic['pre_picture'] = InlineImage(doc, pre_picture, width=Mm(130))
    
    idmin=return_years.index(min(return_years))
    idmmax=return_years.index(max(return_years))

    
    dic['min_pre_p'] = max_values.iloc[idmin,1]
    dic['max_pre_p'] = max_values.iloc[idmmax,1]
    dic['min_pre_g'] = max_values.iloc[idmin,0]
    dic['max_pre_g'] = max_values.iloc[idmmax,0]
    
    if dic['min_pre_g']>= dic['min_pre_p']:
        dic['g_or_p']='耿贝尔法'
        dic['pre_50']=max_values.iloc[return_years.index(50),0]
        dic['pre_100']=max_values.iloc[return_years.index(100),0]        
    else:
        dic['g_or_p']='皮尔逊-Ⅲ法'
        dic['pre_50']=max_values.iloc[return_years.index(50),1]
        dic['pre_100']=max_values.iloc[return_years.index(100),1]
    
    

    dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
    dic['picture_p'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))



    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_PRE.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    
    data_pre=pd.DataFrame(pre_result['PRE_Max_Day']['data'])
    data_pre['年份']=data_pre['年份'].astype(str)

    half_length =int(np.ceil(len(data_pre)/3))
   
    # 分割每个列并创建新列
    data_pre['A1'] = data_pre.iloc[:half_length:,0]
    data_pre['B1'] =data_pre.iloc[:half_length:,1]
    data_pre['A2'] = data_pre.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    data_pre['B2'] =data_pre.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    data_pre['A3'] =data_pre.iloc[half_length*2::,0].reset_index(drop=True)
    data_pre['B3'] = data_pre.iloc[half_length*2::,1].reset_index(drop=True)

    data_pre.drop(['年份', '最大日降水量(mm)'], axis=1, inplace=True)
    
    data_pre.dropna(how='all',inplace=True)
    data_pre.columns=['年份', '最大日降水量','年份', '最大日降水量','年份', '最大日降水量']
    
    # 填充表格数据
    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]='GD'
    max_values.iloc[1,0]='P-Ⅲ'
    creat_table(document,data_pre,'历年最大日降水量')

    creat_table(document,max_values,'日降水量（单位：mm）')

    document.save(report)
    
    
    return report

    
def re_pre_report_pg(pre_result,daily_df,methods,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_PRE_G.docx')
    doc=DocxTemplate(doc_path)
    

    station_name=daily_df['Station_Name'][0]

    return_years=pre_result['return_years']
    
    PRE_Max_Day=dict(pre_result['PRE_Max_Day'])
    img_save_path=dict(PRE_Max_Day['img_save_path'])
    
    return_result=dict(PRE_Max_Day['return_result'])
    max_values=dict(return_result['max_values'])
    max_values=pd.DataFrame(max_values)

    data=PRE_Max_Day['data']
    data=pd.DataFrame(data)



    dic=dict()
    
    dic['methods']=methods

    dic['station_name']=station_name

    dic['max_year']=data[data['最大日降水量(mm)']==data['最大日降水量(mm)'].max()].iloc[0,0]
    dic['max_year_pre']=data['最大日降水量(mm)'].max()
    dic['min_year']=data[data['最大日降水量(mm)']==data['最大日降水量(mm)'].min()].iloc[0,0]
    dic['min_year_pre']=data['最大日降水量(mm)'].min()
    
    dic['min_y']=min(return_years)
    dic['max_y']=max(return_years)

    
    years=data['年份'][:]
    mask = ~np.isnan(data['最大日降水量(mm)'])
    valid_years = data['年份'][mask]
    valid_preperatures = data['最大日降水量(mm)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_preperatures, 1)   
    if slope> 0:
        dic['slope']='上升'
    else:
        dic['slope']='下降'
    
    plt.figure(figsize=(10, 6))
    plt.bar(data['年份'], data['最大日降水量(mm)'], width=0.4, color='skyblue', label='最大日降水量')
    plt.plot(years, slope * years + intercept, color='red', label='最大日降水量')

    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('年')
    plt.ylabel('最大日降水量（mm）')
    plt.xticks(years[::3])
    # plt.ylim(dic['average_pre_1']-5, dic['average_pre_2']+5)
    plt.legend()
    
    pre_picture=os.path.join(data_dir,'最大日降水量年际变化图.png')
    plt.savefig(pre_picture, bbox_inches='tight', dpi=200)
    plt.clf()
    plt.close('all')
    dic['pre_picture'] = InlineImage(doc, pre_picture, width=Mm(130))
    
    idmin=return_years.index(min(return_years))
    idmmax=return_years.index(max(return_years))
    
    dic['min_pre_g'] = max_values.iloc[idmin,0]
    dic['max_pre_g'] = max_values.iloc[idmmax,0]

    
    dic['pre_50']=max_values.iloc[return_years.index(50),0]
    dic['pre_100']=max_values.iloc[return_years.index(100),0]        

    if methods=='Gumbel':
        dic['picture_gd'] = InlineImage(doc, img_save_path['Gumbel_plot'], width=Mm(130))
    else:
        dic['picture_gd'] = InlineImage(doc, img_save_path['P3_plot'], width=Mm(130))
       

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_PRE.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    
    data_pre=pd.DataFrame(pre_result['PRE_Max_Day']['data'])
    data_pre['年份']=data_pre['年份'].astype(str)

    half_length =int(np.ceil(len(data_pre)/3))
   
    # 分割每个列并创建新列
    data_pre['A1'] = data_pre.iloc[:half_length:,0]
    data_pre['B1'] =data_pre.iloc[:half_length:,1]
    data_pre['A2'] = data_pre.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    data_pre['B2'] =data_pre.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    data_pre['A3'] =data_pre.iloc[half_length*2::,0].reset_index(drop=True)
    data_pre['B3'] = data_pre.iloc[half_length*2::,1].reset_index(drop=True)

    data_pre.drop(['年份', '最大日降水量(mm)'], axis=1, inplace=True)
    
    data_pre.dropna(how='all',inplace=True)
    data_pre.columns=['年份', '最大日降水量(mm)','年份', '最大日降水量(mm)','年份', '最大日降水量(mm)']
    
    # 填充表格数据
    max_values=max_values.T.reset_index()
    max_values.columns=['重现期']+[str(year)+'a' for year in return_years]
    max_values.iloc[0,0]=methods
    creat_table(document,data_pre,'历年最大日降水量')

    creat_table(document,max_values,'日降水量（单位：mm）')

    document.save(report)
    
    
    return report

if __name__ == '__main__':

    daily_df=df_sequence
    data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'
    methods=fitting_method[0]
    report=re_pre_report(pre_result,daily_df,methods,data_dir)