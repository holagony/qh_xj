# -*- coding: utf-8 -*-
"""
Created on Tue Jun 11 13:31:10 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# df_sequence=df_sequence
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'

def re_tem_report(result_dict,df_sequence,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_TEM.docx')
    doc=DocxTemplate(doc_path)
    

    station_name=df_sequence['Station_Name'][0]
    
    return_years=result_dict['return_years']
    
    max_tem=result_dict['max_tem']
    data_max_tem=pd.DataFrame(max_tem['data'])
    return_result_max_tem=pd.DataFrame(max_tem['return_result']['max_values'])
    img_save_path_max_tem=max_tem['img_save_path']
    
    min_tem=result_dict['min_tem']
    data_min_tem=pd.DataFrame(min_tem['data'])
    return_result_min_tem=pd.DataFrame(min_tem['return_result']['max_values'])
    img_save_path_min_tem=min_tem['img_save_path']
    
    base_tem_max=result_dict['base_tem_max']
    data_base_tem_max=pd.DataFrame(base_tem_max['data'])
    return_result_data_base_tem_max=pd.DataFrame(base_tem_max['return_result']['max_values'])
    
    base_tem_min=result_dict['base_tem_min']
    data_base_tem_min=pd.DataFrame(base_tem_min['data'])
    return_result_data_base_tem_min=pd.DataFrame(base_tem_min['return_result']['max_values'])



    dic=dict()
    

    dic['station_name']=station_name

    dic['maxtem_50_G']=return_result_max_tem.iloc[return_years.index(50),0]
    dic['maxtem_100_G']=return_result_max_tem.iloc[return_years.index(100),0]   
    dic['maxtem_50_P']=return_result_max_tem.iloc[return_years.index(50),1]
    dic['maxtem_100_P']=return_result_max_tem.iloc[return_years.index(100),1]         
        
    dic['mintem_50_G']=return_result_min_tem.iloc[return_years.index(50),0]
    dic['mintem_100_G']=return_result_min_tem.iloc[return_years.index(100),0]   
    dic['mintem_50_P']=return_result_min_tem.iloc[return_years.index(50),1]
    dic['mintem_100_P']=return_result_min_tem.iloc[return_years.index(100),1]   

    if dic['maxtem_50_G']>= dic['maxtem_50_P']:
        dic['g_or_p']='耿贝尔法'
      
    else:
        dic['g_or_p']='皮尔逊-Ⅲ法'


    dic['picture_gd_max'] = InlineImage(doc, img_save_path_max_tem['Gumbel_plot'], width=Mm(130))
    dic['picture_p_max'] = InlineImage(doc, img_save_path_max_tem['P3_plot'], width=Mm(130))
    dic['picture_gd_min'] = InlineImage(doc, img_save_path_min_tem['Gumbel_plot'], width=Mm(130))
    dic['picture_p_min'] = InlineImage(doc, img_save_path_min_tem['P3_plot'], width=Mm(130))
    

    
    dic['base_maxtem_50_G']=return_result_data_base_tem_max.iloc[return_years.index(50),0]
    dic['base_maxtem_50_P']=return_result_data_base_tem_max.iloc[return_years.index(50),1]   
    dic['base_mintem_50_G']=return_result_data_base_tem_min.iloc[return_years.index(50),0]
    dic['base_mintem_50_P']=return_result_data_base_tem_min.iloc[return_years.index(50),1]   

    if dic['base_maxtem_50_G']>= dic['base_maxtem_50_P']:
        dic['base_g_or_p']='耿贝尔法'
      
    else:
        dic['base_g_or_p']='皮尔逊-Ⅲ法'

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_TEM.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    result_tem=pd.merge(data_max_tem,data_min_tem)
    # result_tem.drop(result_tem.tail(1).index, inplace=True)

    half_length =int(np.ceil(len(result_tem)/2))

    # 分割每个列并创建新列
    result_tem['年份']=result_tem['年份'].astype(str)
    result_tem['年份1'] = result_tem.iloc[:half_length:,0]
    result_tem['最高气温'] =result_tem.iloc[:half_length:,1]
    result_tem['最低气温'] = result_tem.iloc[:half_length:,2]
    result_tem['年份2'] =result_tem.iloc[half_length::,0].reset_index(drop=True)
    result_tem['最高气温1'] =result_tem.iloc[half_length::,1].reset_index(drop=True)
    result_tem['最低气温1'] = result_tem.iloc[half_length::,2].reset_index(drop=True)
    result_tem.drop(['年份', '极端最高气温(°C)', '极端最低气温(°C)'], axis=1, inplace=True)
    
    result_tem.dropna(how='all',inplace=True)
    result_tem.rename(columns={'年份1': '年份','年份2': '年份', '最高气温1': '最高气温', '最低气温1': '最低气温'}, inplace=True)

    
    return_result_max_tem=return_result_max_tem.T.reset_index()
    return_result_max_tem.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_max_tem.iloc[0,0]='GD'
    return_result_max_tem.iloc[1,0]='P-Ⅲ'
    
    return_result_min_tem=return_result_min_tem.T.reset_index()
    return_result_min_tem.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_min_tem.iloc[0,0]='GD'
    return_result_min_tem.iloc[1,0]='P-Ⅲ'
    
    result_base_tem=pd.DataFrame(columns=['最高温度月','基本最高气温','最低温度月','基本最低气温'])
    result_base_tem['最高温度月']=data_base_tem_max.iloc[:,0]
    result_base_tem['基本最高气温']=data_base_tem_max.iloc[:,1]
    result_base_tem['最低温度月']=data_base_tem_min.iloc[:,0]
    result_base_tem['基本最低气温']=data_base_tem_min.iloc[:,1]
    
    return_result_data_base_tem_max=return_result_data_base_tem_max.T.reset_index()
    return_result_data_base_tem_max.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_data_base_tem_max.iloc[0,0]='GD'
    return_result_data_base_tem_max.iloc[1,0]='P-Ⅲ'
    
    return_result_data_base_tem_min=return_result_data_base_tem_min.T.reset_index()
    return_result_data_base_tem_min.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_data_base_tem_min.iloc[0,0]='GD'
    return_result_data_base_tem_min.iloc[1,0]='P-Ⅲ'
    
    creat_table(document,result_tem,'最高气温和极端最低气温（单位：℃）')
    creat_table(document,return_result_max_tem,'不同重现期极端高温（单位：℃）')
    creat_table(document,return_result_min_tem,'不同重现期极端低温（单位：℃）')
    creat_table(document,result_base_tem,'月平均最高气温和最低温度月的月平均最低气温（单位：℃）')
    creat_table(document,return_result_data_base_tem_max,'不同重现期月平均最高气温（单位：℃）')
    creat_table(document,return_result_data_base_tem_min,'不同重现期月平均最低气温（单位：℃）')

    document.save(report)
    
    
    return report

def re_tem_report_pg(result_dict,df_sequence,methods,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module04','RE_TEM_P.docx')
    doc=DocxTemplate(doc_path)
    

    station_name=df_sequence['Station_Name'][0]
    
    return_years=result_dict['return_years']
    
    max_tem=result_dict['max_tem']
    data_max_tem=pd.DataFrame(max_tem['data'])
    return_result_max_tem=pd.DataFrame(max_tem['return_result']['max_values'])
    img_save_path_max_tem=max_tem['img_save_path']
    
    min_tem=result_dict['min_tem']
    data_min_tem=pd.DataFrame(min_tem['data'])
    return_result_min_tem=pd.DataFrame(min_tem['return_result']['max_values'])
    img_save_path_min_tem=min_tem['img_save_path']
    
    base_tem_max=result_dict['base_tem_max']
    data_base_tem_max=pd.DataFrame(base_tem_max['data'])
    return_result_data_base_tem_max=pd.DataFrame(base_tem_max['return_result']['max_values'])
    
    base_tem_min=result_dict['base_tem_min']
    data_base_tem_min=pd.DataFrame(base_tem_min['data'])
    return_result_data_base_tem_min=pd.DataFrame(base_tem_min['return_result']['max_values'])



    dic=dict()
    
    dic['methods']=methods

    dic['station_name']=station_name

    dic['maxtem_50_G']=return_result_max_tem.iloc[return_years.index(50),0]
    dic['maxtem_100_G']=return_result_max_tem.iloc[return_years.index(100),0]   
        
        
    dic['mintem_50_G']=return_result_min_tem.iloc[return_years.index(50),0]
    dic['mintem_100_G']=return_result_min_tem.iloc[return_years.index(100),0]   
  

    if methods=='Gumbel':
        dic['picture_gd_max'] = InlineImage(doc, img_save_path_max_tem['Gumbel_plot'], width=Mm(130))
        dic['picture_gd_min'] = InlineImage(doc, img_save_path_min_tem['Gumbel_plot'], width=Mm(130))

    else:
        dic['picture_gd_max'] = InlineImage(doc, img_save_path_max_tem['P3_plot'], width=Mm(130))
        dic['picture_gd_min'] = InlineImage(doc, img_save_path_min_tem['P3_plot'], width=Mm(130))
    
    dic['base_maxtem_50_G']=return_result_data_base_tem_max.iloc[return_years.index(50),0]
    dic['base_mintem_50_G']=return_result_data_base_tem_min.iloc[return_years.index(50),0]


    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RE_TEM.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    result_tem=pd.merge(data_max_tem,data_min_tem)
    # result_tem.drop(result_tem.tail(1).index, inplace=True)

    half_length =int(np.ceil(len(result_tem)/2))

    # 分割每个列并创建新列
    result_tem['年份']=result_tem['年份'].astype(str)
    result_tem['年份1'] = result_tem.iloc[:half_length:,0]
    result_tem['最高气温'] =result_tem.iloc[:half_length:,1]
    result_tem['最低气温'] = result_tem.iloc[:half_length:,2]
    result_tem['年份2'] =result_tem.iloc[half_length::,0].reset_index(drop=True)
    result_tem['最高气温1'] =result_tem.iloc[half_length::,1].reset_index(drop=True)
    result_tem['最低气温1'] = result_tem.iloc[half_length::,2].reset_index(drop=True)
    result_tem.drop(['年份', '极端最高气温(°C)', '极端最低气温(°C)'], axis=1, inplace=True)
    
    result_tem.dropna(how='all',inplace=True)
    result_tem.rename(columns={'年份1': '年份','年份2': '年份', '最高气温1': '最高气温', '最低气温1': '最低气温'}, inplace=True)

    
    return_result_max_tem=return_result_max_tem.T.reset_index()
    return_result_max_tem.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_max_tem.iloc[0,0]=methods
    
    return_result_min_tem=return_result_min_tem.T.reset_index()
    return_result_min_tem.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_min_tem.iloc[0,0]=methods
    
    result_base_tem=pd.DataFrame(columns=['最高温度月','基本最高气温','最低温度月','基本最低气温'])
    result_base_tem['最高温度月']=data_base_tem_max.iloc[:,0]
    result_base_tem['基本最高气温']=data_base_tem_max.iloc[:,1]
    result_base_tem['最低温度月']=data_base_tem_min.iloc[:,0]
    result_base_tem['基本最低气温']=data_base_tem_min.iloc[:,1]
    
    return_result_data_base_tem_max=return_result_data_base_tem_max.T.reset_index()
    return_result_data_base_tem_max.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_data_base_tem_max.iloc[0,0]=methods
    
    return_result_data_base_tem_min=return_result_data_base_tem_min.T.reset_index()
    return_result_data_base_tem_min.columns=['重现期']+[str(year)+'a' for year in return_years]
    return_result_data_base_tem_min.iloc[0,0]=methods
    
    creat_table(document,result_tem,'最高气温和极端最低气温（单位：℃）')
    creat_table(document,return_result_max_tem,'不同重现期极端高温（单位：℃）')
    creat_table(document,return_result_min_tem,'不同重现期极端低温（单位：℃）')
    creat_table(document,result_base_tem,'月平均最高气温和最低温度月的月平均最低气温（单位：℃）')
    creat_table(document,return_result_data_base_tem_max,'不同重现期月平均最高气温（单位：℃）')
    creat_table(document,return_result_data_base_tem_min,'不同重现期月平均最低气温（单位：℃）')

    document.save(report)
    
    
    return report

if __name__ == '__main__':
    
    df_sequence=df_sequence
    data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'

    methods=fitting_method[0]
    result_dict=tem_result
    report=re_tem_report_pg(result_dict,df_sequence,methods,data_dir)  
    
