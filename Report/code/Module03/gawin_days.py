# -*- coding: utf-8 -*-
"""
Created on Wed Jun  5 16:32:30 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)



# main_sta_ids='54823'
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module03'

def gawin_days_report(table1,table2,table3,month_data,data_dir,main_sta_ids):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module03','GAWIN_DAYS.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=month_data[month_data['Station_Id_C']==main_sta_ids].iloc[0,1]

    table1=pd.DataFrame(table1)
    table2=pd.DataFrame(table2)
    table3=pd.DataFrame(table3)
    
    
    table1u=table1[[station_name+'站','年份']]
    table2u=table2[table2['站点']==station_name]
    table3u=table3[table3['站名']==station_name]
    
    
    dic=dict()
    
    # 年际变化
    dic['station_name']=station_name
    dic['average_pre']=table2u.iloc[0,2]

    
    dic['average_pre_2']=int(table2u.iloc[0,3])
    dic['average_pre_2_year']=table2u.iloc[0,4]
    dic['average_pre_1']=int(table2u.iloc[0,5])
    dic['average_pre_1_year']=table2u.iloc[0,6]

    
    # 图像绘制
    years=table1u['年份'][:]
    mask = ~np.isnan(table1u[station_name+'站'])
    valid_years =table1u['年份'][mask]
    valid_preperatures = table1u[station_name+'站'][mask]
    slope, intercept = np.polyfit(valid_years, valid_preperatures, 1)   
    if slope> 0:
        dic['average_pre_slope']='上升'
    else:
        dic['average_pre_slope']='下降'

    average_pre_picture_hournum=plot_picture(table1u, '年份',station_name+'站','大风日数（d）','d','历年平均大风日数变化.png',10,10,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(years,table1u[station_name+'站'], width=0.4, color='skyblue', label='大风日数')
    # plt.plot(years, slope * years + intercept, color='red', label='线性（大风日数）')

    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('年')
    # plt.ylabel('大风日数（d）')
    # plt.xticks(years[::3])
    # # plt.ylim(dic['average_pre_1'], dic['average_pre_2']+10)
    # plt.legend()
    
    # average_pre_picture_hournum=os.path.join(data_dir,'历年平均大风日数变化.png')
    # plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
    # plt.clf()
    # plt.close('all')
    dic['average_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))

    
    # 逐月变化
    dic['average_pre_m1']=table3u[table3u.iloc[1,2:-1:].astype(float).idxmin()].iloc[1]
    # dic['average_pre_m1_month']=table3u.iloc[1,2:-1:].astype(float).idxmin()
    dic['average_pre_m2']=table3u[table3u.iloc[1,2:-1:].astype(float).idxmax()].iloc[1]
    # dic['average_pre_m2_month']=table3u.iloc[1,2:-1:].astype(float).idxmax()
    
    # 计算春夏秋冬
    spring = table3u.iloc[1,4:7] 
    summer = table3u.iloc[1,7:10]  
    autumn = table3u.iloc[1,10:13] 
    winter = pd.concat([table3u.iloc[1,2:4],table3u.iloc[1,13:14]])  
    
    average_spring = sum(spring) / len(spring)
    average_summer = sum(summer) / len(summer)
    average_autumn = sum(autumn) / len(autumn)
    average_winter = sum(winter) / len(winter)
    
    seasons = {
    "春季": average_spring,
    "夏季": average_summer,
    "秋季": average_autumn,
    "冬季": average_winter}
    
    sorted_seasons = sorted(seasons.items(), key=lambda x: x[1], reverse=True)

    dic['season1'] = sorted_seasons[0][0]  
    dic['season2'] = sorted_seasons[1][0]  
    dic['season3'] = sorted_seasons[-1][0]  

    # 图像绘制
    months=table3u.columns[2::]
    average_pre_picture_month=plot_picture_2(months,table3u.iloc[1,2::],dic,'平均大风日数（d）','average_pre_m1','average_pre_m2','平均大风日数逐月变化.png',100,2,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(months, table3u.iloc[1,2::], width=0.4, color='skyblue')
    
    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('月')
    # plt.ylabel('平均大风日数（d）')
    # plt.xticks(months)
    # # plt.ylim(dic['average_pre_m1'], dic['average_pre_m2']+0.5)
    
    # average_pre_picture_month=os.path.join(data_dir,'平均大风日数逐月变化.png')
    # plt.savefig(average_pre_picture_month, bbox_inches='tight', dpi=200)
    # plt.clf()
    plt.close('all')
    
    dic['average_picture_m'] = InlineImage(doc, average_pre_picture_month, width=Mm(130))

    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'GAWIN_DAYS.docx')
    doc.save(report)
    
    document = Document(report)
    
    
    columns = table1u.columns.tolist()
    columns = [columns[1]] + columns[:1] + columns[2:]    
    table1uu = table1u[columns]

    
    creat_table(document,table1uu,'数年变化（单位：d）')
    creat_table(document,table3u,'数月变化（单位：d）')

    document.save(report)
    
    
    return report

    
