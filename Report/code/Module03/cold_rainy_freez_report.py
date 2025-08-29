# -*- coding: utf-8 -*-
"""
Created on Fri Jun  7 11:39:47 2024

@author: EDY
"""
import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
import os
from Utils.config import cfg
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            

            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)



# main_sta_ids='54823'
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module03'

def cold_rainy_freez_report(cold_rainy, cold_freezing, day_data,data_dir,main_sta_ids):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module03','cold_rainy_freez.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=day_data[day_data['Station_Id_C']==main_sta_ids].iloc[0,1]

    cold_rainy=pd.DataFrame(cold_rainy)
    cold_freezing=pd.DataFrame(cold_freezing)
    
    
    cold_rainyu=cold_rainy[cold_rainy['站名']==station_name]
    cold_freezingu=cold_freezing[cold_freezing['站名']==station_name]
    
    result_rain=pd.DataFrame(columns=['年','过程天数'])
    result_rain2=pd.DataFrame(columns=['年','过程天数'])

    for i in np.arange(len(cold_rainyu)):
        result_rain.at[i,'年']=cold_rainyu.iloc[i,1][:4:]
        result_rain.at[i,'过程天数']=cold_rainyu.iloc[i,6]

    year=result_rain['年'].unique()
    for i in np.arange(len(year)):
        result_rain2.at[i,'年']=year[i]
        result_rain2.at[i,'过程天数']=result_rain[result_rain['年']==year[i]]['过程天数'].sum()
        
        
    result_freez=pd.DataFrame(columns=['年','开始日降水量','结束日降水量','过程天数'])
    result_freez2=pd.DataFrame(columns=['年','过程天数'])

    for i in np.arange(len(cold_freezingu)):
        result_freez.at[i,'年']=cold_freezingu.iloc[i,1][:4:]
        result_freez.at[i,'过程天数']=cold_freezingu.iloc[i,7]
        result_freez.at[i,'开始日降水量']=cold_freezingu.iloc[i,5]
        result_freez.at[i,'结束日降水量']=cold_freezingu.iloc[i,6]

    year=result_freez['年'].unique()
    for i in np.arange(len(year)):
        result_freez2.at[i,'年']=year[i]
        result_freez2.at[i,'过程天数']=result_freez[result_freez['年']==year[i]]['过程天数'].sum()
            
        
    
    dic=dict()
    
    # 年际变化
    dic['station_name']=station_name
    dic['start_year']=day_data.index.year[0]
    dic['end_year']=day_data.index.year[-1]

    dic['average_cold_rain']=round(result_rain2['过程天数'].sum()/(dic['end_year']-dic['start_year']+1),2)
    
    dic['min_cold_rain']=result_rain2['过程天数'].min()
    min_years = result_rain2[result_rain2['过程天数'] ==result_rain2['过程天数'].min()]['年'].tolist()
    min_years_str = '、'.join(map(str, min_years))
    dic['min_cold_rain_year']=min_years_str
    dic['max_cold_rain']=result_rain2['过程天数'].max()
    max_years = result_rain2[result_rain2['过程天数'] ==result_rain2['过程天数'].max()]['年'].tolist()
    max_years_str = '、'.join(map(str, max_years))
    dic['max_cold_rain_year']=max_years_str

    dic['average_cold_freez']=round(result_freez2['过程天数'].sum()/(dic['end_year']-dic['start_year']+1),2)

    dic['min_cold_freez']=result_freez2['过程天数'].min()
    min_years = result_freez2[result_freez2['过程天数'] ==result_freez2['过程天数'].min()]['年'].tolist()
    min_years_str = '、'.join(map(str, min_years))
    dic['min_cold_freez_year']=min_years_str
    dic['max_cold_freez']=result_freez2['过程天数'].max()
    max_years = result_freez2[result_freez2['过程天数'] ==result_freez2['过程天数'].max()]['年'].tolist()
    max_years_str = '、'.join(map(str, max_years))
    dic['max_cold_freez_year']=max_years_str
    
    dic['min_start_rain']=result_freez['开始日降水量'].min()
    min_years = result_freez[result_freez['开始日降水量'] ==result_freez['开始日降水量'].min()]['年'].tolist()
    min_years_str = '、'.join(map(str, min_years))
    dic['min_start_rain_year']=min_years_str
    dic['max_start_rain']=result_freez['开始日降水量'].max()
    max_years = result_freez[result_freez['开始日降水量'] ==result_freez['开始日降水量'].max()]['年'].tolist()
    max_years_str = '、'.join(map(str, max_years))
    dic['max_start_rain_year']=max_years_str
    
    dic['min_end_rain']=result_freez['结束日降水量'].min()
    min_years = result_freez[result_freez['结束日降水量'] ==result_freez['结束日降水量'].min()]['年'].tolist()
    min_years_str = '、'.join(map(str, min_years))
    dic['min_end_rain_year']=min_years_str
    dic['max_end_rain']=result_freez['结束日降水量'].max()
    max_years = result_freez[result_freez['结束日降水量'] ==result_freez['结束日降水量'].max()]['年'].tolist()
    max_years_str = '、'.join(map(str, max_years))
    dic['max_end_rain_year']=max_years_str
    

 
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'cold_rainy_freez.docx')
    doc.save(report)
    
    document = Document(report)
    
    
    creat_table(document,cold_rainyu,'阴雨日参数统计表')
    creat_table(document,cold_freezingu,'冰冻日参数统计表')

    document.save(report)
    
    return report

    
