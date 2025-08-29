# -*- coding: utf-8 -*-
"""
Created on Thu Jun  6 14:38:03 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            

            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)



# main_sta_ids='54823'
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module03'

def cold_wave_report(cold_wave_result, cold_wave_wind, cold_wave_wind_d,day_data,data_dir,main_sta_ids):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module03','cold_wave.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=day_data[day_data['Station_Id_C']==main_sta_ids].iloc[0,1]

    cold_wave_result=pd.DataFrame(cold_wave_result)
    cold_wave_wind=pd.DataFrame(cold_wave_wind)
    cold_wave_wind_d=pd.DataFrame(cold_wave_wind_d)
    
    
    cold_wave_resultu=cold_wave_result[cold_wave_result['站名']==station_name]
    cold_wave_windu=cold_wave_wind[cold_wave_wind['站名']==station_name]
    cold_wave_wind_du=cold_wave_wind_d[cold_wave_wind_d['站名']==station_name]
    
    result=pd.DataFrame(columns=['年','寒潮日','温度降幅'])
    result2=pd.DataFrame(columns=['年','寒潮日','温度降幅'])

    for i in np.arange(len(cold_wave_resultu)):
        result.at[i,'年']=cold_wave_resultu.iloc[i,1][:4:]
        result.at[i,'寒潮日']=cold_wave_resultu.iloc[i,7]
        result.at[i,'温度降幅']=cold_wave_resultu.iloc[i,8]

    year=result['年'].unique()
    for i in np.arange(len(year)):
        result2.at[i,'年']=year[i]
        result2.at[i,'寒潮日']=result[result['年']==year[i]]['寒潮日'].sum()
        result2.at[i,'温度降幅']=result[result['年']==year[i]]['温度降幅'].mean()
    
    dic=dict()
    
    # 年际变化
    dic['station_name']=station_name
    dic['start_year']=day_data.index.year[0]
    dic['end_year']=day_data.index.year[-1]

    dic['average_day']=round(cold_wave_resultu['寒潮天数'].sum()/(dic['end_year']-dic['start_year']+1),2)

    dic['min_day']=result2['寒潮日'].min()
    dic['max_day']=result2['寒潮日'].max()
    
    
    dic['average_tem']=round(result2['温度降幅'].mean(),2)
    dic['min_tem']=result['温度降幅'].min()
    min_years = result[result['温度降幅'] ==result['温度降幅'].min()]['年'].tolist()
    min_years_str = '、'.join(map(str, min_years))
    dic['min_tem_year']=min_years_str
    dic['max_tem']=result['温度降幅'].max()
    max_years = result[result['温度降幅'] ==result['温度降幅'].max()]['年'].tolist()
    max_years_str = '、'.join(map(str, max_years))
    dic['max_tem_year']=max_years_str

    dic['num']=len(cold_wave_windu)
  

    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'cold_wave.docx')
    doc.save(report)
    
    document = Document(report)
    
    
    creat_table(document,cold_wave_resultu,'过程统计表')
    creat_table(document,cold_wave_windu,'大风统计表')
    creat_table(document,cold_wave_wind_du,'风向频数统计表')

    document.save(report)
    
    return report

    
