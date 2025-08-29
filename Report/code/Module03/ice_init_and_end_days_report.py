# -*- coding: utf-8 -*-
"""
Created on Thu Jun  6 14:15:13 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# main_sta_ids='54823'
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module03'

def ice_init_and_end_days_report(table4,table5,day_data,data_dir,main_sta_ids):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module03','ice_init_and_end_days_report.docx')
    doc=DocxTemplate(doc_path)
    
    station_name=day_data[day_data['Station_Id_C']==main_sta_ids].iloc[0,1]

    table4=pd.DataFrame(table4)
    table5=pd.DataFrame(table5)
    
    table4u_columns = [col for col in table4.columns if station_name[:-7:] in col]#['时间段'] + [col for col in table4.columns if station_name[:-7:] in col]
    table4u= table4[table4u_columns]
    table5u_columns = [col for col in table5.columns if station_name[:-7:] in col]#['统计'] + [col for col in table5.columns if station_name[:-7:] in col]
    table5u= table5[table5u_columns]
    
    
    dic=dict()
    

    dic['station_name']=station_name
    dic['start_year']=day_data.index.year.unique()[0]
    dic['end_year']=day_data.index.year.unique()[-1]
    dic['average_frost']=round(np.nanmean(table4u.iloc[:,1].astype(float)),2)

    
    dic['start_frost_average_mon']=table5u.iloc[0,1].split('-')[0]
    dic['start_frost_average_day']=table5u.iloc[0,1].split('-')[1]
    dic['start_frost_start_mon']=table5u.iloc[1,1].split('-')[1]
    dic['start_frost_start_day']=table5u.iloc[1,1].split('-')[2]
    dic['start_frost_start_year']=table5u.iloc[1,1].split('-')[0]
    dic['start_frost_end_mon']=table5u.iloc[2,1].split('-')[1]
    dic['start_frost_end_day']=table5u.iloc[2,1].split('-')[2]
    dic['start_frost_end_year']=table5u.iloc[2,1].split('-')[0]
    dic['end_frost_average_mon']=table5u.iloc[0,2].split('-')[0]
    dic['end_frost_average_day']=table5u.iloc[0,2].split('-')[0]
    dic['end_frost_start_mon']=table5u.iloc[1,2].split('-')[1]
    dic['end_frost_start_day']=table5u.iloc[1,2].split('-')[2]
    dic['end_frost_start_year']=table5u.iloc[1,2].split('-')[0]
    dic['end_frost_end_mon']=table5u.iloc[2,2].split('-')[1]
    dic['end_frost_end_day']=table5u.iloc[2,2].split('-')[2]
    dic['end_frost_end_year']=table5u.iloc[2,2].split('-')[0]
    dic['average_jian_num']=table5u.iloc[0,3]
    dic['max_jian_num']=table5u.iloc[1,3].split('(')[0]
    dic['max_jian_num_year1']=table5u.iloc[1,3].split('(')[1].split('.')[0]
    dic['max_jian_num_year2']=table5u.iloc[1,3].split('-')[1].split('.')[0]
    dic['min_jian_num']=table5u.iloc[2,3].split('(')[0]
    dic['min_jian_num_year1']=table5u.iloc[2,3].split('(')[1].split('.')[0]
    dic['min_jian_num_year2']=table5u.iloc[2,3].split('-')[1].split('.')[0]


    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'ice_init_and_end_days_report.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,table4u,'平均结冰日参数表')
    creat_table(document,table5u,'平均结冰日参数统计表')

    document.save(report)
    
    return report

    
