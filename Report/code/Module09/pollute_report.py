# -*- coding: utf-8 -*-
"""
Created on Mon Jul 29 10:33:12 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                 
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module09'

def pollute_report(post_monthly_df,p_c, depth_mixed_accum, ven_ability_accum, data_asc_accum, data_asi_accum,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module09','pollute_report.docx')
    doc=DocxTemplate(doc_path)
    
    dic=dict()
    dic['station_name']=post_monthly_df.iloc[0,1]
    dic['start_year']=post_monthly_df.index.year[0]
    dic['end_year']=post_monthly_df.index.year[-1]
    
    # 污染系数
    pc_data=p_c.iloc[50,2::].apply(pd.to_numeric, errors='coerce')
    dic['pc_1']=pc_data.max()
    dic['pc_2']=pc_data.idxmax()
    dic['pc_3']=pc_data.min()
    dic['pc_4']=pc_data.idxmin()
  
    # 混合层高度
    depth_mixed_data=depth_mixed_accum.iloc[16,1::]
    dic['depth_mixed']='m、'.join(map(str,depth_mixed_data))

    # 通风量
    ven_ability_data=ven_ability_accum.iloc[16,1::]
    dic['ven_ability']='m²/s、'.join(map(str,ven_ability_data))
    
    # 大气自净能力ASC
    data_asc_data=data_asc_accum.iloc[16,7::]
    dic['data_asc']='、'.join(map(str,data_asc_data))    
    
    # 大气自净能力ASI
    data_asi_data=data_asi_accum.iloc[16,1::]
    dic['data_asi']='t/(d·km²)、'.join(map(str,data_asi_data))    

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'pollute_report.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    creat_table(document,p_c,'污染系数统计表')
    creat_table(document,depth_mixed_accum,'混合层厚度表')
    creat_table(document,ven_ability_accum,'通风量表')
    creat_table(document,data_asc_accum,'大气自净能力ASC表')
    creat_table(document,data_asi_accum,'大气自净能力指数ASI表')


    document.save(report)
    
    
    return report
