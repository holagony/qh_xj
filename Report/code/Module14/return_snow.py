# -*- coding: utf-8 -*-
"""
Created on Mon Feb  2 10:52:04 2026

@author: hx
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2
from scipy import stats

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def pearson_r_sig(y, y_fit, alpha=0.05):
    y = np.asarray(y, dtype=float)
    y_fit = np.asarray(y_fit, dtype=float)
    mask = ~(np.isnan(y) | np.isnan(y_fit))
    y, y_fit = y[mask], y_fit[mask]
    n = len(y)
    if n < 3:
        return np.nan, None, False
    r = float(np.corrcoef(y, y_fit)[0, 1])
    df = n - 2
    t = r * np.sqrt(df / max(1e-12, 1 - r**2))
    p = float(2 * stats.t.sf(np.abs(t), df))
    return r, p, p < alpha

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)
    
def returnSnow_report(years,sta_ids,daily_df,r,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module14','Return.docx')
    
    doc=DocxTemplate(doc_path)
    dic=dict()
    
    sel_years = years.split(',')
    start_year = int(sel_years[0])
    end_year = int(sel_years[1])
    
    #%% Part 1
    dic['element_title']='最大积雪深度极值推算'
    dic['element_subtitle_1']='最大积雪深度'
    dic['start_year']=start_year
    dic['end_year']=end_year
    dic['station_name']=daily_df['Station_Name'][daily_df['Station_Id_C']==sta_ids].iloc[0]
    dic['unit']='mm'
    # 数据
    max_data=pd.DataFrame(r['main_return_result']['max_values'])
    column_max = max_data.max(axis=1)
    max_data.insert(0, '重现期', ['30年一遇','50年一遇','100年一遇'])    
    max_data.columns=['重现期', '极值I型', '皮尔逊Ⅲ型']
    
    dic['data_30']=column_max[0]
    dic['data_50']=column_max[1]
    dic['data_100']=column_max[2]

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'RETURN_Snow.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,max_data,f"气象站各重现期{dic['element_subtitle_1']}估计（{dic['unit']}）")

    document.save(report)

    return report
