# -*- coding: utf-8 -*-
"""
Created on Mon May 27 09:46:35 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2
from scipy import stats

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def pearson_r_sig(y, y_fit, alpha=0.05):
    y = np.asarray(y, dtype=float)
    y_fit = np.asarray(y_fit, dtype=float)
    mask = ~(np.isnan(y) | np.isnan(y_fit))
    y, y_fit = y[mask], y_fit[mask]
    n = len(y)
    if n < 3:
        return np.nan, None, False
    r = float(np.corrcoef(y, y_fit)[0, 1])
    df = n - 2
    t = r * np.sqrt(df / max(1e-12, 1 - r**2))
    p = float(2 * stats.t.sf(np.abs(t), df))
    return r, p, p < alpha

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)
    
# wind_dict = result.copy()
def win_report(years,sta_ids,sub_sta_ids,daily_df,wind_dict,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    if sub_sta_ids is not None:
        doc_path=os.path.join(cfg['report']['template'],'Module14','Base_stationall_template.docx')
    else:
        doc_path=os.path.join(cfg['report']['template'],'Module14','Base_stationall_template_one.docx')
    
    doc=DocxTemplate(doc_path)
    dic=dict()
    
    #%% Part 1
    dic['element_title']='风速'
    dic['element_subtitle_1']='最大风速'
    dic['start_year_1']=years[0]
    dic['end_year_1']=years[1]
    dic['station_name_1']=daily_df['Station_Name'][daily_df['Station_Id_C']==sta_ids].iloc[0]
    dic['station_name_1']=daily_df['Station_Name'][daily_df['Station_Id_C']==sta_ids].iloc[0]
    dic['unit']='m/s'
    # 年数据
    wind_max_year_main=wind_dict['历年最大风速'][['年份',sta_ids]]
    dic['average_1']=wind_max_year_main[sta_ids].mean().round(1)
    dic['max_year_1']=wind_max_year_main['年份'][wind_max_year_main[sta_ids].idxmax()]
    dic['max_data_1']=wind_max_year_main[sta_ids].max().round(1)
    # 年数据趋势
    mask = ~np.isnan(wind_max_year_main[sta_ids])
    valid_years = wind_max_year_main['年份'][mask]
    valid_vals = wind_max_year_main[sta_ids][mask]
    slope, intercept = np.polyfit(valid_years, valid_vals, 1)
    y_fit = slope * valid_years.astype(float) + intercept
    r, p, passed = pearson_r_sig(valid_vals, y_fit, alpha=0.05)
    dic['judge_1'] = '通过' if passed else '未通过'
    dic['slop_data_1'] = np.round(slope*10,2)
    if slope >0:
        dic['slop_1'] = '上升'

    else:
        dic['slop_1'] = '下降'
    
    if passed:
        dic['judge_expalin_1'] = f"{dic['slop_1']}趋势显著,每10a{dic['slop_1']}{abs(dic['slop_data_1'])}{dic['unit']}"
    else:
        dic['judge_expalin_1'] = f"{dic['slop_1']}趋势不显著"

    # 月数据
    win_max_month_main=wind_dict['累年各月平均最大风速'][wind_dict['累年各月平均最大风速']['站号']==sta_ids].iloc[:,1::]
    dic['max_month_1'] = win_max_month_main.idxmax(axis=1).iloc[0][0]
    dic['max_month_data_1'] = win_max_month_main.max(axis=1).iloc[0]
    # 图形绘制
    average_win_picture_year_1=plot_picture(wind_max_year_main, '年份',sta_ids,'最大风速(m/s)','（m/s）','最大风速逐年变化图.png',0.3,0.3,data_dir,R= round(float(r), 3))
    months = list(range(1, win_max_month_main.shape[1] + 1))
    max_win_picture_month_1 = plot_picture_2(months,win_max_month_main.iloc[0].values,dic,'最大风速(m/s)',0,0,'最大风速逐月变化图.png',0,0.5,data_dir,)
    dic['year_picture_1'] = InlineImage(doc, average_win_picture_year_1, width=Mm(130))
    dic['month_picture_1'] = InlineImage(doc, max_win_picture_month_1, width=Mm(130))

    dic['main_explain_1'] =''
    
    #%% Part 3
    dic['element_subtitle_2']='极大风速'
    # 年数据
    wind_inst_year_main=wind_dict['历年极大风速'][['年份',sta_ids]]
    dic['average_2']=wind_inst_year_main[sta_ids].mean().round(1)
    dic['max_year_2']=wind_inst_year_main['年份'][wind_inst_year_main[sta_ids].idxmax()]
    dic['max_data_2']=wind_inst_year_main[sta_ids].max().round(1)
    # 年数据趋势
    mask = ~np.isnan(wind_inst_year_main[sta_ids])
    valid_years = wind_inst_year_main['年份'][mask]
    valid_vals = wind_inst_year_main[sta_ids][mask]
    slope, intercept = np.polyfit(valid_years, valid_vals, 1)
    y_fit = slope * valid_years.astype(float) + intercept
    r, p, passed = pearson_r_sig(valid_vals, y_fit, alpha=0.05)
    dic['judge_2'] = '通过' if passed else '未通过'
    dic['slop_data_2'] = np.round(slope*10,2)
    if slope >0:
        dic['slop_2'] = '上升'
    else:
        dic['slop_2'] = '下降'

    if passed:
        dic['judge_expalin_2'] = f"{dic['slop_1']}趋势显著,每10a{dic['slop_1']}{abs(dic['slop_data_1'])}{dic['unit']}"
    else:
        dic['judge_expalin_2'] = f"{dic['slop_1']}趋势不显著"

    # 月数据
    win_max_month_main=wind_dict['累年各月平均极大风速'][wind_dict['累年各月平均极大风速']['站号']==sta_ids].iloc[:,1::]
    dic['max_month_2'] = win_max_month_main.idxmax(axis=1).iloc[0][0]
    dic['max_month_data_2'] = win_max_month_main.max(axis=1).iloc[0]
    # 图形绘制
    average_win_picture_year_2=plot_picture(wind_inst_year_main, '年份',sta_ids,'极大风速(m/s)','（m/s）','极大风速逐年变化图.png',0.3,0.3,data_dir,R= round(float(r), 3))
    months = list(range(1, win_max_month_main.shape[1] + 1))
    max_win_picture_month_2 = plot_picture_2(months,win_max_month_main.iloc[0].values,dic,'极大风速(m/s)',0,0,'极大风速逐月变化图.png',0,0.5,data_dir,)
    dic['year_picture_2'] = InlineImage(doc, average_win_picture_year_2, width=Mm(130))
    dic['month_picture_2'] = InlineImage(doc, max_win_picture_month_2, width=Mm(130))
    
    dic['main_explain_2'] =''    
    
    
    if sub_sta_ids is not None:
        
        #%% Part 2
        wind_max_year_sub=wind_dict['历年最大风速']
        sub_data = wind_max_year_sub.drop('年份', axis=1)
        year = wind_max_year_sub['年份']

        sub_data = sub_data.apply(pd.to_numeric, errors='coerce')
        sub_data_mean = sub_data.mean(axis=0, skipna=True).dropna()
        sub_data_mean = sub_data_mean.sort_values(ascending=False)
        sub_max_mean_col = sub_data_mean.index[0] if len(sub_data_mean) > 0 else None
        sub_max_mean_val = float(sub_data_mean.iloc[0]) if len(sub_data_mean) > 0 else np.nan
        sub_second_mean_col = sub_data_mean.index[1] if len(sub_data_mean) > 1 else None

        dic['station_name_1_1'] = daily_df['Station_Name'][daily_df['Station_Id_C']==sub_max_mean_col].iloc[0]
        dic['station_name_max_data_1'] = round(sub_max_mean_val, 1) if not np.isnan(sub_max_mean_val) else None
        dic['station_name_2_1'] = daily_df['Station_Name'][daily_df['Station_Id_C']==sub_second_mean_col].iloc[0]
        sub_trends = {}
        for _col in sub_data.columns:
            _vals = sub_data[_col]
            _mask = ~np.isnan(_vals.values)
            if _mask.sum() < 2:
                continue
            _years = year.values[_mask].astype(float)
            _y = _vals.values[_mask].astype(float)
            _slope, _intercept = np.polyfit(_years, _y, 1)
            sub_trends[_col] = '上升' if _slope > 0 else '下降'

        _up_list, _down_list = [], []
        for _sid, _trend in sub_trends.items():
            _name = daily_df['Station_Name'][daily_df['Station_Id_C'].astype(str) == str(_sid)]
            _name = _name.iloc[0] if len(_name) > 0 else str(_sid)
            if _trend == '上升':
                _up_list.append(_name)
            else:
                _down_list.append(_name)
        _up_n = len(_up_list)
        _down_n = len(_down_list)
        if _down_n == 0 and _up_n > 0:
            dic['sub_explain_1'] = "各气象站最大风速趋势均呈上升趋势。"
        elif _up_n == 0 and _down_n > 0:
            dic['sub_explain_1'] = "各气象站最大风速趋势均呈下降趋势。"
        else:
            dic['sub_explain_1'] = (
                f"各气象站最大风速趋势中，{'、'.join(_up_list)}气象站呈上升趋势，"
                f"下降站点为{'、'.join(_down_list)}气象站呈下降趋势。"
            )
        
        _station_cols = list(sub_data.columns)
        _bar_df = pd.concat([year, sub_data[_station_cols]], axis=1)
        _bar_years = _bar_df['年份'].astype(int).values
        _x = np.arange(len(_bar_years))
        _n = max(1, len(_station_cols))
        _width = 0.8 / _n
        _fig, _ax = plt.subplots(figsize=(12, 6))
        for _i, _c in enumerate(_station_cols):
            _y = pd.to_numeric(_bar_df[_c], errors='coerce').values.astype(float)
            _name_s = daily_df['Station_Name'][daily_df['Station_Id_C'].astype(str) == str(_c)]
            _label = _name_s.iloc[0] if len(_name_s) > 0 else str(_c)
            _ax.bar(_x + (_i - (_n - 1) / 2) * _width, _y, width=_width, label=_label)
        _ax.set_xlabel('年', fontsize=14)
        _ax.set_ylabel('最大风速(m/s)', fontsize=14)
        _ax.set_xticks(_x)
        _ax.set_xticklabels(_bar_years, rotation=90 if len(_bar_years) > 15 else 0, fontsize=12)
        _ax.tick_params(axis='y', labelsize=12)
        _ax.grid(axis='y', linestyle='--', alpha=0.3, color='gray')
        _ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.12), ncol=min(6, _n), frameon=False)
        _sub_year_bar_path = os.path.join(data_dir, '各站最大风速逐年对比柱状图.png')
        _fig.tight_layout()
        _fig.savefig(_sub_year_bar_path, bbox_inches='tight', dpi=200)
        plt.close(_fig)
        dic['average_picture_m_1'] = InlineImage(doc, _sub_year_bar_path, width=Mm(130))
    
        win_max_month_sub=wind_dict['累年各月平均最大风速']
        
        #%% Part4
        
        wind_inst_year_sub=wind_dict['历年极大风速']
        sub_data = wind_inst_year_sub.drop('年份', axis=1)
        year = wind_inst_year_sub['年份']

        sub_data = sub_data.apply(pd.to_numeric, errors='coerce')
        sub_data_mean = sub_data.mean(axis=0, skipna=True).dropna()
        sub_data_mean = sub_data_mean.sort_values(ascending=False)
        sub_inst_mean_col = sub_data_mean.index[0] if len(sub_data_mean) > 0 else None
        sub_inst_mean_val = float(sub_data_mean.iloc[0]) if len(sub_data_mean) > 0 else np.nan
        sub_second_mean_col = sub_data_mean.index[1] if len(sub_data_mean) > 1 else None

        dic['station_name_1_2'] = daily_df['Station_Name'][daily_df['Station_Id_C']==sub_inst_mean_col].iloc[0]
        dic['station_name_inst_data_2'] = round(sub_inst_mean_val, 1) if not np.isnan(sub_inst_mean_val) else None
        dic['station_name_2_2'] = daily_df['Station_Name'][daily_df['Station_Id_C']==sub_second_mean_col].iloc[0]
        sub_trends = {}
        for _col in sub_data.columns:
            _vals = sub_data[_col]
            _mask = ~np.isnan(_vals.values)
            if _mask.sum() < 2:
                continue
            _years = year.values[_mask].astype(float)
            _y = _vals.values[_mask].astype(float)
            _slope, _intercept = np.polyfit(_years, _y, 1)
            sub_trends[_col] = '上升' if _slope > 0 else '下降'

        _up_list, _down_list = [], []
        for _sid, _trend in sub_trends.items():
            _name = daily_df['Station_Name'][daily_df['Station_Id_C'].astype(str) == str(_sid)]
            _name = _name.iloc[0] if len(_name) > 0 else str(_sid)
            if _trend == '上升':
                _up_list.append(_name)
            else:
                _down_list.append(_name)
        _up_n = len(_up_list)
        _down_n = len(_down_list)
        if _down_n == 0 and _up_n > 0:
            dic['sub_explain_2'] = "各气象极大风速趋势均呈上升趋势。"
        elif _up_n == 0 and _down_n > 0:
            dic['sub_explain_2'] = "各气象站极大风速趋势均呈下降趋势。"
        else:
            dic['sub_explain_2'] = (
                f"各气象站极大风速趋势中，{'、'.join(_up_list)}气象站呈上升趋势，"
                f"下降站点为{'、'.join(_down_list)}气象站呈下降趋势。"
            )
        
        _station_cols = list(sub_data.columns)
        _bar_df = pd.concat([year, sub_data[_station_cols]], axis=1)
        _bar_years = _bar_df['年份'].astype(int).values
        _x = np.arange(len(_bar_years))
        _n = max(1, len(_station_cols))
        _width = 0.8 / _n
        _fig, _ax = plt.subplots(figsize=(12, 6))
        for _i, _c in enumerate(_station_cols):
            _y = pd.to_numeric(_bar_df[_c], errors='coerce').values.astype(float)
            _name_s = daily_df['Station_Name'][daily_df['Station_Id_C'].astype(str) == str(_c)]
            _label = _name_s.iloc[0] if len(_name_s) > 0 else str(_c)
            _ax.bar(_x + (_i - (_n - 1) / 2) * _width, _y, width=_width, label=_label)
        _ax.set_xlabel('年', fontsize=14)
        _ax.set_ylabel('极大风速(m/s)', fontsize=14)
        _ax.set_xticks(_x)
        _ax.set_xticklabels(_bar_years, rotation=90 if len(_bar_years) > 15 else 0, fontsize=12)
        _ax.tick_params(axis='y', labelsize=12)
        _ax.grid(axis='y', linestyle='--', alpha=0.3, color='gray')
        _ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.12), ncol=min(6, _n), frameon=False)
        _sub_year_bar_path = os.path.join(data_dir, '各站极大风速逐年对比柱状图.png')
        _fig.tight_layout()
        _fig.savefig(_sub_year_bar_path, bbox_inches='tight', dpi=200)
        plt.close(_fig)
        dic['average_picture_m_2'] = InlineImage(doc, _sub_year_bar_path, width=Mm(130))
    
        win_inst_month_sub=wind_dict['累年各月平均极大风速']
    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'WIND.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,win_max_month_sub,f"参证站和专用站{dic['element_subtitle_1']}统计（单位：{dic['unit']}）")
    creat_table(document,win_inst_month_sub,f"参证站和专用站{dic['element_subtitle_2']}统计（单位：{dic['unit']}）")

    document.save(report)

    return report
