# -*- coding: utf-8 -*-
"""
Created on Mon May 27 16:09:38 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# post_daily_df=vapor_day

def vapor_report(basic_vapor_yearly,basic_vapor_accum,post_daily_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module02','VAPOR.docx')
    doc=DocxTemplate(doc_path)
    
    
    basic_vapor_yearly=pd.DataFrame(basic_vapor_yearly)
    basic_vapor_accum=pd.DataFrame(basic_vapor_accum)
    
    
    dic=dict()
    
    ## 水汽压
    # 年际变化
    dic['num_years']=len(basic_vapor_yearly)
    dic['station_name']=post_daily_df['Station_Name'][0]
    dic['start_year']=basic_vapor_yearly['年份'][0]
    dic['end_year']=basic_vapor_yearly.iloc[-1,0]
    dic['average_vapor']=basic_vapor_accum['年'][0]
    dic['max_vapor']=basic_vapor_accum['年'][1]
    dic['min_vapor']=basic_vapor_accum['年'][4]
    
    dic['average_vapor_1']=basic_vapor_yearly['平均水汽压(hPa)'][basic_vapor_yearly['平均水汽压(hPa)'].astype(float).idxmin()]
    dic['average_vapor_1_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['平均水汽压(hPa)'].astype(float).idxmin()]
    dic['average_vapor_2']=basic_vapor_yearly['平均水汽压(hPa)'][basic_vapor_yearly['平均水汽压(hPa)'].astype(float).idxmax()]
    dic['average_vapor_2_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['平均水汽压(hPa)'].astype(float).idxmax()]
    dic['max_vapor_1']=basic_vapor_yearly['最大水汽压(hPa)'][basic_vapor_yearly['最大水汽压(hPa)'].astype(float).idxmin()]
    dic['max_vapor_1_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['最大水汽压(hPa)'].astype(float).idxmin()]
    dic['max_vapor_2']=basic_vapor_yearly['最大水汽压(hPa)'][basic_vapor_yearly['最大水汽压(hPa)'].astype(float).idxmax()]
    dic['max_vapor_2_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['最大水汽压(hPa)'].astype(float).idxmax()]
    dic['min_vapor_1']=basic_vapor_yearly['最小水汽压(hPa)'][basic_vapor_yearly['最小水汽压(hPa)'].astype(float).idxmin()]
    dic['min_vapor_1_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['最小水汽压(hPa)'].astype(float).idxmin()]
    dic['min_vapor_2']=basic_vapor_yearly['最小水汽压(hPa)'][basic_vapor_yearly['最小水汽压(hPa)'].astype(float).idxmax()]
    dic['min_vapor_2_year']=basic_vapor_yearly['年份'][basic_vapor_yearly['最小水汽压(hPa)'].astype(float).idxmax()]
    
    # 图像绘制
    # 平均水汽压
    average_vapor_picture_hournum=plot_picture(basic_vapor_yearly, '年份','平均水汽压(hPa)','平均水汽压(hPa)','hPa','历年平均水汽压变化.png',0.5,0.5,data_dir)

    # 平均最大水汽压
    max_vapor_picture_hournum=plot_picture(basic_vapor_yearly, '年份','最大水汽压(hPa)','平均最大水汽压(hPa)','hPa','历年平均最大水汽压变化.png',1,1,data_dir)

    # 平均最小水汽压
    min_vapor_picture_hournum=plot_picture(basic_vapor_yearly, '年份','最小水汽压(hPa)','平均最小水汽压(hPa)','hPa','历年最小平均水汽压变化.png',0.5,0.5,data_dir)

    dic['average_picture'] = InlineImage(doc, average_vapor_picture_hournum, width=Mm(130))
    dic['max_picture'] = InlineImage(doc, max_vapor_picture_hournum, width=Mm(130))
    dic['min_picture'] = InlineImage(doc,min_vapor_picture_hournum, width=Mm(130))
    
    # 逐月变化
    dic['average_vapor_m1']=basic_vapor_accum[basic_vapor_accum.iloc[0,1:-1:].astype(float).idxmin()][0]
    dic['average_vapor_m1_month']=basic_vapor_accum.iloc[0,1:-1:].astype(float).idxmin()
    dic['average_vapor_m2']=basic_vapor_accum[basic_vapor_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['average_vapor_m2_month']=basic_vapor_accum.iloc[0,1:-1:].astype(float).idxmax()
    
    dic['max_vapor_m1']=basic_vapor_accum[basic_vapor_accum.iloc[1,1:-1:].astype(float).idxmin()][1]
    dic['max_vapor_m1_month']=basic_vapor_accum.iloc[1,1:-1:].astype(float).idxmin()
    dic['max_vapor_m2']=basic_vapor_accum[basic_vapor_accum.iloc[1,1:-1:].astype(float).idxmax()][1]
    dic['max_vapor_m2_month']=basic_vapor_accum.iloc[1,1:-1:].astype(float).idxmax()
    
    dic['min_vapor_m1']=basic_vapor_accum[basic_vapor_accum.iloc[4,1:-1:].astype(float).idxmin()][4]
    dic['min_vapor_m1_month']=basic_vapor_accum.iloc[4,1:-1:].astype(float).idxmin()
    dic['min_vapor_m2']=basic_vapor_accum[basic_vapor_accum.iloc[4,1:-1:].astype(float).idxmax()][4]
    dic['min_vapor_m2_month']=basic_vapor_accum.iloc[4,1:-1:].astype(float).idxmax()
    
    # 图像绘制
    # 平均水汽压
    months=basic_vapor_accum.columns[1:-1:]
    average_vapor_picture_month=plot_picture_2(months,basic_vapor_accum.iloc[0, 1:-1],dic,'平均水汽压(hPa)','average_vapor_m1','average_vapor_m2','平均水汽压逐月变化.png',0.5,0.5,data_dir)

    # 平均最大水汽压
    max_vapor_picture_month=plot_picture_2(months,basic_vapor_accum.iloc[1, 1:-1],dic,'平均最大水汽压(hPa)','max_vapor_m1','max_vapor_m2','最大水汽压逐月变化.png',0.5,1,data_dir)

    # 平均最小水汽压
    min_vapor_picture_month=plot_picture_2(months,basic_vapor_accum.iloc[4, 1:-1],dic,'平均最小水汽压(hPa)','min_vapor_m1','min_vapor_m2','最小水汽压逐月变化.png',0.5,1,data_dir)

    dic['average_picture_m'] = InlineImage(doc, average_vapor_picture_month, width=Mm(130))
    dic['max_picture_m'] = InlineImage(doc, max_vapor_picture_month, width=Mm(130))
    dic['min_picture_m'] = InlineImage(doc,min_vapor_picture_month, width=Mm(130))
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'VAPOR.docx')
    doc.save(report)
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,basic_vapor_yearly,'站水汽压年变化')
    creat_table(document,basic_vapor_accum,'站水汽压月变化')

    document.save(report)


    return report