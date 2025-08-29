# -*- coding: utf-8 -*-
"""
Created on Mon May 27 09:46:35 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from Report.code.Module02.Function.rose import rose_picture as rose
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)
    

def win_report(basic_win_yearly,basic_win_accum,post_yearly_df,basic_win_d_accum,basic_win_s_accum,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module02','WIND.docx')
    doc=DocxTemplate(doc_path)
    
    
    basic_win_yearly=pd.DataFrame(basic_win_yearly)
    basic_win_accum=pd.DataFrame(basic_win_accum)
    basic_win_d_accum=pd.DataFrame(basic_win_d_accum)
    basic_win_s_accum=pd.DataFrame(basic_win_s_accum)
    
    dic=dict()
    
    ## 相对湿度
    # 年际变化
    dic['num_years']=len(basic_win_yearly)
    dic['station_name']=post_yearly_df['Station_Name'][0]
    dic['start_year']=basic_win_yearly['年份'][0]
    dic['end_year']=basic_win_yearly.iloc[-1,0]
    dic['average_win']=basic_win_accum['年'][0]

    
    dic['average_win_1']=basic_win_yearly['平均风速(m/s)'][basic_win_yearly['平均风速(m/s)'].astype(float).idxmin()]
    dic['average_win_1_year']=basic_win_yearly['年份'][basic_win_yearly['平均风速(m/s)'].astype(float).idxmin()]
    dic['average_win_2']=basic_win_yearly['平均风速(m/s)'][basic_win_yearly['平均风速(m/s)'].astype(float).idxmax()]
    dic['average_win_2_year']=basic_win_yearly['年份'][basic_win_yearly['平均风速(m/s)'].astype(float).idxmax()]

    
    # 图像绘制
    # 平均风速
    average_win_picture_hournum=plot_picture(basic_win_yearly, '年份','平均风速(m/s)','平均风速(m/s)','m/s','历年平均风速变化.png',0.3,0.3,data_dir)
    dic['average_picture'] = InlineImage(doc, average_win_picture_hournum, width=Mm(130))

    # 逐月变化
    dic['average_win_m1']=basic_win_accum[basic_win_accum.iloc[0,1:-1:].astype(float).idxmin()][0]
    dic['average_win_m1_month']=basic_win_accum.iloc[0,1:-1:].astype(float).idxmin()
    dic['average_win_m2']=basic_win_accum[basic_win_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['average_win_m2_month']=basic_win_accum.iloc[0,1:-1:].astype(float).idxmax()
    
    # 计算春夏秋冬
    spring = basic_win_accum.iloc[0,3:6] 
    summer = basic_win_accum.iloc[0,6:9]  
    autumn = basic_win_accum.iloc[0,9:12] 
    winter = pd.concat([basic_win_accum.iloc[0,1:3],basic_win_accum.iloc[0,12:13]])  
    
    # 计算每个季节的平均降水量
    average_spring = sum(spring) / len(spring)
    average_summer = sum(summer) / len(summer)
    average_autumn = sum(autumn) / len(autumn)
    average_winter = sum(winter) / len(winter)
    
    seasons = {
    "春季": average_spring,
    "夏季": average_summer,
    "秋季": average_autumn,
    "冬季": average_winter}
    
    sorted_seasons = sorted(seasons.items(), key=lambda x: x[1], reverse=True)

    dic['season1'] = sorted_seasons[0][0]  
    dic['season2'] = sorted_seasons[1][0]  
    dic['season3'] = sorted_seasons[2][0]  

    # 图像绘制
    # 平均降水量
    months=basic_win_accum.columns[1:-1:]
    average_win_picture_month=plot_picture_2(months,basic_win_accum.iloc[0, 1:-1],dic,'平均风速(m/s)','average_win_m1','average_win_m2','平均风速逐月变化.png',dic['average_win_m1'],0.5,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(months, basic_win_accum.iloc[0,1:-1:], width=0.4, color='skyblue')
    
    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('月')
    # plt.ylabel('平均风速(m/s)')
    # plt.xticks(months)
    # plt.ylim(0, dic['average_win_m2']+0.5)
    
    # average_win_picture_month=os.path.join(data_dir,'平均风速逐月变化.png')
    # plt.savefig(average_win_picture_month, bbox_inches='tight', dpi=200)
    # plt.clf()
    # plt.close('all')
    
    dic['average_picture_m'] = InlineImage(doc, average_win_picture_month, width=Mm(130))

    
    ## 风向频率
    dic['win_d_year_max_num'] = basic_win_d_accum[basic_win_d_accum.iloc[-1,1:-1:].astype(float).idxmax()][16]
    dic['win_d_year_max'] = basic_win_d_accum.iloc[-1,1:-1:].astype(float).idxmax()
    dic['win_d_year_min_num'] = basic_win_d_accum[basic_win_d_accum.iloc[-1,1:-1:].astype(float).idxmax()][16]
    dic['win_d_year_min'] = basic_win_d_accum.iloc[-1,1:-1:].astype(float).idxmax()
    
    dic['spring_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[12,1:-1:].astype(float).idxmax()][12]
    dic['spring_d'] =basic_win_d_accum.iloc[12,1:-1:].astype(float).idxmax()
    dic['summer_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[13,1:-1:].astype(float).idxmax()][13]
    dic['summer_d'] =   basic_win_d_accum.iloc[13,1:-1:].astype(float).idxmax() 
    dic['autumn_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[14,1:-1:].astype(float).idxmax()][14]
    dic['autumn_d'] = basic_win_d_accum.iloc[14,1:-1:].astype(float).idxmax()   
    dic['winter_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[15,1:-1:].astype(float).idxmax()][15]
    dic['winter_d'] =basic_win_d_accum.iloc[15,1:-1:].astype(float).idxmax()

    dic['d1_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['d1_d'] =basic_win_d_accum.iloc[0,1:-1:].astype(float).idxmax()
    dic['d2_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[1,1:-1:].astype(float).idxmax()][1]
    dic['d2_d'] =   basic_win_d_accum.iloc[1,1:-1:].astype(float).idxmax() 
    dic['d3_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[2,1:-1:].astype(float).idxmax()][2]
    dic['d3_d'] = basic_win_d_accum.iloc[2,1:-1:].astype(float).idxmax()   
    dic['d4_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[3,1:-1:].astype(float).idxmax()][3]
    dic['d4_d'] =basic_win_d_accum.iloc[3,1:-1:].astype(float).idxmax()
    dic['d5_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[4,1:-1:].astype(float).idxmax()][4]
    dic['d5_d'] =basic_win_d_accum.iloc[4,1:-1:].astype(float).idxmax()
    dic['d6_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[5,1:-1:].astype(float).idxmax()][5]
    dic['d6_d'] =basic_win_d_accum.iloc[5,1:-1:].astype(float).idxmax()
    dic['d7_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[6,1:-1:].astype(float).idxmax()][6]
    dic['d7_d'] =basic_win_d_accum.iloc[6,1:-1:].astype(float).idxmax()
    dic['d8_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[7,1:-1:].astype(float).idxmax()][7]
    dic['d8_d'] =basic_win_d_accum.iloc[7,1:-1:].astype(float).idxmax()
    dic['d9_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[8,1:-1:].astype(float).idxmax()][8]
    dic['d9_d'] =basic_win_d_accum.iloc[8,1:-1:].astype(float).idxmax()
    dic['d10_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[9,1:-1:].astype(float).idxmax()][9]
    dic['d10_d'] =basic_win_d_accum.iloc[9,1:-1:].astype(float).idxmax()
    dic['d11_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[10,1:-1:].astype(float).idxmax()][10]
    dic['d11_d'] =basic_win_d_accum.iloc[10,1:-1:].astype(float).idxmax()
    dic['d12_d_num'] =basic_win_d_accum[basic_win_d_accum.iloc[11,1:-1:].astype(float).idxmax()][11]
    dic['d12_d'] =basic_win_d_accum.iloc[11,1:-1:].astype(float).idxmax()
    
    
    fig_path=os.path.join(data_dir)
    title_spring=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'春季风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[12,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[12,1:17]),title_spring,'春季',fig_path)
    title_summer=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'夏季风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[13,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[13,1:17]),title_summer,'夏季',fig_path)
    title_autumn=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'秋季风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[14,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[14,1:17]),title_autumn,'秋季',fig_path)
    title_winter=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'冬季风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[15,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[15,1:17]),title_winter,'冬季',fig_path)
    title_year=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'全年风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[16,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[16,1:17]),title_year,'全年',fig_path)
    
    
    dic['sp']=os.path.join(fig_path,'春季.png')
    dic['spring_picture'] = InlineImage(doc, dic['sp'], width=Mm(70))
    dic['su']=os.path.join(fig_path,'夏季.png')
    dic["summer_picture"] = InlineImage(doc, dic['su'], width=Mm(70))
    dic['a']=os.path.join(fig_path,'秋季.png')
    dic["autumn_picture"] = InlineImage(doc, dic['a'], width=Mm(70))
    dic['w']=os.path.join(fig_path,'冬季.png')
    dic["winter_picture"] = InlineImage(doc, dic['w'], width=Mm(70))
    dic['q']=os.path.join(fig_path,'全年.png')
    dic["qn_picture"] = InlineImage(doc, dic['q'], width=Mm(70))
    
    title_1=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'1月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[0,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[0,1:17]),title_1,'1月',fig_path)
    title_2=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'2月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[1,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[1,1:17]),title_2,'2月',fig_path)
    title_3=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'3月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[2,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[2,1:17]),title_3,'3月',fig_path)
    title_4=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'4月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[3,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[3,1:17]),title_4,'4月',fig_path)
    title_5=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'5月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[4,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[4,1:17]),title_5,'5月',fig_path)
    title_6=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'6月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[5,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[5,1:17]),title_6,'6月',fig_path)
    title_7=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'7月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[6,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[6,1:17]),title_7,'7月',fig_path)
    title_8=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'8月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[7,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[7,1:17]),title_8,'8月',fig_path)
    title_9=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'9月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[8,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[8,1:17]),title_9,'9月',fig_path)
    title_10=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'10月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[9,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[9,1:17]),title_10,'10月',fig_path)
    title_11=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'11月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[10,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[10,1:17]),title_11,'11月',fig_path)
    title_12=dic['station_name']+'站累年'+'('+str(dic['start_year'])+'-'+str(dic['end_year'])+'年）'+'12月风向频率（%）'+'  '+'C='+' '+str(basic_win_d_accum.iloc[11,-1])+'%'
    rose(pd.DataFrame(basic_win_d_accum.iloc[11,1:17]),title_12,'12月',fig_path)

    
    dic['1']=os.path.join(fig_path,'1月.png')
    dic['d1_picture'] = InlineImage(doc, dic['1'], width=Mm(70))
    dic['2']=os.path.join(fig_path,'2月.png')
    dic["d2_picture"] = InlineImage(doc, dic['2'], width=Mm(70))
    dic['3']=os.path.join(fig_path,'3月.png')
    dic["d3_picture"] = InlineImage(doc, dic['3'], width=Mm(70))
    dic['4']=os.path.join(fig_path,'4月.png')
    dic["d4_picture"] = InlineImage(doc, dic['4'], width=Mm(70))
    dic['5']=os.path.join(fig_path,'5月.png')
    dic["d5_picture"] = InlineImage(doc, dic['5'], width=Mm(70))
    dic['6']=os.path.join(fig_path,'6月.png')
    dic["d6_picture"] = InlineImage(doc, dic['6'], width=Mm(70))
    dic['7']=os.path.join(fig_path,'7月.png')
    dic["d7_picture"] = InlineImage(doc, dic['7'], width=Mm(70))
    dic['8']=os.path.join(fig_path,'8月.png')
    dic["d8_picture"] = InlineImage(doc, dic['8'], width=Mm(70))
    dic['9']=os.path.join(fig_path,'9月.png')
    dic["d9_picture"] = InlineImage(doc, dic['9'], width=Mm(70))
    dic['10']=os.path.join(fig_path,'10月.png')
    dic["d10_picture"] = InlineImage(doc, dic['10'], width=Mm(70))
    dic['11']=os.path.join(fig_path,'11月.png')
    dic["d11_picture"] = InlineImage(doc, dic['11'], width=Mm(70))
    dic['12']=os.path.join(fig_path,'12月.png')
    dic["d12_picture"] = InlineImage(doc, dic['12'], width=Mm(70))
    
    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'WIND.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,basic_win_d_accum,'频率统计表（单位：%）')
    
    document.save(report)

    return report