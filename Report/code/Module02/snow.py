# -*- coding: utf-8 -*-
"""
Created on Thu May 23 15:42:15 2024

@author: EDY
"""
import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# post_yearly_df=snow_year
def snow_report(basic_snow_yearly,basic_snow_accum,post_yearly_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
      
    doc_path=os.path.join(cfg['report']['template'],'Module02','SNOW.docx')
    doc=DocxTemplate(doc_path)
    
    
    basic_snow_yearly=pd.DataFrame(basic_snow_yearly)
    basic_snow_accum=pd.DataFrame(basic_snow_accum)
    
    
    dic=dict()
    
    ## 相对湿度
    # 年际变化
    dic['num_years']=len(basic_snow_yearly)
    dic['station_name']=post_yearly_df['Station_Name'][0]
    dic['start_year']=basic_snow_yearly['年份'][0]
    dic['end_year']=basic_snow_yearly.iloc[-1,0]
    dic['average_snow']=basic_snow_accum['年'][0]

    
    dic['average_snow_1']=basic_snow_yearly['最大雪深(cm)'][basic_snow_yearly['最大雪深(cm)'].astype(float).idxmin()]
    dic['average_snow_1_year']=basic_snow_yearly['年份'][basic_snow_yearly['最大雪深(cm)'].astype(float).idxmin()]
    dic['average_snow_2']=basic_snow_yearly['最大雪深(cm)'][basic_snow_yearly['最大雪深(cm)'].astype(float).idxmax()]
    dic['average_snow_2_year']=basic_snow_yearly['年份'][basic_snow_yearly['最大雪深(cm)'].astype(float).idxmax()]

    
    # 图像绘制
    # 最大雪深
    mask = ~np.isnan(basic_snow_yearly['最大雪深(cm)'])
    valid_years = basic_snow_yearly['年份'][mask]
    valid_snowperatures = basic_snow_yearly['最大雪深(cm)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_snowperatures, 1)   
    if slope> 0:
        dic['average_snow_slope']='上升'
    else:
        dic['average_snow_slope']='下降'
    average_snow_picture_hournum=plot_picture(basic_snow_yearly, '年份','最大雪深(cm)','最大雪深（cm）','cm','历年最大雪深变化.png',2,2,data_dir)
    dic['average_picture'] = InlineImage(doc, average_snow_picture_hournum, width=Mm(130))

    # 逐月变化
    dic['average_snow_m1']=basic_snow_accum[basic_snow_accum.iloc[0,1:-1:].astype(float).idxmin()][0]
    dic['average_snow_m1_month']=basic_snow_accum.iloc[0,1:-1:].astype(float).idxmin()
    dic['average_snow_m2']=basic_snow_accum[basic_snow_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['average_snow_m2_month']=basic_snow_accum.iloc[0,1:-1:].astype(float).idxmax()
    
    dic['max_snow_m1']=basic_snow_accum[basic_snow_accum.iloc[1,1:-1:].astype(float).idxmin()][1]
    dic['max_snow_m1_month']=basic_snow_accum.iloc[1,1:-1:].astype(float).idxmin()
    dic['max_snow_m2']=basic_snow_accum[basic_snow_accum.iloc[1,1:-1:].astype(float).idxmax()][1]
    dic['max_snow_m2_month']=basic_snow_accum.iloc[1,1:-1:].astype(float).idxmax()
    
    # 图像绘制
    # 最大雪深
    months=basic_snow_accum.columns[1:-1:]
    average_snow_picture_month=plot_picture_2(months,basic_snow_accum.iloc[0, 1:-1],dic,'平均最大雪深(cm)','average_snow_m1','average_snow_m2','月平均最大雪深逐月变化.png',dic['average_snow_m1'],0.5,data_dir)
    dic['average_picture_m'] = InlineImage(doc, average_snow_picture_month, width=Mm(130))
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'SNOW.docx')
    doc.save(report)
    
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,basic_snow_yearly,'度年变化（单位：cm）')
    creat_table(document,basic_snow_accum,'度月变化（单位：cm）')

    document.save(report)
    

    return report