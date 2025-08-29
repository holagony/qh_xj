# -*- coding: utf-8 -*-
"""
Created on Mon May 27 15:39:57 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# post_yearly_df=gst_year
def gst_report(basic_gst_yearly,basic_gst_accum,post_yearly_df,data_dir):


    doc_path=os.path.join(cfg['report']['template'],'Module02','GST.docx')
    doc=DocxTemplate(doc_path)
    
    
    basic_gst_yearly=pd.DataFrame(basic_gst_yearly)
    basic_gst_accum=pd.DataFrame(basic_gst_accum)
    
    
    dic=dict()
    
    ## 地面温度
    # 年际变化
    dic['num_years']=len(basic_gst_yearly)
    dic['station_name']=post_yearly_df['Station_Name'][0]
    dic['start_year']=basic_gst_yearly['年份'][0]
    dic['end_year']=basic_gst_yearly.iloc[-1,0]
    dic['average_gst']=basic_gst_accum['年'][0]
    dic['max_gst']=basic_gst_accum['年'][1]
    dic['min_gst']=basic_gst_accum['年'][2]
    
    dic['average_gst_1']=basic_gst_yearly['平均地面温度(℃)'][basic_gst_yearly['平均地面温度(℃)'].astype(float).idxmin()]
    dic['average_gst_1_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均地面温度(℃)'].astype(float).idxmin()]
    dic['average_gst_2']=basic_gst_yearly['平均地面温度(℃)'][basic_gst_yearly['平均地面温度(℃)'].astype(float).idxmax()]
    dic['average_gst_2_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均地面温度(℃)'].astype(float).idxmax()]
    dic['max_gst_1']=basic_gst_yearly['平均最高地面温度(℃)'][basic_gst_yearly['平均最高地面温度(℃)'].astype(float).idxmin()]
    dic['max_gst_1_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均最高地面温度(℃)'].astype(float).idxmin()]
    dic['max_gst_2']=basic_gst_yearly['平均最高地面温度(℃)'][basic_gst_yearly['平均最高地面温度(℃)'].astype(float).idxmax()]
    dic['max_gst_2_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均最高地面温度(℃)'].astype(float).idxmax()]
    dic['min_gst_1']=basic_gst_yearly['平均最低地面温度(℃)'][basic_gst_yearly['平均最低地面温度(℃)'].astype(float).idxmin()]
    dic['min_gst_1_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均最低地面温度(℃)'].astype(float).idxmin()]
    dic['min_gst_2']=basic_gst_yearly['平均最低地面温度(℃)'][basic_gst_yearly['平均最低地面温度(℃)'].astype(float).idxmax()]
    dic['min_gst_2_year']=basic_gst_yearly['年份'][basic_gst_yearly['平均最低地面温度(℃)'].astype(float).idxmax()]
    
    
    
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    # 图像绘制
    # years=basic_gst_yearly['年份'][:]
    
    # 平均地面温度
    mask = ~np.isnan(basic_gst_yearly['平均地面温度(℃)'])
    valid_years = basic_gst_yearly['年份'][mask]
    valid_gstperatures = basic_gst_yearly['平均地面温度(℃)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_gstperatures, 1)   
    if slope> 0:
        dic['average_gst_slope']='上升'
    else:
        dic['average_gst_slope']='下降'
    
    average_gst_picture_hournum=plot_picture(basic_gst_yearly, '年份','平均地面温度(℃)','平均地面温度(℃)','℃','历年平均地面温度变化.png',0.5,0.5,data_dir)
    
    # 平均最高地面温度
    mask = ~np.isnan(basic_gst_yearly['平均最高地面温度(℃)'])
    valid_years = basic_gst_yearly['年份'][mask]
    valid_gstperatures = basic_gst_yearly['平均最高地面温度(℃)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_gstperatures, 1)    
    if slope> 0:
        dic['max_gst_slope']='上升'
    else:
        dic['max_gst_slope']='下降'
    max_gst_picture_hournum=plot_picture(basic_gst_yearly, '年份','平均最高地面温度(℃)','地面温度(℃)','℃','历年平均最高地面温度变化.png',0.5,0.5,data_dir)

    # 平均最低地面温度
    mask = ~np.isnan(basic_gst_yearly['平均最低地面温度(℃)'])
    valid_years = basic_gst_yearly['年份'][mask]
    valid_gstperatures = basic_gst_yearly['平均最低地面温度(℃)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_gstperatures, 1)
    if slope> 0:
        dic['min_gst_slope']='上升'
    else:
        dic['min_gst_slope']='下降'
    min_gst_picture_hournum=plot_picture(basic_gst_yearly, '年份','平均最低地面温度(℃)','地面温度(℃)','℃','历年最低平均地面温度变化.png',0.5,0.5,data_dir)

    
    dic['average_picture'] = InlineImage(doc, average_gst_picture_hournum, width=Mm(130))
    dic['max_picture'] = InlineImage(doc, max_gst_picture_hournum, width=Mm(130))
    dic['min_picture'] = InlineImage(doc,min_gst_picture_hournum, width=Mm(130))
    
    # 逐月变化
    dic['average_gst_m1']=basic_gst_accum[basic_gst_accum.iloc[0,1:-1:].astype(float).idxmin()][0]
    dic['average_gst_m1_month']=basic_gst_accum.iloc[0,1:-1:].astype(float).idxmin()
    dic['average_gst_m2']=basic_gst_accum[basic_gst_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['average_gst_m2_month']=basic_gst_accum.iloc[0,1:-1:].astype(float).idxmax()
    
    dic['max_gst_m1']=basic_gst_accum[basic_gst_accum.iloc[1,1:-1:].astype(float).idxmin()][1]
    dic['max_gst_m1_month']=basic_gst_accum.iloc[1,1:-1:].astype(float).idxmin()
    dic['max_gst_m2']=basic_gst_accum[basic_gst_accum.iloc[1,1:-1:].astype(float).idxmax()][1]
    dic['max_gst_m2_month']=basic_gst_accum.iloc[1,1:-1:].astype(float).idxmax()
    
    dic['min_gst_m1']=basic_gst_accum[basic_gst_accum.iloc[2,1:-1:].astype(float).idxmin()][2]
    dic['min_gst_m1_month']=basic_gst_accum.iloc[2,1:-1:].astype(float).idxmin()
    dic['min_gst_m2']=basic_gst_accum[basic_gst_accum.iloc[2,1:-1:].astype(float).idxmax()][2]
    dic['min_gst_m2_month']=basic_gst_accum.iloc[2,1:-1:].astype(float).idxmax()
    
    # 图像绘制
    # 平均地面温度
    months=basic_gst_accum.columns[1:-1:]
    average_gst_picture_month=plot_picture_2(months,basic_gst_accum.iloc[0, 1:-1],dic,'平均地面温度(℃)','average_gst_m1','average_gst_m2','平均地面温度逐月变化.png',2,2,data_dir)
    
    # 平均最高地面温度
    max_gst_picture_month=plot_picture_2(months,basic_gst_accum.iloc[1, 1:-1],dic,'平均最高地面温度(℃)','max_gst_m1','max_gst_m2','最高地面温度逐月变化.png',2,2,data_dir)

    # 平均最低地面温度
    min_gst_picture_month=plot_picture_2(months,basic_gst_accum.iloc[2, 1:-1],dic,'平均最低地面温度(℃)','min_gst_m1','min_gst_m2','最低地面温度逐月变化.png',2,2,data_dir)

    dic['average_picture_m'] = InlineImage(doc, average_gst_picture_month, width=Mm(130))
    dic['max_picture_m'] = InlineImage(doc, max_gst_picture_month, width=Mm(130))
    dic['min_picture_m'] = InlineImage(doc,min_gst_picture_month, width=Mm(130))
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'GST.docx')
    doc.save(report)
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,basic_gst_yearly,'年地面温度年变化')
    creat_table(document,basic_gst_accum,'年地面温度月变化')

    document.save(report)

    return report