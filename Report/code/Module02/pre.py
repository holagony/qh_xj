# -*- coding: utf-8 -*-
"""
Created on Thu May 23 15:42:15 2024

@author: EDY
"""
import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)



def pre_report(basic_pre_yearly,basic_pre_accum,post_yearly_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module02','PRE.docx')
    doc=DocxTemplate(doc_path)
    
    
    basic_pre_yearly=pd.DataFrame(basic_pre_yearly)
    basic_pre_accum=pd.DataFrame(basic_pre_accum)
    
    
    dic=dict()
    
    ## 相对湿度
    # 年际变化
    dic['num_years']=len(basic_pre_yearly)
    dic['station_name']=post_yearly_df['Station_Name'][0]
    dic['start_year']=basic_pre_yearly['年份'][0]
    dic['end_year']=basic_pre_yearly.iloc[-1,0]
    dic['average_pre']=basic_pre_accum['年'][0]

    
    dic['average_pre_1']=basic_pre_yearly['总降水量(mm)'][basic_pre_yearly['总降水量(mm)'].astype(float).idxmin()]
    dic['average_pre_1_year']=basic_pre_yearly['年份'][basic_pre_yearly['总降水量(mm)'].astype(float).idxmin()]
    dic['average_pre_2']=basic_pre_yearly['总降水量(mm)'][basic_pre_yearly['总降水量(mm)'].astype(float).idxmax()]
    dic['average_pre_2_year']=basic_pre_yearly['年份'][basic_pre_yearly['总降水量(mm)'].astype(float).idxmax()]

    
    # 图像绘制
    # 平均降水量
    mask = ~np.isnan(basic_pre_yearly['总降水量(mm)'])
    valid_years = basic_pre_yearly['年份'][mask]
    valid_preperatures = basic_pre_yearly['总降水量(mm)'][mask]
    slope, intercept = np.polyfit(valid_years, valid_preperatures, 1)   
    if slope> 0:
        dic['average_pre_slope']='上升'
    else:
        dic['average_pre_slope']='下降'
        
    average_pre_picture_hournum=plot_picture(basic_pre_yearly, '年份','总降水量(mm)','降水量(mm)','mm','历年平均降水量变化.png',40,40,data_dir)
    
    # 逐月变化
    dic['average_pre_m1']=basic_pre_accum[basic_pre_accum.iloc[0,1:-1:].astype(float).idxmin()][0]
    dic['average_pre_m1_month']=basic_pre_accum.iloc[0,1:-1:].astype(float).idxmin()
    dic['average_pre_m2']=basic_pre_accum[basic_pre_accum.iloc[0,1:-1:].astype(float).idxmax()][0]
    dic['average_pre_m2_month']=basic_pre_accum.iloc[0,1:-1:].astype(float).idxmax()
    
    # 计算春夏秋冬
    spring = basic_pre_accum.iloc[0,3:6] 
    summer = basic_pre_accum.iloc[0,6:9]  
    autumn = basic_pre_accum.iloc[0,9:12] 
    winter = pd.concat([basic_pre_accum.iloc[0,1:3],basic_pre_accum.iloc[0,12:13]])  
    
    # 计算每个季节的平均降水量
    average_spring = sum(spring) / len(spring)
    average_summer = sum(summer) / len(summer)
    average_autumn = sum(autumn) / len(autumn)
    average_winter = sum(winter) / len(winter)
    
    seasons = {
    "春季": average_spring,
    "夏季": average_summer,
    "秋季": average_autumn,
    "冬季": average_winter}
    
    sorted_seasons = sorted(seasons.items(), key=lambda x: x[1], reverse=True)

    dic['season1'] = sorted_seasons[0][0]  
    dic['season2'] = sorted_seasons[1][0]  
    dic['season3'] = sorted_seasons[2][0]  

    # 图像绘制
    # 平均降水量
    months=basic_pre_accum.columns[1:-1:]

    average_pre_picture_month=plot_picture_2(months,basic_pre_accum.iloc[0, 1:-1],dic,'平均降水量(℃)','average_pre_m1','average_pre_m2','平均降水量逐月变化.png',2,10,data_dir)

    
    dic['average_picture'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
    dic['average_picture_m'] = InlineImage(doc, average_pre_picture_month, width=Mm(130))

    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'PRE.docx')
    doc.save(report)
    
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,basic_pre_yearly,'年变化（单位：mm）')
    creat_table(document,basic_pre_accum,'月变化（单位：mm）')

    document.save(report)

    return report