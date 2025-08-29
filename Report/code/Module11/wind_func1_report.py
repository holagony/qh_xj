# -*- coding: utf-8 -*-
"""
Created on Thu Jul 11 09:45:34 2024

@author: EDY
"""


import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
from datetime import datetime
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module11'
# time_range='20220801,20240830'
def wind_func1_report(check_result1,check_result2,time_range,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module11','wind_1_report.docx')
    doc=DocxTemplate(doc_path)
    
    
    # 处理一下数据
    check_result=dict(check_result1)
    ele_name=dict()
    ele_name['风速']='风速'
    ele_name['风向']='风向'
    ele_name['最大风速']='最大风速'
    ele_name['极大风速']='极大风速'
    
    df=[]
    for key_1 in check_result.keys():
        for key_2 in check_result[key_1].keys():
            for key_3 in check_result[key_1][key_2].keys():
                data_3=pd.DataFrame(check_result[key_1][key_2][key_3])
                data_3.insert(0, '高度',key_3)
                data_3.insert(0, '要素',ele_name[key_2])
                data_3.insert(0, '测风塔编号',key_1)
                df.append(data_3)
    df= pd.concat(df)

    # 总时长
    start_date_str, end_date_str = time_range.split(',')
    start_date = datetime.strptime(start_date_str, '%Y%m%d')
    end_date = datetime.strptime(end_date_str, '%Y%m%d')  
    hours_difference = (end_date - start_date).days*24
    
    # 报告填充
    dic=dict()
    dic['name']=key_1
    dic['start_year']=start_date_str[0:4:]
    dic['start_month']=start_date_str[4:6:]
    dic['end_year']=end_date_str[0:4:]
    dic['end_month']=end_date_str[4:6:]
    dic['num_1']=hours_difference
    
    dic['qc']=''
    for key_2 in check_result[key_1].keys():
        num=1
        for key_3 in check_result[key_1][key_2].keys():
                data = df[(df['要素'] == ele_name[key_2]) & (df['高度'] == key_3)]
                if num==1:
                    str_1=ele_name[key_2]+','+key_3+'高度，共缺测数据'+str(data['缺测时间总计(小时)'].sum())+',占比'+str(round(data['缺测时间总计(小时)'].sum()/hours_difference*100,2))+'%。'
                else:
                    str_1=key_3+'高度，共缺测数据'+str(data['缺测时间总计(小时)'].sum())+',占比'+str(round(data['缺测时间总计(小时)'].sum()/hours_difference*100,2))+'%。'
                                    
                dic['qc']=dic['qc']+str_1
                
                num=num+1
                
    check_result2=dict(check_result2)
    df_2=[]
    for key_1 in check_result2.keys():
        for key_2 in check_result2[key_1].keys():
            for key_3 in check_result2[key_1][key_2].keys():
                data_3=pd.DataFrame(check_result2[key_1][key_2][key_3])
                data_3.insert(0, '高度',key_3)
                data_3.insert(0, '要素',ele_name[key_2])
                data_3.insert(0, '测风塔编号',key_1)
                df_2.append(data_3)
    df_2= pd.concat(df_2)

    
    dic['qc_2']=''
    for key_2 in check_result2[key_1].keys():
        num=1
        for key_3 in check_result2[key_1][key_2].keys():
                data = df_2[(df_2['要素'] == ele_name[key_2]) & (df_2['高度'] == key_3)]
                if num==1:
                    str_1=ele_name[key_2]+','+key_3+'高度，有效完整率'+str(round(data['实有样本数'].sum()/data['应有样本数'].sum(),2))+'%。'
                else:
                    str_1=key_3+'高度，有效完整率'+str(round(data['实有样本数'].sum()/data['应有样本数'].sum(),2))+'%。'

                dic['qc_2']=dic['qc_2']+str_1
                num=num+1
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'wind_1_report.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    
    creat_table(document,df,'测风塔评估年缺测时间')
    creat_table(document,df_2,'测风塔各月测风数据及缺测和无效次数统计表')


    document.save(report)
    
    
    return report
