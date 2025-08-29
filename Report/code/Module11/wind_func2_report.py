# -*- coding: utf-8 -*-
"""
Created on Thu Jul 11 13:29:03 2024

@author: EDY
"""
import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module11'
# time_range='20220801,20240830'
def wind_func2_report(result,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module11','wind_2_report.docx')
    doc=DocxTemplate(doc_path)
    
    dic=dict()
    #%% 空气密度
    p=pd.DataFrame(result['气象站空气密度'])
    

    data_to_plot = p.iloc[:, 1:13]  
    legends = p.iloc[:, 0]          
    
    fig, ax = plt.subplots()
    
    x_labels = data_to_plot.columns
    for i in range(len(data_to_plot)):
        ax.plot(x_labels, data_to_plot.iloc[i], label=legends[i])
    ax.legend()   
    ax.set_xticks(range(len(x_labels)))
    ax.set_xticklabels(x_labels)
    ax.set_xlabel('月份')
    ax.set_ylabel('空气密度$(kg/m^3)$')
    ax.grid()
    
    average_pre_picture_hournum=os.path.join(data_dir,'测风塔评估年各月平均空气密度.png')
    plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
    plt.clf()
    plt.close('all')
    dic['figure_1'] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
    plt.close()

    dic['p_80'] =p.iloc[7,13]
    dic['p_90'] =p.iloc[8,13]
    dic['p_100'] =p.iloc[9,13]

    #%% 平均风速和平均风功率密度
    key_1=next(iter(result))
    
    data_1=pd.DataFrame(result[key_1]['综合参数统计'])
    dic['data_1_2'] ='m/s、'.join(map(str, data_1['平均风速(m/s)']))
    dic['data_1_3'] ='(W/m²)、'.join(map(str, data_1['平均风功率密度(W/m²)']))
    dic['data_1_4'] ='h、'.join(map(str, data_1['有效风速小时数']))
    dic['data_1_5'] ='%、'.join(map(str, data_1['有效风速小时数百分率%']))
    dic['data_1_6'] ='(W/m²)、'.join(map(str, data_1['有效风功率密度(W/m²)']))
    dic['data_1_7'] ='m/s、'.join(map(str, data_1['最大风速(m/s)']))
    dic['data_1_8'] ='m/s、'.join(map(str, data_1['极大风速(m/s)']))

    
    #%% 月变化和年变化
    data_21=pd.DataFrame(result[key_1]['逐月平均风速'])
    data_22=pd.DataFrame(result[key_1]['逐月平均风功率密度'])
    data_23=pd.DataFrame(result[key_1]['逐月平均有效风功率密度'])
    data_24=pd.DataFrame(result[key_1]['平均风速日变化'])
    data_25=pd.DataFrame(result[key_1]['平均风功率密度日变化'])
    data_26=pd.DataFrame(result[key_1]['平均有效风功率密度日变化'])

    data_211 = data_21.copy()
    data_211 = data_211.sort_values(by=data_21.columns[-1], ascending=False)
    data_211.reset_index(inplace=True, drop=True)

    
    def data_2_f1(x):
        return int(x[6::])
    
    dic['data_2_1'] ='、'.join(map(str,sorted(map(data_2_f1,data_211['时间'][:3:]))))
    dic['data_2_2'] =data_21.columns[-1]
    dic['data_2_3'] =data_211[data_21.columns[-1]].min()
    dic['data_2_4'] =data_211[data_21.columns[-1]].max()
    dic['data_2_5'] =int(data_211['时间'][0][5::])
    dic['data_2_6'] =data_211[data_21.columns[-1]].max()
    dic['data_2_7'] =int(data_211['时间'][len(data_211)-1][5::])
    dic['data_2_8'] =data_211[data_21.columns[-1]].min()  
    
    data_221 = data_22.copy()
    data_221 = data_221.sort_values(by=data_22.columns[-1], ascending=False)
    dic['data_2_9'] ='、'.join(map(str,sorted(map(data_2_f1,data_221['时间'][:3:]))))
    dic['data_2_10'] =data_221.columns[-1]
    dic['data_2_11'] =data_221[data_221.columns[-1]].min()
    dic['data_2_12'] =data_221[data_221.columns[-1]].max()
    
    data_231 = data_23.copy()
    data_231 = data_231.sort_values(by=data_23.columns[-1], ascending=False)
    dic['data_2_13'] ='、'.join(map(str,sorted(map(data_2_f1,data_221['时间'][:3:]))))
    dic['data_2_14'] =data_231.columns[-1]
    dic['data_2_15'] =data_231[data_231.columns[-1]].min()
    dic['data_2_16'] =data_231[data_231.columns[-1]].max()
    
    #%% 图片绘制
    str_2=''
    for name in data_21.columns[1::]:
        str_2=str_2+'{{picture_2_'+name+'}}'
    dic['figure_2'] =str_2
    
    str_3=''
    for name in data_21.columns[1::]:
        str_3=str_3+'{{picture_3_'+name+'}}'
    dic['figure_3'] =str_3
    
    doc.render(dic)
    report=os.path.join(data_dir,'wind_2_1_report.docx')
    doc.save(report)

    doc_path=os.path.join(data_dir,'wind_2_1_report.docx')
    doc=DocxTemplate(doc_path)
    
    dic=dict()
    
    for name in data_21.columns[1::]:
        
        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(111)    
        x_labels = data_21['时间']
        
        color = 'tab:red'
        ax.set_xlabel('时间')
        ax.set_ylabel('风速（m/s）')
        line1,=ax.plot(x_labels, data_21[name], color=color,label='平均风速')
        ax.tick_params(axis='y', labelcolor=color)
        
        ax2 = ax.twinx()  
        color = 'tab:blue'
        ax2.set_ylabel('风功率密度$(W/m^2)$')
        line2,=ax2.plot(x_labels, data_22[name], color=color, label='平均风功率密度')
        line3,=ax2.plot(x_labels, data_23[name], color='k', label='平均风功率密度')

        ax2.tick_params(axis='y', labelcolor=color)
        fig.tight_layout()  
        
        lines = [line1, line2, line3]
        labels = [l.get_label() for l in lines]
        fig.legend(lines, labels, loc='upper center', bbox_to_anchor=(0.5, 0.98), ncol=3)

        ax.set_xticks(range(len(x_labels)))
        ax.set_xticklabels(x_labels)
        ax.grid()
        plt.title(name+'高度')
        average_pre_picture_hournum=os.path.join(data_dir,name+'米月平均风.png')
        plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        plt.clf()
        plt.close('all')
        dic['picture_2_'+name] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
        plt.close()
        
    for name in data_24['高度']:
        
        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(111)    
        x_labels = data_24.columns[1::]
        
        color = 'tab:red'
        ax.set_xlabel('时间')
        ax.set_ylabel('风速（m/s）')
        line1,=ax.plot(x_labels, data_24[data_24['高度']==name].iloc[0,1::], color=color,label='平均风速')
        ax.tick_params(axis='y', labelcolor=color)
        
        ax2 = ax.twinx()  
        color = 'tab:blue'
        ax2.set_ylabel('风功率密度$(W/m^2)$')
        line2,=ax2.plot(x_labels, data_25[data_25['高度']==name].iloc[0,1::], color=color, label='平均风功率密度')
        line3,=ax2.plot(x_labels, data_26[data_26['高度']==name].iloc[0,1::], color='k', label='平均风功率密度')

        ax2.tick_params(axis='y', labelcolor=color)
        fig.tight_layout()  
        
        lines = [line1, line2, line3]
        labels = [l.get_label() for l in lines]
        fig.legend(lines, labels, loc='upper center', bbox_to_anchor=(0.5, 0.98), ncol=3)

        ax.set_xticks(range(len(x_labels)))
        ax.set_xticklabels(x_labels)
        ax.grid()
        plt.title(name+'高度')

        average_pre_picture_hournum=os.path.join(data_dir,name+'米日平均风.png')
        plt.savefig(average_pre_picture_hournum, bbox_inches='tight', dpi=200)
        plt.clf()
        plt.close('all')
        dic['picture_3_'+name] = InlineImage(doc, average_pre_picture_hournum, width=Mm(130))
        plt.close()
    
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'wind_2_report.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,p,'测风塔评估年各月平均空气密度')
    creat_table(document,data_1,'测风塔评估年风能参数表')
    creat_table(document,data_21,'测风塔评估年实测序列平均风速月变化值')
    creat_table(document,data_22,'测风塔评估年实测序列平均风功率密度月变化值')
    creat_table(document,data_23,'测风塔评估年实测序列平均有效风功率密度月变化值')
    creat_table(document,data_24,'测风塔评估年实测序列平均风速日变化值')
    creat_table(document,data_25,'测风塔评估年实测序列平均风功率密度日变化值')
    creat_table(document,data_26,'测风塔评估年实测序列平均有效风功率密度日变化值')

    document.save(report)
    
    
    return report
