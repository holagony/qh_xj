# -*- coding: utf-8 -*-
"""
Created on Thu Jul 11 17:53:35 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module11'
def wind_func3_report(result4,result5,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module11','wind_3_report.docx')
    doc=DocxTemplate(doc_path)
    
    key_1=next(iter(result4))
    dic=dict()


    #%% 
    data_11=pd.DataFrame(result4[key_1]['风向频率'])
    data_12=pd.DataFrame(result4[key_1]['风能密度方向频率'])
    data_13=pd.DataFrame(result4[key_1]['逐月有效风速小时数'])
    data_14=pd.DataFrame(result4[key_1]['各风速等级小时数'])
    data_15=pd.DataFrame(result4[key_1]['风速频率分布'])
    data_16=pd.DataFrame(result4[key_1]['各风速区间风能频率分布'])

    data_11['高度']=data_12['高度']
    nan_counts=data_11.iloc[:,1:17].isna().sum(axis=1)
    data_11=data_11[~(nan_counts>=15)]
    data_12=data_12[~(nan_counts>=15)]

    dic['data_1']='、'.join(data_11['高度'])
    dic['data_2']='、'.join(data_11.iloc[:,1:17:].idxmax(axis=1))
    dic['data_3']='%、'.join(map(str,data_11.iloc[:,1:17:].max(axis=1)))
    dic['data_4']='、'.join(data_12.iloc[:,1:17:].idxmax(axis=1))
    dic['data_5']='小时、'.join(map(str,data_13.iloc[:,1::].sum(axis=0)))
    dic['data_6']=int(data_13['时间'][data_13.iloc[:,-1].idxmax()][5::])
    dic['data_7']=int(data_13['时间'][data_13.iloc[:,-1].idxmin()][5::])
    dic['data_8']=data_14.iloc[-1,0]
    dic['data_9']='、'.join(data_14.columns[1::])
    dic['data_10']='小时、'.join(map(str,data_14.iloc[-1,1:-1:]))
    dic['data_11']=data_14.iloc[-1,-1]
    dic['data_12']=data_15.columns[np.where(data_15.iloc[-1,::]==data_15.iloc[-1,1:].max())][0]
    dic['data_13']=data_15.iloc[-1,1:].max()
    dic['data_14']=data_16.columns[np.where(data_16.iloc[-1,::]==data_16.iloc[-1,1:].max())][0]
    dic['data_15']=data_16.iloc[-1,1:].max()

    str_2=''
    for name in data_16['高度']:
        str_2=str_2+'{{picture_'+name+'}}'
    dic['figure'] =str_2
    
    
    doc.render(dic)
    report=os.path.join(data_dir,'wind_3_1_report.docx')
    doc.save(report)

    doc_path=os.path.join(data_dir,'wind_3_1_report.docx')
    doc=DocxTemplate(doc_path)
    
    dic=dict()

    for name in data_16['高度']:
        try:
            dic['picture_'+name] = InlineImage(doc, result5[key_1]['img_save_path'][name], width=Mm(130))
        except:
            dic['picture_'+name] =None
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'wind_3_report.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    creat_table(document,data_11,'测风塔评估年各高度风向频率')
    creat_table(document,data_12,'测风塔评估年各高度风能密度方向频率')
    creat_table(document,data_13,'测风塔有效风速小时数')
    creat_table(document,data_14,'测风塔各高度各风速等级小时数（h）')
    creat_table(document,data_15,'测风塔评估年各高度风速频率分布表')
    creat_table(document,data_16,'测风塔评估年各高度风能频率分布表')
    creat_table(document,pd.DataFrame(result5[key_1]['weibull_params']),'测风塔评估年各高度风速Weibull分布参数')

    document.save(report)
    
    
    return report
