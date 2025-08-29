# -*- coding: utf-8 -*-
"""
Created on Thu Jun 13 09:44:04 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'

def water_circulation_report(result,df_day,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','water_circulation.docx')
    doc=DocxTemplate(doc_path)
    

    dic=dict()
    dic['station_name']=df_day['Station_Name'][0]
    dic['start_year']=df_day.index.year[0]
    dic['end_year']=df_day.index.year[-1]
    dic['num_years']=len(df_day.index.year.unique())

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'water_circulation.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据

    result=dict(result)
    result1=pd.DataFrame(result['result1'])
    result2=pd.DataFrame(result['result2'])
    result3=pd.DataFrame(result['result3'])
    result4=pd.DataFrame(result['result4'])
    result5=pd.DataFrame(result['result5'])
    result6=pd.DataFrame(result['result6'])
    result7=pd.DataFrame(result['result7'])
    
    result1.columns=['干球温度（°C）','气压（hPa）','相对湿度','湿球温度（°C）','10%累积频率湿球温度','日期']
    result2.columns=['干球温度（°C）','气压（hPa）','相对湿度','湿球温度（°C）','5%累积频率湿球温度','日期']
    result3.columns=['干球温度（°C）','气压（hPa）','相对湿度','湿球温度（°C）','1%累积频率湿球温度','日期']
    result4.columns=['湿球温度（°C）','干球温度（°C）','气压（hPa）','相对湿度']


    column_names = result5.columns.tolist()
    insert_data = ['日期', '湿球温度（°C）', '干球温度（°C）','气压（hPa）','相对湿度']
    new_result5 = pd.DataFrame(columns=[str(i) for i in np.arange(np.size(result5,1))])
    new_result5['0']=insert_data
    for i in np.arange(np.size(result5,1)-1):
        new_result5.at[0,str(i+1)]=column_names[i]
        new_result5.loc[1:4,str(i+1)]=np.array(result5.iloc[:,i+1])
    new_result5.columns = new_result5.iloc[0]
    new_result5 = new_result5.drop(new_result5.index[0])

    column_names = result6.columns.tolist()
    insert_data = ['日期', '湿球温度（°C）', '干球温度（°C）','气压（hPa）','相对湿度']
    new_result6 = pd.DataFrame(columns=[str(i) for i in np.arange(np.size(result6,1))])
    new_result6['0']=insert_data
    for i in np.arange(np.size(result5,1)-1):
        new_result6.at[0,str(i+1)]=column_names[i]
        new_result6.loc[1:4,str(i+1)]=np.array(result6.iloc[:,i+1])
    new_result6.columns = new_result6.iloc[0]
    new_result6 = new_result6.drop(new_result6.index[0])    
        
    column_names = result7.columns.tolist()
    insert_data = ['日期', '湿球温度（°C）', '干球温度（°C）','气压（hPa）','相对湿度']
    new_result7 = pd.DataFrame(columns=[str(i) for i in np.arange(np.size(result7,1))])
    new_result7['0']=insert_data
    for i in np.arange(np.size(result7,1)-1):
        new_result7.at[0,str(i+1)]=column_names[i]
        new_result7.loc[1:4,str(i+1)]=np.array(result7.iloc[:,i+1])
    new_result7.columns = new_result7.iloc[0]
    new_result7 = new_result7.drop(new_result7.index[0])
    
    creat_table(document,result3,'夏季1%累积频率的日湿球温度及对应的干球温度、气压、相对湿度')
    creat_table(document,result2,'夏季5%累积频率的日湿球温度及对应的干球温度、气压、相对湿度')
    creat_table(document,result1,'夏季10%累积频率的日湿球温度及对应的干球温度、气压、相对湿度')
    creat_table(document,result4,'最高第七个日平均湿球温度的历年平均值及对应的干球温度、气压、相对湿度')
    creat_table(document,new_result5,'历年最高湿球温度及对应的干球温度、气压、相对湿度')
    creat_table(document,new_result6,'历年最低湿球温度及对应的干球温度、气压、相对湿度')
    creat_table(document,new_result7,'历年平均湿球温度及对应的干球温度、气压、相对湿度')


    document.save(report)
    
    
    return report
