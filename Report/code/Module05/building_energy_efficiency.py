# -*- coding: utf-8 -*-
"""
Created on Thu Jun 13 13:49:14 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'
def building_energy_efficiency(result, post_daily_df, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','building_energy_efficiency.docx')
    doc=DocxTemplate(doc_path)
    
    
    dic=dict()
    dic['station_name']=post_daily_df.iloc[0,1]
    dic['start_year']=post_daily_df.index.year[0]
    dic['end_year']=post_daily_df.index.year[-1]
    
    dic['HDD18']=result['HDD18']
    dic['CDD26']=result['CDD26']
    dic['Z']=result['Z']
    dic['Z_start']=result['Z_start']
    dic['Z_end']=result['Z_end']
    
    table1=pd.DataFrame(result['table1'])
    table2=pd.DataFrame(result['table2'])
    table3=pd.DataFrame(result['table3'])
    table4=pd.DataFrame(result['table4'])
    
    plt.figure(figsize=(10, 6))
    plt.bar(table1['date'], table1['历年采暖日数'], width=0.4, color='skyblue')
    
    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('时间')
    plt.ylabel('历年采暖日数')
    plt.xticks(table1['date'][::3], rotation=45) 
    save_path = data_dir + '/历年采暖日数.png'
    plt.savefig(save_path, dpi=200, bbox_inches='tight')
    plt.clf()
    plt.close('all')
    dic['历年采暖日数'] = InlineImage(doc, save_path, width=Mm(130))
    
    plt.figure(figsize=(10, 6))
    plt.bar(table2['date'], table2['累年各月平均采暖日数'], width=0.4, color='skyblue')
    
    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('时间')
    plt.ylabel('累年各月平均采暖日数')
    plt.xticks(table2['date'][::], rotation=45) 
    save_path = data_dir + '/累年各月平均采暖日数.png'
    plt.savefig(save_path, dpi=200, bbox_inches='tight')
    plt.clf()
    plt.close('all')
    dic['累年各月平均采暖日数'] = InlineImage(doc, save_path, width=Mm(130))
    
    plt.figure(figsize=(10, 6))
    plt.bar(table3['date'], table3['历年空调日数'], width=0.4, color='skyblue')
    
    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('时间')
    plt.ylabel('历年空调日数')
    plt.xticks(table3['date'][::3], rotation=45) 
    save_path = data_dir + '/历年空调日数.png'
    plt.savefig(save_path, dpi=200, bbox_inches='tight')
    plt.clf()
    plt.close('all')
    dic['历年空调日数'] = InlineImage(doc, save_path, width=Mm(130))
    
    plt.figure(figsize=(10, 6))
    plt.bar(table4['date'], table4['累年各月平均空调日数'], width=0.4, color='skyblue')
    
    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('时间')
    plt.ylabel('历年空调日数')
    plt.xticks(table4['date'][::], rotation=45) 
    save_path = data_dir + '/累年各月平均空调日数.png'
    plt.savefig(save_path, dpi=200, bbox_inches='tight')
    plt.clf()
    plt.close('all')
    dic['累年各月平均空调日数'] = InlineImage(doc, save_path, width=Mm(130))

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'building_energy_efficiency.docx')
    doc.save(report)
    
    ## 插入表格
    # document = Document(report)
    

    # document.save(report)
    
    
    return report
