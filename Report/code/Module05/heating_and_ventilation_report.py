# -*- coding: utf-8 -*-
"""
Created on Fri May 31 14:35:06 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# df=df_dict
# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'
# daily_df=post_daily_df
# monthly_df=post_monthly_df
def heating_and_ventilation_report(df,t_sh,enthalpy,daily_df,monthly_df,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','heating_and_ventilation.docx')
    doc=DocxTemplate(doc_path)
    
    # result=pd.DataFrame(columns=['参数','数值'])
    # result['参数']=df['name']  
    # result['数值']=df['value']  
    
    dic=dict()
    dic['start_year']=daily_df.index.year[0]
    dic['start_month']=daily_df.index.month[0]
    dic['start_day']=daily_df.index.day[0]
    dic['end_year']=daily_df.index.year[-1]
    dic['end_month']=daily_df.index.month[-1]
    dic['end_day']=daily_df.index.day[-1]
    
    monthly_means = monthly_df[['TEM_Avg']].groupby(monthly_df[['TEM_Avg']].index.month).mean().round(2)
    sorted_monthly_means = monthly_means.sort_values('TEM_Avg',ascending=False).reset_index()
      
    dic['heat_month']=sorted_monthly_means.iloc[0,0]
    dic['heat_month_1']=sorted_monthly_means.iloc[0,0]
    dic['heat_month_2']=sorted_monthly_means.iloc[1,0]
    dic['heat_month_3']=sorted_monthly_means.iloc[2,0]
    dic['cold_month']=sorted_monthly_means.iloc[-1,0]
    dic['cold_month_1']=sorted_monthly_means.iloc[-1,0]
    dic['cold_month_2']=sorted_monthly_means.iloc[-2,0]
    dic['cold_month_3']=sorted_monthly_means.iloc[-3,0]
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'heating_and_ventilation.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    enthalpy=pd.DataFrame(enthalpy)
    enthalpy=enthalpy.T.reset_index()
    enthalpy_result=pd.DataFrame(columns=['时刻']+[str(num)+':00' for num in np.arange(12)])
    columns=['时刻']+[str(num)+':00' for num in np.arange(12)]
    for num in np.arange(13):
        enthalpy_result.at[0,columns[num]]=enthalpy.iloc[1,num]
        if num ==0:
            enthalpy_result.at[1,columns[num]]='时刻'
        else:
            enthalpy_result.at[1,columns[num]]=str(num+11)+':00'
        enthalpy_result.at[2,columns[num]]=enthalpy.iloc[1,num+12]

    enthalpy_result.iloc[0,0]='焓值'
    enthalpy_result.iloc[2,0]='焓值'
    
    t_sh=pd.DataFrame(t_sh)
    t_sh=t_sh.T.reset_index()
    t_sh_result=pd.DataFrame(columns=['时刻']+[str(num)+':00' for num in np.arange(12)])
    columns=['时刻']+[str(num)+':00' for num in np.arange(12)]
    for num in np.arange(13):
        t_sh_result.at[0,columns[num]]=t_sh.iloc[1,num]
        if num ==0:
            t_sh_result.at[1,columns[num]]='时刻'
        else:
            t_sh_result.at[1,columns[num]]=str(num+11)+':00'
        t_sh_result.at[2,columns[num]]=t_sh.iloc[1,num+12]

    t_sh_result.iloc[0,0]='温度'
    t_sh_result.iloc[2,0]='温度'
        
    creat_table(document,df,'年）')
    creat_table(document,enthalpy_result,'kJ/kg）')
    creat_table(document,t_sh_result,'℃）')

    document.save(report)
    
    
    return report

if __name__ == '__main__':
    pass
    # df=df_dict
    # data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module04'
    # daily_df=post_daily_df
    # monthly_df=post_monthly_df
    # report= heating_and_ventilation_report(df,t_sh,enthalpy,daily_df,monthly_df,data_dir)