# -*- coding: utf-8 -*-
"""
Created on Thu Jun 13 14:51:03 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'

def water_supply_report(result,df_day,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','water_supply.docx')
    doc=DocxTemplate(doc_path)
    

    dic=dict()
    dic['station_name']=df_day['Station_Name'][0]
    dic['start_year']=df_day.index.year[0]
    dic['end_year']=df_day.index.year[-1]
    dic['num_years']=len(df_day.index.year.unique())

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'water_supply.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据

    result=dict(result)
    result1=pd.DataFrame(result['result1'])
    result2=pd.DataFrame(result['result2'])
    result3=pd.DataFrame(result['result3'])
    result4=pd.DataFrame(result['result4'])
    result5=pd.DataFrame(result['result5'])
    result6=pd.DataFrame(result['result6'])
    
    result1.columns=['日平均温度(°C)','日平均气压(hap)','日平均相对湿度','湿球温度（°C）','10%累积频率温度','日期']
    result2.columns=['日平均温度(°C)','日平均气压(hap)','日平均相对湿度','湿球温度（°C）','5%累积频率温度','日期']
    result3.columns=['日平均温度(°C)','日平均气压(hap)','日平均相对湿度','湿球温度（°C）','1%累积频率温度','日期']
    result4.columns=['湿球温度（°C）','干球温度（°C）','气压（hPa）','湿球温度（°C）','99%累积频率温度（°C）','日期']


    half_length =int(np.ceil(len(result5)/2))
      
    # 分割每个列并创建新列
    result5['A1'] = result5.iloc[:half_length:,0]
    result5['A2'] =result5.iloc[:half_length:,1]
    result5['A3'] = result5.iloc[:half_length:,2]
    result5['A4'] = result5.iloc[:half_length:,3]
    result5['B1'] =result5.iloc[half_length::,0].reset_index(drop=True)
    result5['B2'] =result5.iloc[half_length::,1].reset_index(drop=True)
    result5['B3'] =result5.iloc[half_length::,2].reset_index(drop=True)
    result5['B4'] =result5.iloc[half_length::,3].reset_index(drop=True)

    result5.drop(['date', 'TEM_Avg','PRS_Avg','RHU_Avg'], axis=1, inplace=True)
    
    result5.dropna(how='all',inplace=True)
    result5.columns=['日期','日平均温度(°C)','日平均气压(hap)','日平均相对湿度','日期','日平均温度(°C)','日平均气压(hap)','日平均相对湿度']
     

    half_length =int(np.ceil(len(result6)/2))
      
    # 分割每个列并创建新列
    result6['A1'] = result6.iloc[:half_length:,0]
    result6['A2'] =result6.iloc[:half_length:,1]
    result6['A3'] = result6.iloc[:half_length:,2]
    result6['A4'] = result6.iloc[:half_length:,3]
    result6['B1'] =result6.iloc[half_length::,0].reset_index(drop=True)
    result6['B2'] =result6.iloc[half_length::,1].reset_index(drop=True)
    result6['B3'] =result6.iloc[half_length::,2].reset_index(drop=True)
    result6['B4'] =result6.iloc[half_length::,3].reset_index(drop=True)

    result6.drop(['date', 'TEM_Avg','PRS_Avg','RHU_Avg'], axis=1, inplace=True)
    
    result6.dropna(how='all',inplace=True)
    result6.columns=['日期','日平均温度(°C)','日平均气压(hap)','日平均相对湿度','日期','日平均温度(°C)','日平均气压(hap)','日平均相对湿度']
  

    creat_table(document,result3,'夏季1%累积频率的平均气温，及对应时间/真实气温/气压/湿度/湿球温度')
    creat_table(document,result2,'夏季5%累积频率的平均气温，及对应时间/真实气温/气压/湿度/湿球温度')
    creat_table(document,result1,'夏季10%累积频率的平均气温，及对应时间/真实气温/气压/湿度/湿球温度')
    creat_table(document,result4,'冬季最冷三个月累积频率为99%日平均气温，及对应时间/真实气温/气压/湿度/湿球温度')
    creat_table(document,result5,'最热三个月温度气压湿度')
    creat_table(document,result6,'最冷三个月温度气压湿度')


    document.save(report)
    
    
    return report
