# -*- coding: utf-8 -*-
"""
Created on Wed Jun 12 09:39:13 2024

@author: EDY
"""


import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'
def rain_runoff_report(pre, pre_points, table, save_path1, post_daily_df, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    pre=pd.DataFrame(pre)
    pre_points=pd.DataFrame(pre_points)
    table=pd.DataFrame(table)

    doc_path=os.path.join(cfg['report']['template'],'Module05','rain_runoff.docx')
    doc=DocxTemplate(doc_path)
    
    
    dic=dict()
    dic['station_name']=post_daily_df.iloc[0,1]
    dic['start_year']=post_daily_df.index.year[0]
    dic['end_year']=post_daily_df.index.year[-1]

    plt.figure(figsize=(10, 6))
    plt.plot(pre['降水事件'], pre['降水量(mm)'],color='blue')
    
    plt.grid(axis='y', linestyle='--', alpha=0.7)    
    plt.xlabel('降水事件')
    plt.ylabel('降水量(mm)')
    save_path = data_dir + '/降水事件_降水量.png'
    plt.savefig(save_path, dpi=200, bbox_inches='tight')
    plt.clf()
    plt.close('all')
    dic['picture_1'] = InlineImage(doc, save_path, width=Mm(130))
    dic['picture_2'] = InlineImage(doc, save_path1, width=Mm(130))

        
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'rain_runoff.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    table['年径流总量控制率(%)'] = table['年径流总量控制率(%)'].astype(int)

    half_length =int(np.ceil(len(table)/4))
   
    # 分割每个列并创建新列
    table['A1'] = table.iloc[:half_length:,0]
    table['B1'] =table.iloc[:half_length:,1]
    table['A2'] = table.iloc[half_length:half_length*2:,0].reset_index(drop=True)
    table['B2'] =table.iloc[half_length:half_length*2:,1].reset_index(drop=True)
    table['A3'] =table.iloc[half_length*2:half_length*3:,0].reset_index(drop=True)
    table['B3'] = table.iloc[half_length*2:half_length*3:,1].reset_index(drop=True)
    table['A4'] =table.iloc[half_length*3:half_length*4:,0].reset_index(drop=True)
    table['B4'] = table.iloc[half_length*3:half_length*4:,1].reset_index(drop=True)
    table.drop(['年径流总量控制率(%)', '设计降雨量(mm)'], axis=1, inplace=True)
    
    table.dropna(how='all',inplace=True)
    table.columns=['雨水年径流总量控制率Alpha(%)','设计降雨深度H','雨水年径流总量控制率Alpha','设计降雨深度H','雨水年径流总量控制率Alpha','设计降雨深度H','雨水年径流总量控制率Alpha','设计降雨深度H']
     
    
    
    creat_table(document,pre_points,'降水事件和降水量关系统计表')
    creat_table(document,table,'雨水年径流总量控制率与设计降雨深度统计表')

    document.save(report)
    
    
    return report
