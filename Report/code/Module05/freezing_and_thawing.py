# -*- coding: utf-8 -*-
"""
Created on Wed Jun 12 10:59:09 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Report.code.Function.plot_picture import plot_picture_2

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'
def freezing_and_thawing_report(num,result2, post_daily_df, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','freezing_and_thawing_2.docx')
    doc=DocxTemplate(doc_path)
    
    result2=pd.DataFrame(result2)
    result2=result2.dropna()

    dic=dict()
    dic['station_name']=post_daily_df.iloc[0,1]


    if num==1:
        dic['description']='融交替循环次数'
        dic['description_1']='基于《核电厂工程气象技术规范》GRIT 50674—2013，判定标准为从逐时的干球温度中计算，逐时气温从 3℃以上降至-3℃以下，然后再回升到 3℃ 以上算1次冻融交替循环'
    elif num==2:
        dic['description']='冻融循环日数'
        dic['description_1']='基于《军用设备气候极值》GJB 1172.11-1991，判定标准为以一年内的逐日的日最低气温判断，若日最低气温从+3℃以上降至-3℃以下，然后回升到+3℃以上的过程'
    elif num==3:
        dic['description']='冻融次数'
        dic['description_1']='基于《水工建筑物抗冻设计规范》SL 211-98计算'

    save_path=plot_picture_2(result2['时间段'],result2['冻融交替次数'],dic,'冻融交替次数',1,'average_pre_m2','冻融交替次数.png',100,2,data_dir)

    # plt.figure(figsize=(10, 6))
    # plt.bar(result2['时间段'], result2['冻融交替次数'], width=0.4, color='skyblue')
    
    # plt.grid(axis='y', linestyle='--', alpha=0.7)    
    # plt.xlabel('时间段')
    # plt.ylabel('冻融交替次数')
    # plt.xticks(rotation=90)
    # plt.tight_layout()
    # save_path = data_dir + '/冻融交替次数.png'
    # plt.savefig(save_path, dpi=200, bbox_inches='tight')
    # plt.clf()
    # plt.close('all')
    dic['picture_1'] = InlineImage(doc, save_path, width=Mm(130))

        
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'freezing_and_thawing_'+str(num)+'.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据
    half_length =int(np.ceil(len(result2)/2))
   
    # 分割每个列并创建新列
    result2['A1'] = result2.iloc[:half_length:,0]
    result2['B1'] =result2.iloc[:half_length:,1]
    result2['A2'] = result2.iloc[half_length::,0].reset_index(drop=True)
    result2['B2'] =result2.iloc[half_length::,1].reset_index(drop=True)

    result2.drop(['时间段', '冻融交替次数'], axis=1, inplace=True)
    
    result2.dropna(how='all',inplace=True)
    result2.columns=['时间段(%)','冻融交替次数','时间段','冻融交替次数']
     
    creat_table(document,result2,'次数表')

    document.save(report)
    
    
    return report
