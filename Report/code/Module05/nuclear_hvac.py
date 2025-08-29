# -*- coding: utf-8 -*-
"""
Created on Tue Jun 11 17:44:14 2024

@author: EDY
"""

import pandas as pd
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        # print(paragraph_text)
        # print('----------')
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Module05'
# df_hour=hourly_df
def nuclear_hvac_report(result,df_hour,data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module05','nuclear_hvac.docx')
    doc=DocxTemplate(doc_path)
    
    result=dict(result)
    
    dic=dict()
    dic['num_years']=len(df_hour.index.year.unique())
    dic['start_year']=df_hour.index.year[0]
    dic['start_month']=df_hour.index.month[0]
    dic['start_day']=df_hour.index.day[0]
    dic['end_year']=df_hour.index.year[-1]
    dic['end_month']=df_hour.index.month[-1]
    dic['end_day']=df_hour.index.day[-1]
    
    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'nuclear_hvac.docx')
    doc.save(report)
    
    ## 插入表格
    document = Document(report)
    
    # 填充表格数据

    result_1=pd.DataFrame(columns=['项目','不保证2小时'])
    result_1.at[0,'项目']='最高安全设计干球温度(TCS)'
    result_1.at[1,'项目']='最高安全设计湿球温度(TCH)'
    result_1.at[2,'项目']='最低安全设计干球温度(TDS)'
    result_1.at[3,'项目']='最低安全设计湿球温度(TDH)'
    result_1.at[0,'不保证2小时']=result['TCS']
    result_1.at[1,'不保证2小时']=result['TCH']
    result_1.at[2,'不保证2小时']=result['TDS']
    result_1.at[3,'不保证2小时']=result['TDH']
    
    result_2=pd.DataFrame(result['result3'].copy())
    result_2.columns=['出现时间','最高安全设计干球温度(TCS)','最高安全设计湿球温度(TCH)(对应)']
    
    result_3=pd.DataFrame(result['result4'].copy())
    result_3.columns=['出现时间','最低安全设计干球温度(TDS)','最低安全设计湿球温度(TDH)(对应)']
    
    result_4=pd.DataFrame(columns=['项目','不保证1%'])
    result_4.at[0,'项目']='最高正常设计干球温度(TAS)'
    result_4.at[1,'项目']='最高正常设计湿球温度(TAH)'
    result_4.at[2,'项目']='最低正常设计干球温度(TBS)'
    result_4.at[3,'项目']='最低正常设计湿球温度(TBH)'
    result_4.at[0,'不保证1%']=result['TAS']
    result_4.at[1,'不保证1%']=result['TAH']
    result_4.at[2,'不保证1%']=result['TBS']
    result_4.at[3,'不保证1%']=result['TBH']
    
    result_5=pd.DataFrame(result['result1'].copy())
    result_5.columns=['出现时间','最高正常设计干球温度(TAS)','最高正常设计湿球温度(TAH)']
    
    result_6=pd.DataFrame(result['result2'].copy())
    result_6.columns=['出现时间','最低正常设计干球温度(TBS)','最低正常设计湿球温度(TBH)']

    result_7=pd.DataFrame(result['result5'].copy())
    result_7.columns=['出现时间','干球温度(℃)','湿球温度(℃)']

    result_8=pd.DataFrame(result['result6'].copy())
    result_8.columns=['出现时间','干球温度(℃)','湿球温度(℃)']    
    
    
    creat_table(document,result_1,'最高和最低安全设计温度(℃)')
    creat_table(document,result_2,'最高安全设计干球温度对应湿球温度(℃)及出现时间')
    creat_table(document,result_3,'最低安全设计干球温度对应湿球温度(℃)及出现时间')
    creat_table(document,result_4,'最高和最低正常设计温度(℃)')
    creat_table(document,result_5,'最高正常设计干球温度对应湿球温度(℃)及出现时间')
    creat_table(document,result_6,'最低正常设计干球温度对应湿球温度(℃)及出现时间')
    creat_table(document,result_7,'不保证5%最高干球温度对应湿球温度(℃)及出现时间')
    creat_table(document,result_8,'表8 不保证5%最低干球温度对应湿球温度(℃)及出现时间')


    document.save(report)
    
    
    return report
