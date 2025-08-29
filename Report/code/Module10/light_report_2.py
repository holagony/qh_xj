# -*- coding: utf-8 -*-
"""
Created on Thu Jul 18 13:16:29 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
import os
from Utils.config import cfg
import numpy as np
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text

        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules10'
def light_report_2(result,factors, data_dir):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        
    doc_path=os.path.join(cfg['report']['template'],'Module10','light_report_2.docx')
    doc=DocxTemplate(doc_path)
    
    data=pd.DataFrame(result['各因子隶属度结果'])
    data.set_index(data.columns[0], inplace=True)

    data=data.T
    name=['雷击密度','雷电流强度','土壤电阻率','土壤垂直分层','土壤水平分层','地形地貌','安全距离','相对高度','电磁环境',
          '使用性质','人员数量','影响程度','占地面积','材料结构','等效高度','电子系统','电气系统','土壤结构','周边环境',
          '项目属性','建筑特征','电子电气系统','雷电风险','地域风险','承载体风险']
    
    dic=dict()

    for i in np.arange(len(name)):
        dic['b_'+str(i+1)+'_1']=round(data[name[i]][0],2)
        dic['b_'+str(i+1)+'_2']=round(data[name[i]][1],2)
        dic['b_'+str(i+1)+'_3']=round(data[name[i]][2],2)
        dic['b_'+str(i+1)+'_4']=round(data[name[i]][3],2)
        dic['b_'+str(i+1)+'_5']=round(data[name[i]][4],2)
        dic['b_'+str(i+1)+'_6']=round(data[name[i]][5],2)

    data_2=pd.DataFrame(result['区域雷电灾害风险'])
    i=25
    dic['b_'+str(i+1)+'_1']=round(data_2.iloc[0,0],2)
    dic['b_'+str(i+1)+'_2']=round(data_2.iloc[0,1],2)
    dic['b_'+str(i+1)+'_3']=round(data_2.iloc[0,2],2)
    dic['b_'+str(i+1)+'_4']=round(data_2.iloc[0,3],2)
    dic['b_'+str(i+1)+'_5']=round(data_2.iloc[0,4],2)
    
    dic['b_27']=result['雷电灾害风险值']
    dic['b_29']=result['雷电灾害风险等级']
    
    dic['data_1']=round(factors['雷电流强度']['value'],2)
    dic['data_2']=round(factors['土壤电阻率']['value'],2)
    dic['data_3']=round(factors['土壤垂直分层']['value'],2)
    dic['data_4']=round(factors['土壤水平分层']['value'],2)
    
    name_5=['平原','丘陵','山地','河流、湖泊以及低洼潮湿地区、山间风口等','旷野孤立或突出区域']
    dic['data_5']=name_5[int(factors['地形地貌']['value']-1)]

    name_6=['1km范围内无易燃易爆场所','500m范围内存在易燃易爆场所','300m范围内存在易燃易爆场所','100m范围内存在易燃易爆场所','100m范围内无易燃易爆场所且后果严重']
    dic['data_6']=name_6[int(factors['安全距离']['value']-1)]
    
    name_7=['评估区域被比区域内项目高的外部建(构)筑物或其他雷击可接闪物所环绕','评估区域外局部方向有高于评估区域内项目的建(构)筑物或其他雷击可接闪物',
            '评估区域外建(构)筑物或其他雷击可接闪物与评估区域内项目高度基本持平','评估区域外建(构)筑物或其他雷击可接闪物低于区域内项目高度',
            '评估区域外无建(构)筑物或其他雷击可接闪物']
    dic['data_7']=name_7[int(factors['相对高度']['value']-1)]
    dic['data_8']=round(factors['电磁环境']['value'],2)
    dic['data_9']=int(factors['电磁环境']['value'])
    dic['data_10']=int(factors['人员数量']['value'])
    dic['data_11']=int(factors['影响程度']['value'])
    dic['data_12']=round(factors['占地面积']['value'],2)

    name_13=['木结构','评砖木结构','砖混结构','钢筋混凝土结构','评钢结构']
    dic['data_13']=name_13[int(factors['相对高度']['value']-1)]
    dic['data_14']=round(factors['等效高度']['value'],2)
    
    dic['data_15']=int(factors['电气系统']['value'])

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'light_report_2.docx')
    doc.save(report)
    
    
    
    return report
