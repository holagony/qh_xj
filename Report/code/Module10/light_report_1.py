# -*- coding: utf-8 -*-
"""
Created on Mon Jul  1 14:55:03 2024

@author: EDY
"""

import matplotlib
matplotlib.use('Agg')
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import os
from Utils.config import cfg
from docx.shared import Mm
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False 

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def creat_table(document,data,expect_text):
    data =data.astype(str)
    data = data.transpose()
    data=data.reset_index()
    data =data.transpose()
    
    table = document.add_table(rows=data.shape[0], cols=data.shape[1])
    
    for i in range(data.shape[0]):
        row = table.rows[i]
        for j in range(data.shape[1]):
            cell = row.cells[j]
            cell.text = data.iloc[i,j]
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
    # 设置表格样式
    table.style = document.styles['Table Grid']
    table.autofit = True
    table.allow_autofit = True

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text

        if paragraph_text.endswith(expect_text):
            target = paragraph
            break
    
    move_table_after(table, target)


# data_dir=r'D:\Project\3_项目\2_气候评估和气候可行性论证\qhkxxlz\Report\report\Modules10'
def light_report_1(result, light_density_picture,data_dir,lonlat=None, save_path_picture_p=None):

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    if lonlat is not None:
        doc_path=os.path.join(cfg['report']['template'],'Module10','light_report_1.docx')
    else:
        doc_path=os.path.join(cfg['report']['template'],'Module10','light_report_12.docx')
        
    doc=DocxTemplate(doc_path)
    
    
    dic=dict()

    # 年次数
    year_f=pd.DataFrame(result['次数统计']['年'])
    row_sum = year_f[year_f.columns].assign(Z=year_f['正闪数量'] + year_f['负闪数量'])

    dic['start_year']=year_f['年份'].iloc[0]
    dic['end_year']=year_f['年份'].iloc[-1]
    dic['total_light']=int(row_sum['Z'].sum())
    dic['max_light_year']=year_f['年份'].iloc[row_sum['Z'].idxmax()]
    dic['max_lighty']=row_sum['Z'].max()
    dic['min_light_year']=year_f['年份'].iloc[row_sum['Z'].idxmin()]
    dic['min_lighty']=row_sum['Z'].min()
    dic['picture_1']=InlineImage(doc, result['图片路径']['年次数'], width=Mm(130))

    # 年内次数
    month_f=pd.DataFrame(result['次数统计']['月'])
    row_sum =  month_f[month_f.columns].assign(Z=month_f['正闪数量'] + month_f['负闪数量'])
    dic['max_light_monthe']=month_f['月份'].iloc[row_sum['Z'].idxmax()]
    dic['max_lightm']=int(row_sum['Z'].max())
    dic['picture_2']=InlineImage(doc, result['图片路径']['月次数'], width=Mm(130))

    # 日次数
    hour_f=pd.DataFrame(result['次数统计']['小时'])
    row_sum = hour_f[hour_f.columns].assign(Z=hour_f['正闪数量'] + hour_f['负闪数量'])
    
    dic['num_years']=len(year_f['年份'])
    dic['max_light_hour']=hour_f['小时'].iloc[row_sum['Z'].idxmax()]
    dic['max_lighth']=int(row_sum['Z'].max())
    dic['min_light_hour']=hour_f['小时'].iloc[row_sum['Z'].idxmin()]
    dic['min_lighth']=int(row_sum['Z'].min())
    dic['picture_3']=InlineImage(doc, result['图片路径']['小时次数'], width=Mm(130))

    # 年日数
    year_f=pd.DataFrame(result['天数统计']['年'])
    row_sum = year_f[year_f.columns].assign(Z=year_f['正闪日数'] + year_f['负闪日数'])
    dic['mean_light_days']=round(row_sum['Z'].mean(),2)
    dic['max_light_days_years']=year_f['年份'].iloc[row_sum['Z'].idxmax()]
    dic['max_light_days']=row_sum['Z'].max()
    dic['min_light_days_years']=year_f['年份'].iloc[row_sum['Z'].idxmin()]
    dic['min_light_days']=row_sum['Z'].min()
    dic['picture_4']=InlineImage(doc, result['图片路径']['年天数'], width=Mm(130))
    
    # 年内日数
    month_f=pd.DataFrame(result['天数统计']['月'])
    row_sum = month_f[month_f.columns].assign(Z=month_f['正闪日数'] + month_f['负闪日数'])
    dic['max_light_day_monthe']=month_f['月份'].iloc[row_sum['Z'].idxmax()]
    dic['max__day_lightm']=row_sum['Z'].max()
    dic['min_light_day_monthe']=month_f['月份'].iloc[row_sum['Z'].idxmin()]
    dic['min__day_lightm']=row_sum['Z'].min()    
    dic['picture_5']=InlineImage(doc, result['图片路径']['月天数'], width=Mm(130))

    
    dic['data_1']=result['累积概率曲线']['a'] 
    dic['data_2']=result['累积概率曲线']['b']
    dic['picture_7']=InlineImage(doc, result['图片路径']['累积概率曲线'], width=Mm(130))
    dic['picture_6']=InlineImage(doc, light_density_picture, width=Mm(130))

    if lonlat is not None:
        lonlat=pd.DataFrame(lonlat)
        dic['data_3']='A/m、'.join(map(str,lonlat['电磁强度']))
        dic['picture_8']=InlineImage(doc, save_path_picture_p, width=Mm(130))

    # 模版文件读取写入字典
    doc.render(dic)
    # 保存结果到新的docx文件
    report=os.path.join(data_dir,'light_report_1.docx')
    doc.save(report)
    
    
    
    return report
