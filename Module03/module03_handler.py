import os
import logging
import uuid
import numpy as np
import pandas as pd
from collections import OrderedDict
from Module00.wrapped.check import check
from Module03.wrapped.table_stats_part1 import table_stats_part1
from Module03.wrapped.table_stats_part2 import table_stats_part2
from Module03.wrapped.weather_stats import all_weather_statistics_accum
from Module03.wrapped.cold_wave_stats import cold_wave_statistics
from Module03.wrapped.cold_rainy_days_stats import cold_rainy_days_statistics
from Module03.wrapped.cold_freezing_days_stats import cold_freeing_days_statistics
from Utils.data_processing import monthly_data_processing, daily_data_processing
from Utils.get_url_path import get_url_path
from Utils.get_local_data import get_local_data
from Utils.config import cfg
from Utils.ordered_easydict import OrderedEasyDict as edict
from Utils.data_loader_with_threads import get_cmadaas_monthly_data, get_cmadaas_daily_data

# report
from Report.code.Module03.pre_days import pre_days_report
from Report.code.Module03.hail_days import hail_days_report
from Report.code.Module03.frost_init_and_end_days_report import frost_init_and_end_days_report
from Report.code.Module03.cold_rainy_freez_report import cold_rainy_freez_report
from Report.code.Module03.cold_wave import cold_wave_report
from Report.code.Module03.drsnow_days import drsnow_days_report
from Report.code.Module03.duwhr_days import duwhr_days_report
from Report.code.Module03.fldu_days import fldu_days_report
from Report.code.Module03.flsa_days import flsa_days_report
from Report.code.Module03.fog_days import fog_days_report
from Report.code.Module03.frost_days import frost_days_report
from Report.code.Module03.gawin_days import gawin_days_report
from Report.code.Module03.glaze_days import glaze_days_report
from Report.code.Module03.gss_days import gss_days_report
from Report.code.Module03.gss_init_and_end_days_report import gss_init_and_end_days_report
from Report.code.Module03.gst_under0_init_and_end_days_report import gst_under0_init_and_end_days_report
from Report.code.Module03.haze_days import haze_days_report
from Report.code.Module03.ice_days import ice_days_report
from Report.code.Module03.ice_init_and_end_days_report import ice_init_and_end_days_report
from Report.code.Module03.lit_days import lit_days_report
from Report.code.Module03.mist_days import mist_days_report
from Report.code.Module03.sast_days import sast_days_report
from Report.code.Module03.snow_days import snow_days_report
from Report.code.Module03.snow_init_and_end_days_report import snow_init_and_end_days_report
from Report.code.Module03.sori_days import sori_days_report
from Report.code.Module03.squa_days import squa_days_report
from Report.code.Module03.tem_under0_init_and_end_days_report import tem_under0_init_and_end_days_report
from Report.code.Module03.thund_days import thund_days_report
from Report.code.Module03.thund_init_and_end_days_report import thund_init_and_end_days_report
from Report.code.Module03.tord_days import tord_days_report
from Report.code.Module03.tem_max_days_report import tem_max_days_report
from Report.code.Module03.tem_min_days_report import tem_min_days_report
from Report.code.Module03.pre_strong_report import pre_strong_report
from Report.code.Module03.eice_days import eice_days_report

from docx import Document
from docxcompose.composer import Composer
from Utils.get_url_path import save_cmadaas_data

def weather_phenomena_days(data_json):
    '''
    天气现象日数统计组件
    '''
    result_dict = edict()

    # 1.参数读取
    years = data_json['years']  # 选择的数据年份
    main_sta_ids = data_json['main_sta_ids']  # 主站
    sub_sta_ids = data_json.get('sub_sta_ids')  # 对比站
    elements = data_json['elements']  # 选择的气象要素

    # 2.参数处理
    uuid4 = uuid.uuid4().hex
    result_dict['uuid'] = uuid4

    data_dir = os.path.join(cfg.INFO.IN_DATA_DIR, uuid4)
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        os.chmod(data_dir, 0o007 | 0o070 | 0o700)

    if isinstance(main_sta_ids, int):
        main_sta_ids = str(main_sta_ids)

    if isinstance(sub_sta_ids, list):
        sub_sta_ids = [str(ids) for ids in sub_sta_ids]
        sub_sta_ids = ','.join(sub_sta_ids)

    elif isinstance(sub_sta_ids, int):  # 只有单个值的情况
        sub_sta_ids = str(sub_sta_ids)

    # 3.拼接需要下载的参数
    monthly_elements = ''
    daily_elements = ''
    daily_df = None
    monthly_df = None

    for ele in elements:
        if ele == 'PRE_Days':
            monthly_elements += 'PRE_Days,'

        elif ele == 'Hail_Days':
            monthly_elements += 'Hail_Days,'

        elif ele == 'Fog_Days':
            monthly_elements += 'Fog_Days,'

        elif ele == 'Mist_Days':
            monthly_elements += 'Mist_Days,'

        elif ele == 'Glaze_Days':
            monthly_elements += 'Glaze_Days,'

        elif ele == 'Tord_Days':
            monthly_elements += 'Tord_Days,'

        elif ele == 'SoRi_Days':
            monthly_elements += 'SoRi_Days,'

        elif ele == 'SaSt_Days':
            monthly_elements += 'SaSt_Days,'

        elif ele == 'FlSa_Days':
            monthly_elements += 'FlSa_Days,'

        elif ele == 'FlDu_Days':
            monthly_elements += 'FlDu_Days,'

        elif ele == 'Haze_Days':
            monthly_elements += 'Haze_Days,'

        elif ele == 'GaWIN_Days':
            monthly_elements += 'GaWIN_Days,'
        
        elif ele == 'EICE_Days':
            monthly_elements += 'EICE_Days,'

        elif ele == 'Squa':
            daily_elements += 'Squa,'

        elif ele == 'Lit':
            daily_elements += 'Lit,'

        elif ele == 'DuWhr':
            daily_elements += 'DuWhr,'

        elif ele == 'DrSnow':
            daily_elements += 'DrSnow,'

        elif ele == 'Snow':
            daily_elements += 'Snow,'

        elif ele == 'Frost':
            daily_elements += 'Frost,'

        elif ele == 'GSS':
            daily_elements += 'GSS,'

        elif ele == 'ICE':
            daily_elements += 'ICE,'

        elif ele == 'Thund':
            daily_elements += 'Thund,'
        elif ele == 'TEM_Max':
            daily_elements += 'TEM_Max,'
        elif ele == 'TEM_Min':
            daily_elements += 'TEM_Min,'
        elif ele == 'PRE_Time_2020':
            daily_elements += 'PRE_Time_2020,'
            
            
    # 4.数据获取
    month_ele_list = ['PRE_Days', 'Hail_Days', 'Fog_Days', 'Mist_Days', 'Glaze_Days', 'Tord_Days', 'SoRi_Days', 'SaSt_Days', 'FlSa_Days', 'FlDu_Days', 'Haze_Days', 'GaWIN_Days', 'EICE_Days']
    day_ele_list = ['Squa', 'Lit', 'DuWhr', 'DrSnow', 'Snow', 'Frost', 'GSS', 'ICE', 'Thund','TEM_Max','TEM_Min','PRE_Time_2020']
    
    if cfg.INFO.READ_LOCAL:

        if sub_sta_ids == None:
            sta_ids = main_sta_ids
        else:
            sta_ids = main_sta_ids + ',' + sub_sta_ids

        monthly_elements_tmp = monthly_elements[:-1].split(',')
        if set(month_ele_list) & set(monthly_elements_tmp):
            
            monthly_df = pd.read_csv(cfg.FILES.QH_DATA_MONTH, low_memory=False)
            if 'EICE_Days' in monthly_elements:
                monthly_df['EICE_Days']=15
            
            month_eles = ('Station_Id_C,Station_Name,Lat,Lon,Datetime,Year,Mon,' + monthly_elements[:-1]).split(',')
            monthly_df = get_local_data(monthly_df, sta_ids, month_eles, years, 'Month')

        daily_elements_tmp = daily_elements[:-1].split(',')
        if set(day_ele_list) & set(daily_elements_tmp):
            daily_df = pd.read_csv(cfg.FILES.QH_DATA_DAY)
            day_eles = ('Station_Id_C,Station_Name,Lat,Lon,Datetime,Year,Mon,Day,' + daily_elements[:-1]).split(',')
            daily_df = get_local_data(daily_df, sta_ids, day_eles, years, 'Day')

        # daily_df = daily_df[~daily_df['Station_Name'].isna()]  # 测试数据问题 去掉station_name为null的行

    else:
        # 天擎数据下载 and 数据前处理
        try:
            if sub_sta_ids is not None:
                sta_ids = main_sta_ids + ',' + sub_sta_ids
            else:
                sta_ids = main_sta_ids

            if len(set(month_ele_list) & set(elements)) != 0:
                monthly_df = get_cmadaas_monthly_data(years, monthly_elements, sta_ids)
                monthly_df = monthly_data_processing(monthly_df, years)

            # 日数据
            if len(set(day_ele_list) & set(elements)) != 0:
                daily_df = get_cmadaas_daily_data(years, daily_elements, sta_ids)
                daily_df = daily_data_processing(daily_df, years)
            else:
                daily_df = None

        except Exception as e:
            raise Exception('天擎数据下载或处理失败')

    # 5.计算之前先检测数据完整率 check H小时 D天 MS月 YS年
    years = years.split(',')
    result_dict.check_result = edict()
    monthly_elements = monthly_elements[:-1]
    daily_elements = daily_elements[:-1]

    if monthly_df is not None and len(monthly_df) != 0:
        checker = check(monthly_df, 'MS', monthly_elements.split(','), sta_ids.split(','), years[0], years[1])
        check_result = checker.run()
        result_dict.check_result['使用的天擎月要素'] = check_result

    if daily_df is not None and len(daily_df) != 0:
        checker = check(daily_df, 'D', daily_elements.split(','), sta_ids.split(','), years[0], years[1])
        check_result = checker.run()
        result_dict.check_result['使用的天擎日要素'] = check_result

    # 6.结果生成
    ele_list = []  # 记录出现的要素，用于累年各月总表合成(总表1)
    result_list = []  # 用于将保存的csv，进行url转换，然后输出路径
    result_path = []
    for ele in elements:
        if ele == 'PRE_Days':
            result_dict.PRE_Days = edict()
            ele_list.append('PRE_Days')
            PRE_Days_tab1, PRE_Days_tab2, PRE_Days_tab3 = table_stats_part1(monthly_df, 'PRE_Days')
            result_dict.PRE_Days['table1'] = PRE_Days_tab1
            result_dict.PRE_Days['table2'] = PRE_Days_tab2
            result_dict.PRE_Days['table3'] = PRE_Days_tab3

            try:
                report_path = pre_days_report(PRE_Days_tab1, PRE_Days_tab2, PRE_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.PRE_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.PRE_Days['report'] = None

            result_list.append(OrderedDict(zip(['历年雨日统计', '累年雨日统计', '累年各月雨日统计'], [PRE_Days_tab1, PRE_Days_tab2, PRE_Days_tab3])))

        elif ele == 'Hail_Days':
            result_dict.Hail_Days = edict()
            ele_list.append('Hail_Days')
            Hail_Days_tab1, Hail_Days_tab2, Hail_Days_tab3 = table_stats_part1(monthly_df, 'Hail_Days')
            result_dict.Hail_Days['table1'] = Hail_Days_tab1
            result_dict.Hail_Days['table2'] = Hail_Days_tab2
            result_dict.Hail_Days['table3'] = Hail_Days_tab3

            try:
                report_path = hail_days_report(Hail_Days_tab1, Hail_Days_tab2, Hail_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Hail_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Hail_Days['report'] = None

            result_list.append(OrderedDict(zip(['历年冰雹日统计', '累年冰雹日统计', '累年各月冰雹日统计'], [Hail_Days_tab1, Hail_Days_tab2, Hail_Days_tab3])))

        elif ele == 'Fog_Days':
            result_dict.Fog_Days = edict()
            ele_list.append('Fog_Days')
            Fog_Days_tab1, Fog_Days_tab2, Fog_Days_tab3 = table_stats_part1(monthly_df, 'Fog_Days')
            result_list.append(OrderedDict(zip(['历年雾日统计', '累年雾日统计', '累年各月雾日统计'], [Fog_Days_tab1, Fog_Days_tab2, Fog_Days_tab3])))

            result_dict.Fog_Days['table1'] = Fog_Days_tab1
            result_dict.Fog_Days['table2'] = Fog_Days_tab2
            result_dict.Fog_Days['table3'] = Fog_Days_tab3

            try:
                report_path = fog_days_report(Fog_Days_tab1, Fog_Days_tab2, Fog_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Fog_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Fog_Days['report'] = None

        elif ele == 'Mist_Days':
            result_dict.Mist_Days = edict()
            ele_list.append('Mist_Days')
            Mist_Days_tab1, Mist_Days_tab2, Mist_Days_tab3 = table_stats_part1(monthly_df, 'Mist_Days')
            result_list.append(OrderedDict(zip(['历年轻雾日统计', '累年轻雾日统计', '累年各月轻雾日统计'], [Mist_Days_tab1, Mist_Days_tab2, Mist_Days_tab3])))

            result_dict.Mist_Days['table1'] = Mist_Days_tab1
            result_dict.Mist_Days['table2'] = Mist_Days_tab2
            result_dict.Mist_Days['table3'] = Mist_Days_tab3

            try:
                report_path = mist_days_report(Mist_Days_tab1, Mist_Days_tab2, Mist_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)

                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Mist_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Mist_Days['report'] = None

        elif ele == 'Glaze_Days':
            result_dict.Glaze_Days = edict()
            ele_list.append('Glaze_Days')
            Glaze_Days_tab1, Glaze_Days_tab2, Glaze_Days_tab3 = table_stats_part1(monthly_df, 'Glaze_Days')
            result_list.append(OrderedDict(zip(['历年雨凇日统计', '累年雨凇日统计', '累年各月雨凇日统计'], [Glaze_Days_tab1, Glaze_Days_tab2, Glaze_Days_tab3])))

            result_dict.Glaze_Days['table1'] = Glaze_Days_tab1
            result_dict.Glaze_Days['table2'] = Glaze_Days_tab2
            result_dict.Glaze_Days['table3'] = Glaze_Days_tab3

            try:
                report_path = glaze_days_report(Glaze_Days_tab1, Glaze_Days_tab2, Glaze_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Glaze_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Glaze_Days['report'] = None

        elif ele == 'SoRi_Days':
            result_dict.SoRi_Days = edict()
            ele_list.append('SoRi_Days')
            SoRi_Days_tab1, SoRi_Days_tab2, SoRi_Days_tab3 = table_stats_part1(monthly_df, 'SoRi_Days')
            result_list.append(OrderedDict(zip(['历年雾凇日统计', '累年雾凇日统计', '累年各月雾凇日统计'], [SoRi_Days_tab1, SoRi_Days_tab2, SoRi_Days_tab3])))

            result_dict.SoRi_Days['table1'] = SoRi_Days_tab1
            result_dict.SoRi_Days['table2'] = SoRi_Days_tab2
            result_dict.SoRi_Days['table3'] = SoRi_Days_tab3

            try:
                report_path = sori_days_report(SoRi_Days_tab1, SoRi_Days_tab2, SoRi_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.SoRi_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.SoRi_Days['report'] = None

        elif ele == 'Tord_Days':
            result_dict.Tord_Days = edict()
            ele_list.append('Tord_Days')
            Tord_Days_tab1, Tord_Days_tab2, Tord_Days_tab3 = table_stats_part1(monthly_df, 'Tord_Days')
            result_list.append(OrderedDict(zip(['历年龙卷日统计', '累年龙卷日统计', '累年各月龙卷日统计'], [Tord_Days_tab1, Tord_Days_tab2, Tord_Days_tab3])))

            result_dict.Tord_Days['table1'] = Tord_Days_tab1
            result_dict.Tord_Days['table2'] = Tord_Days_tab2
            result_dict.Tord_Days['table3'] = Tord_Days_tab3

            try:
                report_path = tord_days_report(Tord_Days_tab1, Tord_Days_tab2, Tord_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Tord_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Tord_Days['report'] = None

        elif ele == 'SaSt_Days':
            result_dict.SaSt_Days = edict()
            ele_list.append('SaSt_Days')
            SaSt_Days_tab1, SaSt_Days_tab2, SaSt_Days_tab3 = table_stats_part1(monthly_df, 'SaSt_Days')
            result_list.append(OrderedDict(zip(['历年沙尘暴日统计', '累年沙尘暴日统计', '累年各月沙尘暴日统计'], [SaSt_Days_tab1, SaSt_Days_tab2, SaSt_Days_tab3])))

            result_dict.SaSt_Days['table1'] = SaSt_Days_tab1
            result_dict.SaSt_Days['table2'] = SaSt_Days_tab2
            result_dict.SaSt_Days['table3'] = SaSt_Days_tab3

            try:
                report_path = sast_days_report(SaSt_Days_tab1, SaSt_Days_tab2, SaSt_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.SaSt_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.SaSt_Days['report'] = None

        elif ele == 'FlSa_Days':
            result_dict.FlSa_Days = edict()
            ele_list.append('FlSa_Days')
            FlSa_Days_tab1, FlSa_Days_tab2, FlSa_Days_tab3 = table_stats_part1(monthly_df, 'FlSa_Days')
            result_list.append(OrderedDict(zip(['历年扬沙日统计', '累年扬沙日统计', '累年各月扬沙日统计'], [FlSa_Days_tab1, FlSa_Days_tab2, FlSa_Days_tab3])))

            result_dict.FlSa_Days['table1'] = FlSa_Days_tab1
            result_dict.FlSa_Days['table2'] = FlSa_Days_tab2
            result_dict.FlSa_Days['table3'] = FlSa_Days_tab3

            try:
                report_path = flsa_days_report(FlSa_Days_tab1, FlSa_Days_tab2, FlSa_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.FlSa_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.FlSa_Days['report'] = None

        elif ele == 'FlDu_Days':
            result_dict.FlDu_Days = edict()
            ele_list.append('FlDu_Days')
            FlDu_Days_tab1, FlDu_Days_tab2, FlDu_Days_tab3 = table_stats_part1(monthly_df, 'FlDu_Days')
            result_list.append(OrderedDict(zip(['历年浮尘日统计', '累年浮尘日统计', '累年各月浮尘日统计'], [FlDu_Days_tab1, FlDu_Days_tab2, FlDu_Days_tab3])))

            result_dict.FlDu_Days['table1'] = FlDu_Days_tab1
            result_dict.FlDu_Days['table2'] = FlDu_Days_tab2
            result_dict.FlDu_Days['table3'] = FlDu_Days_tab3

            try:
                report_path = fldu_days_report(FlDu_Days_tab1, FlDu_Days_tab2, FlDu_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.FlDu_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.FlDu_Days['report'] = None

        elif ele == 'Haze_Days':
            result_dict.Haze_Days = edict()
            ele_list.append('Haze_Days')
            Haze_Days_tab1, Haze_Days_tab2, Haze_Days_tab3 = table_stats_part1(monthly_df, 'Haze_Days')
            result_list.append(OrderedDict(zip(['历年霾日统计', '累年霾日统计', '累年各月霾日统计'], [Haze_Days_tab1, Haze_Days_tab2, Haze_Days_tab3])))

            result_dict.Haze_Days['table1'] = Haze_Days_tab1
            result_dict.Haze_Days['table2'] = Haze_Days_tab2
            result_dict.Haze_Days['table3'] = Haze_Days_tab3

            try:
                report_path = haze_days_report(Haze_Days_tab1, Haze_Days_tab2, Haze_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Haze_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Haze_Days['report'] = None

        elif ele == 'GaWIN_Days':
            result_dict.GaWIN_Days = edict()
            ele_list.append('GaWIN_Days')
            GaWIN_Days_tab1, GaWIN_Days_tab2, GaWIN_Days_tab3 = table_stats_part1(monthly_df, 'GaWIN_Days')
            result_list.append(OrderedDict(zip(['历年大风日统计', '累年大风日统计', '累年各月大风日统计'], [GaWIN_Days_tab1, GaWIN_Days_tab2, GaWIN_Days_tab3])))

            result_dict.GaWIN_Days['table1'] = GaWIN_Days_tab1
            result_dict.GaWIN_Days['table2'] = GaWIN_Days_tab2
            result_dict.GaWIN_Days['table3'] = GaWIN_Days_tab3

            try:
                report_path = gawin_days_report(GaWIN_Days_tab1, GaWIN_Days_tab2, GaWIN_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.GaWIN_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.GaWIN_Days['report'] = None
                
        # 电线结冰
        elif ele == 'EICE_Days':
            result_dict.EICE_Days = edict()
            ele_list.append('EICE_Days')
            EICE_Days_tab1, EICE_Days_tab2, EICE_Days_tab3 = table_stats_part1(monthly_df, 'EICE_Days')
            result_list.append(OrderedDict(zip(['历年电线结冰日统计', '累年电线结冰日统计', '累年各月电线结冰日统计'], [EICE_Days_tab1, EICE_Days_tab2, EICE_Days_tab3])))

            result_dict.EICE_Days['table1'] = EICE_Days_tab1
            result_dict.EICE_Days['table2'] = EICE_Days_tab2
            result_dict.EICE_Days['table3'] = EICE_Days_tab3

            try:
                report_path = eice_days_report(EICE_Days_tab1, EICE_Days_tab2, EICE_Days_tab3, monthly_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.EICE_Days['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.EICE_Days['report'] = None

        elif ele == 'Squa':
            result_dict.Squa = edict()
            ele_list.append('Squa')
            Squa_tab1, Squa_tab2, Squa_tab3 = table_stats_part1(daily_df, 'Squa')
            result_list.append(OrderedDict(zip(['历年飑日统计', '累年飑日统计', '累年各月飑日统计'], [Squa_tab1, Squa_tab2, Squa_tab3])))

            result_dict.Squa['table1'] = Squa_tab1
            result_dict.Squa['table2'] = Squa_tab2
            result_dict.Squa['table3'] = Squa_tab3

            try:
                report_path = squa_days_report(Squa_tab1, Squa_tab2, Squa_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Squa['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Squa['report'] = None

        elif ele == 'Lit':
            result_dict.Lit = edict()
            ele_list.append('Lit')
            Lit_tab1, Lit_tab2, Lit_tab3 = table_stats_part1(daily_df, 'Lit')
            result_list.append(OrderedDict(zip(['历年闪电日统计', '累年闪电日统计', '累年各月闪电日统计'], [Lit_tab1, Lit_tab2, Lit_tab3])))

            result_dict.Lit['table1'] = Lit_tab1
            result_dict.Lit['table2'] = Lit_tab2
            result_dict.Lit['table3'] = Lit_tab3

            try:
                report_path = lit_days_report(Lit_tab1, Lit_tab2, Lit_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Lit['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Lit['report'] = None

        elif ele == 'DuWhr':
            result_dict.DuWhr = edict()
            ele_list.append('DuWhr')
            DuWhr_tab1, DuWhr_tab2, DuWhr_tab3 = table_stats_part1(daily_df, 'DuWhr')
            result_list.append(OrderedDict(zip(['历年尘卷风日统计', '累年尘卷风日统计', '累年各月尘卷风日统计'], [DuWhr_tab1, DuWhr_tab2, DuWhr_tab3])))

            result_dict.DuWhr['table1'] = DuWhr_tab1
            result_dict.DuWhr['table2'] = DuWhr_tab2
            result_dict.DuWhr['table3'] = DuWhr_tab3

            try:
                report_path = duwhr_days_report(DuWhr_tab1, DuWhr_tab2, DuWhr_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.DuWhr['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.DuWhr['report'] = None

        elif ele == 'DrSnow':
            result_dict.DrSnow = edict()
            ele_list.append('DrSnow')
            DrSnow_tab1, DrSnow_tab2, DrSnow_tab3 = table_stats_part1(daily_df, 'DrSnow')
            result_list.append(OrderedDict(zip(['历年吹雪日统计', '累年吹雪日统计', '累年各月吹雪日统计'], [DrSnow_tab1, DrSnow_tab2, DrSnow_tab3])))

            result_dict.DrSnow['table1'] = DrSnow_tab1
            result_dict.DrSnow['table2'] = DrSnow_tab2
            result_dict.DrSnow['table3'] = DrSnow_tab3

            # try:
            report_path = drsnow_days_report(DrSnow_tab1, DrSnow_tab2, DrSnow_tab3, daily_df, data_dir, main_sta_ids)
            result_path.append(report_path)
            report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
            result_dict.DrSnow['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            # except:
            #     result_dict.DrSnow['report'] = None

        elif ele == 'Snow':
            result_dict.Snow = edict()
            ele_list.append('Snow')
            Snow_tab1, Snow_tab2, Snow_tab3 = table_stats_part1(daily_df, 'Snow')
            result_list.append(OrderedDict(zip(['历年雪日统计', '累年雪日统计', '累年各月雪日统计'], [Snow_tab1, Snow_tab2, Snow_tab3])))

            result_dict.Snow['table1'] = Snow_tab1
            result_dict.Snow['table2'] = Snow_tab2
            result_dict.Snow['table3'] = Snow_tab3

            try:
                report_path = snow_days_report(Snow_tab1, Snow_tab2, Snow_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Snow['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Snow['report'] = None

        elif ele == 'Frost':
            result_dict.Frost = edict()
            ele_list.append('Frost')
            Frost_tab1, Frost_tab2, Frost_tab3 = table_stats_part1(daily_df, 'Frost')
            result_list.append(OrderedDict(zip(['历年霜日统计', '累年霜日统计', '累年各月霜日统计'], [Frost_tab1, Frost_tab2, Frost_tab3])))

            result_dict.Frost['table1'] = Frost_tab1
            result_dict.Frost['table2'] = Frost_tab2
            result_dict.Frost['table3'] = Frost_tab3

            try:
                report_path = frost_days_report(Frost_tab1, Frost_tab2, Frost_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Frost['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Frost['report'] = None

        elif ele == 'GSS':
            result_dict.GSS = edict()
            ele_list.append('GSS')
            GSS_tab1, GSS_tab2, GSS_tab3 = table_stats_part1(daily_df, 'GSS')
            result_list.append(OrderedDict(zip(['历年积雪日统计', '累年积雪日统计', '累年各月积雪日统计'], [GSS_tab1, GSS_tab2, GSS_tab3])))

            result_dict.GSS['table1'] = GSS_tab1
            result_dict.GSS['table2'] = GSS_tab2
            result_dict.GSS['table3'] = GSS_tab3

            try:
                report_path = gss_days_report(GSS_tab1, GSS_tab2, GSS_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.GSS['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.GSS['report'] = None

        elif ele == 'ICE':
            result_dict.ICE = edict()
            ele_list.append('ICE')
            ICE_tab1, ICE_tab2, ICE_tab3 = table_stats_part1(daily_df, 'ICE')
            result_list.append(OrderedDict(zip(['历年结冰日统计', '累年结冰日统计', '累年各月结冰日统计'], [ICE_tab1, ICE_tab2, ICE_tab3])))

            result_dict.ICE['table1'] = ICE_tab1
            result_dict.ICE['table2'] = ICE_tab2
            result_dict.ICE['table3'] = ICE_tab3

            try:
                report_path = ice_days_report(ICE_tab1, ICE_tab2, ICE_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.ICE['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.ICE['report'] = None

        elif ele == 'Thund':
            result_dict.Thund = edict()
            ele_list.append('Thund')
            Thund_tab1, Thund_tab2, Thund_tab3 = table_stats_part1(daily_df, 'Thund')
            result_list.append(OrderedDict(zip(['历年雷暴日统计', '累年雷暴日统计', '累年各月雷暴日统计'], [Thund_tab1, Thund_tab2, Thund_tab3])))

            result_dict.Thund['table1'] = Thund_tab1
            result_dict.Thund['table2'] = Thund_tab2
            result_dict.Thund['table3'] = Thund_tab3

            try:
                report_path = thund_days_report(Thund_tab1, Thund_tab2, Thund_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Thund['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Thund['report'] = None
                
        elif ele == 'TEM_Max':
            result_dict.tem_max = edict()
            ele_list.append('TEM_Max')
            TEM_Max_tab1, TEM_Max_tab2, TEM_Max_tab3 = table_stats_part1(daily_df, 'TEM_Max')
            result_list.append(OrderedDict(zip(['历年高温日统计', '累年高温日统计', '累年各月高温日统计'], [TEM_Max_tab1, TEM_Max_tab2, TEM_Max_tab3])))

            result_dict.tem_max['table1'] = TEM_Max_tab1
            result_dict.tem_max['table2'] = TEM_Max_tab2
            result_dict.tem_max['table3'] = TEM_Max_tab3

            try:
                report_path = tem_max_days_report(TEM_Max_tab1, TEM_Max_tab2, TEM_Max_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.tem_max['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.tem_max['report'] = None
                
        elif ele == 'TEM_Min':
            result_dict.tem_min = edict()
            ele_list.append('TEM_Min')
            TEM_Min_tab1, TEM_Min_tab2, TEM_Min_tab3 = table_stats_part1(daily_df, 'TEM_Min')
            result_list.append(OrderedDict(zip(['历年低温日统计', '累年低温日统计', '累年各月低温日统计'], [TEM_Min_tab1, TEM_Min_tab2, TEM_Min_tab3])))

            result_dict.tem_min['table1'] = TEM_Min_tab1
            result_dict.tem_min['table2'] = TEM_Min_tab2
            result_dict.tem_min['table3'] = TEM_Min_tab3

            try:
                report_path =tem_min_days_report(TEM_Min_tab1, TEM_Min_tab2, TEM_Min_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.tem_min['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.tem_min['report'] = None                

        elif ele == 'PRE_Time_2020':
            result_dict.pre_strong = edict()
            ele_list.append('PRE_Time_2020')
            PRE_Time_2020_tab1, PRE_Time_2020_tab2, PRE_Time_2020_tab3 = table_stats_part1(daily_df, 'PRE_Time_2020')
            result_list.append(OrderedDict(zip(['历年强降雨日统计', '累年强降雨日统计', '累年各月强降雨日统计'], [PRE_Time_2020_tab1, PRE_Time_2020_tab2, PRE_Time_2020_tab3])))

            result_dict.pre_strong['table1'] = PRE_Time_2020_tab1
            result_dict.pre_strong['table2'] = PRE_Time_2020_tab2
            result_dict.pre_strong['table3'] = PRE_Time_2020_tab3

            try:
                report_path = pre_strong_report(PRE_Time_2020_tab1, PRE_Time_2020_tab2, PRE_Time_2020_tab3, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.pre_strong['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.pre_strong['report'] = None   
                
    if len(result_path) == 0:
        result_dict['report'] = None
    else:
        try:
            new_docx_path = os.path.join(data_dir, 'weather_phenomena_days.docx')
            master = Document(result_path[0])
            middle_new_docx = Composer(master)
            for word in result_path[1:]:  # 从第二个文档开始追加
                word_document = Document(word)
                middle_new_docx.append(word_document)
            middle_new_docx.save(new_docx_path)
            new_docx_path = new_docx_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
            result_dict['report'] = new_docx_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)

        except Exception as e:
            print(f"发生错误：{e}")
            result_dict['report'] = None

    # 总表1
    if daily_df is not None:
        main_station_name = daily_df[daily_df['Station_Id_C'] == main_sta_ids]['Station_Name'][0]
    else:
        main_station_name = monthly_df[monthly_df['Station_Id_C'] == main_sta_ids]['Station_Name'][0]

    if len(ele_list) != 0:
        result_dict.total_table = edict()
        all_accum_result = all_weather_statistics_accum(monthly_df, daily_df, ele_list, main_station_name)
        result_list.append(OrderedDict(zip(['总表1'], [all_accum_result])))
        result_dict.total_table['总表1'] = all_accum_result

    # 7.结果保存
    if cfg.INFO.SAVE_RESULT:
        result_dict['csv'] = save_cmadaas_data(data_dir, mon_data=monthly_df, day_data=daily_df)

    return result_dict


def init_and_end_days(data_json):
    '''
    天气现象初终日统计组件
    '''
    # 1.参数读取
    years = data_json['years']  # 选择的数据年份
    main_sta_ids = data_json['main_sta_ids']  # 主站
    sub_sta_ids = data_json.get('sub_sta_ids')  # 对比站
    elements = data_json['elements']  # 选择的气象要素

    # 2.参数处理
    uuid4 = uuid.uuid4().hex
    data_dir = os.path.join(cfg.INFO.IN_DATA_DIR, uuid4)
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        os.chmod(data_dir, 0o007 | 0o070 | 0o700)

    if isinstance(main_sta_ids, int):
        main_sta_ids = str(main_sta_ids)

    if isinstance(sub_sta_ids, list):
        sub_sta_ids = [str(ids) for ids in sub_sta_ids]
        sub_sta_ids = ','.join(sub_sta_ids)

    elif isinstance(sub_sta_ids, int):  # 只有单个值的情况
        sub_sta_ids = str(sub_sta_ids)

    if 'Frost' in elements:
        elements.remove('Frost')
        elements.insert(0, 'Frost')

    # 3.拼接需要下载的参数
    daily_elements = ''
    daily_df = None

    for ele in elements:
        if ele == 'Snow':
            daily_elements += 'Snow,'

        elif ele == 'Frost':
            daily_elements += 'Frost,'

        elif ele == 'GSS':
            daily_elements += 'GSS,'

        elif ele == 'ICE':
            daily_elements += 'ICE,'

        elif ele == 'Thund':
            daily_elements += 'Thund,'

        elif ele == 'GST_Under0':
            daily_elements += 'GST_Min,'

        elif ele == 'Tem_Under0':
            daily_elements += 'TEM_Min,'

    # 4.数据获取
    day_ele_list = ['Snow', 'Frost', 'GSS', 'ICE', 'Thund', 'GST_Min', 'TEM_Min']

    if cfg.INFO.READ_LOCAL:
        daily_df = pd.read_csv(cfg.FILES.QH_DATA_DAY)

        if sub_sta_ids == None:
            sta_ids = main_sta_ids
        else:
            sta_ids = main_sta_ids + ',' + sub_sta_ids

        day_eles = ('Station_Id_C,Station_Name,Lat,Lon,Datetime,Year,Mon,Day,' + daily_elements[:-1]).split(',')
        daily_df = get_local_data(daily_df, sta_ids, day_eles, years, 'Day')       
        # daily_df = daily_df[~daily_df['Station_Name'].isna()]  # 测试数据问题 去掉station_name为null的行

    else:
        # 天擎数据下载 and 数据前处理
        try:
            if sub_sta_ids is not None:
                sta_ids = main_sta_ids + ',' + sub_sta_ids
            else:
                sta_ids = main_sta_ids

            if len(set(day_ele_list) & set(elements)) != 0:
                daily_df = get_cmadaas_daily_data(years, daily_elements, sta_ids)
                daily_df = daily_data_processing(daily_df, years)
            else:
                daily_df = None

        except Exception as e:
            raise Exception('天擎数据下载或处理失败')

    # 5.计算之前先检测数据完整率 check H小时 D天 MS月 YS年
    years = years.split(',')
    result_dict = edict()
    result_dict['uuid'] = uuid4
    result_dict.check_result = edict()
    daily_elements = daily_elements[:-1]

    if daily_df is not None and len(daily_df) != 0:
        checker = check(daily_df, 'D', daily_elements.split(','), sta_ids.split(','), years[0], years[1])
        check_result = checker.run()
        result_dict.check_result['使用的天擎日要素'] = check_result

    # 6.结果生成
    accum_ele_list = []  # 记录出现的要素名称，用于生成初终日总表(总表2)
    table5_list = []  # 记录出现的要素统计结果，用于生成初终日总表(总表2)
    result_list = []  # 用于将保存的csv，进行url转换，然后输出路径
    result_path = []
    for ele in elements:
        if ele == 'Snow':
            result_dict.Snow = edict()
            Snow_tab4, Snow_tab5 = table_stats_part2(daily_df, 'Snow')
            result_list.append(OrderedDict(zip(['历年雪的初终间日', '累年雪的初终间日'], [Snow_tab4, Snow_tab5])))

            result_dict.Snow['table4'] = Snow_tab4
            result_dict.Snow['table5'] = Snow_tab5

            if Snow_tab5 is not None:
                accum_ele_list.append('雪')
                table5_list.append(pd.DataFrame(Snow_tab5))

            try:
                report_path = snow_init_and_end_days_report(Snow_tab4, Snow_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Snow['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Snow['report'] = None

        elif ele == 'Frost':
            result_dict.Frost = edict()
            Frost_tab4, Frost_tab5 = table_stats_part2(daily_df, 'Frost')
            result_list.append(OrderedDict(zip(['历年霜的初终间日', '累年霜的初终间日'], [Frost_tab4, Frost_tab5])))

            result_dict.Frost['table4'] = Frost_tab4
            result_dict.Frost['table5'] = Frost_tab5

            try:
                report_path = frost_init_and_end_days_report(Frost_tab4, Frost_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Frost['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Frost['report'] = None

            if Frost_tab5 is not None:
                accum_ele_list.append('霜')
                table5_list.append(pd.DataFrame(Frost_tab5))

        elif ele == 'GSS':
            result_dict.GSS = edict()
            GSS_tab4, GSS_tab5 = table_stats_part2(daily_df, 'GSS')
            result_list.append(OrderedDict(zip(['历年积雪的初终间日', '累年积雪的初终间日'], [GSS_tab4, GSS_tab5])))

            result_dict.GSS['table4'] = GSS_tab4
            result_dict.GSS['table5'] = GSS_tab5

            if GSS_tab5 is not None:
                accum_ele_list.append('积雪')
                table5_list.append(pd.DataFrame(GSS_tab5))

            try:
                report_path = gss_init_and_end_days_report(GSS_tab4, GSS_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.GSS['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.GSS['report'] = None

        elif ele == 'ICE':
            result_dict.ICE = edict()
            ICE_tab4, ICE_tab5 = table_stats_part2(daily_df, 'ICE')
            result_list.append(OrderedDict(zip(['历年结冰的初终间日', '累年结冰的初终间日'], [ICE_tab4, ICE_tab5])))

            result_dict.ICE['table4'] = ICE_tab4
            result_dict.ICE['table5'] = ICE_tab5

            if ICE_tab5 is not None:
                accum_ele_list.append('结冰')
                table5_list.append(pd.DataFrame(ICE_tab5))

            try:
                report_path = ice_init_and_end_days_report(ICE_tab4, ICE_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.ICE['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.ICE['report'] = None

        elif ele == 'Thund':
            result_dict.Thund = edict()
            Thund_tab4, Thund_tab5 = table_stats_part2(daily_df, 'Thund')
            result_list.append(OrderedDict(zip(['历年雷暴的初终间日', '累年雷暴的初终间日'], [Thund_tab4, Thund_tab5])))

            result_dict.Thund['table4'] = Thund_tab4
            result_dict.Thund['table5'] = Thund_tab5

            if Thund_tab5 is not None:
                accum_ele_list.append('雷暴')
                table5_list.append(pd.DataFrame(Thund_tab5))

            try:
                report_path = thund_init_and_end_days_report(Thund_tab4, Thund_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Thund['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Thund['report'] = None

        elif ele == 'Tem_Under0':
            result_dict.Tem_Under0 = edict()
            daily_df['TEM_under'] = np.where(daily_df['TEM_Min'].values <= 0, 1, 0)
            Tem_Under0_tab4, Tem_Under0_tab5 = table_stats_part2(daily_df, 'TEM_under')
            result_list.append(OrderedDict(zip(['历年最低气温小于等于0度的初终间日', '累年最低气温小于等于0度的初终间日'], [Tem_Under0_tab4, Tem_Under0_tab5])))

            result_dict.Tem_Under0['table4'] = Tem_Under0_tab4
            result_dict.Tem_Under0['table5'] = Tem_Under0_tab5

            if Tem_Under0_tab5 is not None:
                accum_ele_list.append('最低气温小于等于0度')
                table5_list.append(pd.DataFrame(Tem_Under0_tab5))

            try:
                report_path = tem_under0_init_and_end_days_report(Tem_Under0_tab4, Tem_Under0_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Tem_Under0['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Tem_Under0['report'] = None

        elif ele == 'GST_Under0':
            result_dict.GST_Under0 = edict()
            daily_df['GST_under'] = np.where(daily_df['GST_Min'].values <= 0, 1, 0)
            GST_under_Under0_tab4, GST_under_Under0_tab5 = table_stats_part2(daily_df, 'GST_under')
            result_list.append(OrderedDict(zip(['历年地面最低气温小于等于0度的初终间日', '累年地面最低气温小于等于0度的初终间日'], [GST_under_Under0_tab4, GST_under_Under0_tab5])))

            result_dict.GST_Under0['table4'] = GST_under_Under0_tab4
            result_dict.GST_Under0['table5'] = GST_under_Under0_tab5

            if GST_under_Under0_tab5 is not None:
                accum_ele_list.append('地面最低气温小于等于0度')
                table5_list.append(pd.DataFrame(GST_under_Under0_tab5))

            try:
                report_path = gst_under0_init_and_end_days_report(GST_under_Under0_tab4, GST_under_Under0_tab5, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.GST_Under0['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except Exception as e:
                logging.exception(e)
                result_dict.GST_Under0['report'] = None

    if len(result_path) == 0:
        result_dict['report'] = None
    else:
        try:
            new_docx_path = os.path.join(data_dir, 'init_and_end_days.docx')
            master = Document(result_path[0])
            middle_new_docx = Composer(master)
            for word in result_path[1:]:  # 从第二个文档开始追加
                word_document = Document(word)
                middle_new_docx.append(word_document)
            middle_new_docx.save(new_docx_path)
            new_docx_path = new_docx_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
            result_dict['report'] = new_docx_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)

        except Exception as e:
            print(f"发生错误：{e}")
            result_dict['report'] = None

    # 生成主站号对应的主站名称，用于生成总表1和总表2
    if daily_df is not None:
        main_station_name = daily_df[daily_df['Station_Id_C'] == main_sta_ids]['Station_Name'][0]

    # 总表2
    require_data_list = list(zip(accum_ele_list, table5_list))
    if len(require_data_list) != 0:
        result_dict.total_table = edict()
        for i, require in enumerate(require_data_list):
            name = require[0]
            table = require[1]
            cols = table.columns[table.columns.str.contains(main_station_name)].tolist()
            row = table[cols].to_numpy().T.reshape(1, -1)
            df = pd.DataFrame(row)

            if name == '霜':
                df.columns = ['平均初日', '最早初日', '最晚初日', '平均终日', '最早终日', '最晚终日', '平均间日', '最小间日', '最大间日', '平均无霜日', '最小无霜日', '最大无霜日']

            else:
                df.columns = ['平均初日', '最早初日', '最晚初日', '平均终日', '最早终日', '最晚终日', '平均间日', '最小间日', '最大间日']

            if i == 0:
                total = df

            else:
                total = pd.concat([total, df], axis=0)

        total.insert(loc=0, column='名称', value=accum_ele_list)
        total.reset_index(drop=True, inplace=True)
        result_dict.total_table['总表2'] = total.to_dict(orient='records')
        result_list.append(OrderedDict(zip(['总表2'], [total])))

    # 7.结果保存
    if cfg.INFO.SAVE_RESULT:
        result_dict['csv'] = save_cmadaas_data(data_dir, day_data=daily_df)

    return result_dict


def weather_process_stats(data_json):
    '''
    天气过程（寒潮、低温连阴雨、低温冰冻）统计组件
    '''
    # 1.参数读取
    years = data_json['years']  # 选择的数据年份
    main_sta_ids = data_json['main_sta_ids']  # 主站
    sub_sta_ids = data_json.get('sub_sta_ids')  # 对比站
    elements = data_json['elements']  # 选择的气象要素

    # 2.参数处理
    uuid4 = uuid.uuid4().hex
    data_dir = os.path.join(cfg.INFO.IN_DATA_DIR, uuid4)
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        os.chmod(data_dir, 0o007 | 0o070 | 0o700)

    if isinstance(main_sta_ids, int):
        main_sta_ids = str(main_sta_ids)

    if isinstance(sub_sta_ids, list):
        sub_sta_ids = [str(ids) for ids in sub_sta_ids]
        sub_sta_ids = ','.join(sub_sta_ids)

    elif isinstance(sub_sta_ids, int):  # 只有单个值的情况
        sub_sta_ids = str(sub_sta_ids)

    # 3.拼接需要下载的参数
    daily_elements = 'TEM_Avg,TEM_Min,'
    daily_df = None

    for ele in elements:
        if ele == 'Cold_Wave':
            daily_elements += 'WIN_S_Max,WIN_D_S_Max,'

        elif ele == 'Low_TEM':
            daily_elements += 'SSH,PRE_Time_2020,'

    # 4.数据获取
    if cfg.INFO.READ_LOCAL:
        daily_df = pd.read_csv(cfg.FILES.QH_DATA_DAY)

        if sub_sta_ids == None:
            sta_ids = main_sta_ids
        else:
            sta_ids = main_sta_ids + ',' + sub_sta_ids
            
        day_eles = ('Station_Id_C,Station_Name,Lat,Lon,Datetime,Year,Mon,Day,' + daily_elements[:-1]).split(',')
        daily_df = get_local_data(daily_df, sta_ids, day_eles, years, 'Day')
        # daily_df = daily_df[~daily_df['Station_Name'].isna()]  # 测试数据问题 去掉station_name为null的行

    else:
        # 天擎数据下载 and 数据前处理
        try:
            if sub_sta_ids is not None:
                sta_ids = main_sta_ids + ',' + sub_sta_ids
            else:
                sta_ids = main_sta_ids

            daily_df = get_cmadaas_daily_data(years, daily_elements, sta_ids)
            daily_df = daily_data_processing(daily_df, years)

        except Exception as e:
            raise Exception('天擎数据下载或处理失败')

    # 5.计算之前先检测数据完整率 check H小时 D天 MS月 YS年
    years = years.split(',')
    result_dict = edict()
    result_dict['uuid'] = uuid4
    result_dict.check_result = edict()
    daily_elements = daily_elements[:-1]

    if daily_df is not None and len(daily_df) != 0:
        checker = check(daily_df, 'D', daily_elements.split(','), sta_ids.split(','), years[0], years[1])
        check_result = checker.run()
        result_dict.check_result['使用的天擎日要素'] = check_result

    # 6.结果生成
    result_list = []
    result_path = []
    for ele in elements:
        if ele == 'Cold_Wave':
            result_dict.Cold_Wave = dict()
            cold_wave_result, cold_wave_wind, cold_wave_wind_d = cold_wave_statistics(daily_df)
            result_list.append(OrderedDict(zip(['寒潮过程统计', '寒潮大风统计', '寒潮大风风向频数统计'], [cold_wave_result, cold_wave_wind, cold_wave_wind_d])))

            result_dict.Cold_Wave['寒潮'] = cold_wave_result
            result_dict.Cold_Wave['寒潮大风'] = cold_wave_wind
            result_dict.Cold_Wave['寒潮风向频数'] = cold_wave_wind_d

            try:
                report_path = cold_wave_report(cold_wave_result, cold_wave_wind, cold_wave_wind_d, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Cold_Wave['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Cold_Wave['report'] = None

        elif ele == 'Low_TEM':
            result_dict.Low_TEM = edict()
            cold_rainy = cold_rainy_days_statistics(daily_df)
            cold_freezing = cold_freeing_days_statistics(daily_df)
            result_list.append(OrderedDict(zip(['低温连阴雨统计', '低温冰冻统计'], [cold_rainy, cold_freezing])))

            result_dict.Low_TEM['低温连阴雨'] = cold_rainy
            result_dict.Low_TEM['低温冰冻'] = cold_freezing

            try:
                report_path = cold_rainy_freez_report(cold_rainy, cold_freezing, daily_df, data_dir, main_sta_ids)
                result_path.append(report_path)
                report_path = report_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
                result_dict.Low_TEM['report'] = report_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)
            except:
                result_dict.Low_TEM['report'] = None

    if len(result_path) == 0:
        result_dict['report'] = None
    else:
        try:
            new_docx_path = os.path.join(data_dir, 'weather_process_stats.docx')
            master = Document(result_path[0])
            middle_new_docx = Composer(master)
            for word in result_path[1:]:  # 从第二个文档开始追加
                word_document = Document(word)
                middle_new_docx.append(word_document)
            middle_new_docx.save(new_docx_path)
            new_docx_path = new_docx_path.replace(cfg.INFO.IN_DATA_DIR, cfg.INFO.OUT_DATA_DIR)
            result_dict['report'] = new_docx_path.replace(cfg.INFO.OUT_DATA_DIR, cfg.INFO.OUT_DATA_URL)

        except Exception as e:
            print(f"发生错误：{e}")
            result_dict['report'] = None

    # 7.结果保存
    if cfg.INFO.SAVE_RESULT:
        result_dict['csv'] = save_cmadaas_data(data_dir, day_data=daily_df)

    return result_dict

if __name__=='__main__':
    years="1970,2023"
    sub_sta_ids=["56151"]
    main_sta_ids="56067"
    elements=["Thund","GSS","TEM_Min","TEM_Max","PRE_Time_2020"]